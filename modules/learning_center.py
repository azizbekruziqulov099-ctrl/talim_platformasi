"""Tenant-safe learning/repetitorlik centre API.

The guided avatar is deliberately a UI and draft assistant.  It cannot publish
teaching content, enrol a learner, assign roles, score an exam, approve payroll
or record money.  Those actions remain authenticated, permission checked and,
where irreversible, explicitly human-confirmed.
"""

from __future__ import annotations

import json
import io
import re
from contextlib import contextmanager
from datetime import date, datetime, time, timedelta, timezone
from decimal import Decimal
from typing import Any, Callable, Iterator, Literal
from zoneinfo import ZoneInfo

import psycopg2
from fastapi import APIRouter, Depends, Header, HTTPException, Query
from fastapi.responses import JSONResponse, Response
from pydantic import BaseModel, Field

from platform_core.database import DatabaseBusyError, db_session


ROLE_LABELS = {
    "owner": "Mulkdor",
    "founder": "Ta'sischi",
    "director": "Direktor",
    "administrator": "Administrator",
    "academic_manager": "O'quv ishlari rahbari",
    "receptionist": "Qabulxona xodimi",
    "accountant": "Hisobchi",
    "methodist": "Metodist",
    "teacher": "O'qituvchi/repetitor",
    "student": "O'quvchi",
    "parent": "Ota-ona",
}
MANAGER_ROLES = {
    "owner", "founder", "director", "administrator", "academic_manager",
}
ACADEMIC_ROLES = MANAGER_ROLES | {"methodist", "teacher"}
ENROLLMENT_ROLES = {
    "owner", "founder", "director", "administrator", "academic_manager",
    "receptionist",
}
SCHEDULE_ROLES = ENROLLMENT_ROLES | {"methodist", "teacher"}
FINANCE_WRITE_ROLES = {"owner", "founder", "director", "accountant"}
FINANCE_VIEW_ROLES = FINANCE_WRITE_ROLES | {
    "administrator", "receptionist", "student", "parent",
}
FINANCE_VIEW_STAFF_ROLES = FINANCE_VIEW_ROLES - {"student", "parent"}
ROLE_GRANT_MATRIX = {
    "owner": {
        "director", "administrator", "academic_manager", "receptionist",
        "accountant", "methodist", "teacher",
    },
    "founder": {
        "director", "administrator", "academic_manager", "receptionist",
        "accountant", "methodist", "teacher",
    },
    "director": {
        "administrator", "academic_manager", "receptionist", "accountant",
        "methodist", "teacher",
    },
    "administrator": {"receptionist", "methodist", "teacher"},
    "academic_manager": {"methodist", "teacher"},
}
VIEW_ROLES = set(ROLE_LABELS)
STAFF_ROLES = VIEW_ROLES - {"student", "parent"}
PRIVILEGED_ROLES = {
    "owner", "founder", "director", "administrator", "academic_manager",
    "accountant",
}

ALL_PERMISSIONS = {
    "center.view", "branches.manage", "rooms.manage", "subjects.manage",
    "staff.manage", "courses.manage", "enrollments.manage", "schedule.manage",
    "attendance.write", "grades.write", "plans.write", "homework.write",
    "assessments.write", "billing.view", "billing.manage", "workload.view",
    "assistant.use",
}
ROLE_PERMISSIONS = {
    "owner": ALL_PERMISSIONS,
    "founder": ALL_PERMISSIONS,
    "director": ALL_PERMISSIONS,
    "administrator": ALL_PERMISSIONS - {"billing.manage"},
    "academic_manager": {
        "center.view", "subjects.manage", "staff.manage", "courses.manage",
        "enrollments.manage", "schedule.manage", "attendance.write",
        "grades.write", "plans.write", "homework.write", "assessments.write",
        "workload.view", "assistant.use",
    },
    "receptionist": {
        "center.view", "enrollments.manage", "schedule.manage", "billing.view",
        "assistant.use",
    },
    "accountant": {
        "center.view", "billing.view", "billing.manage", "workload.view",
        "assistant.use",
    },
    "methodist": {
        "center.view", "courses.manage", "schedule.manage", "attendance.write",
        "grades.write", "plans.write", "homework.write", "assessments.write",
        "workload.view", "assistant.use",
    },
    "teacher": {
        "center.view", "courses.manage", "schedule.manage", "attendance.write",
        "grades.write", "plans.write", "homework.write", "assessments.write",
        "workload.view", "assistant.use",
    },
    "student": {"center.view", "billing.view", "assistant.use"},
    "parent": {"center.view", "billing.view", "assistant.use"},
}
ASSISTANT_ACTIONS = {
    "SHOW_MENU", "FOCUS_FIELD", "SET_DRAFT_VALUE", "NEXT_STEP",
    "PREVIOUS_STEP", "PAUSE", "RESUME", "UNDO", "MINIMIZE",
    "RESTORE", "SPEAK", "COMPLETE_TOUR",
}
REVERSIBLE_ACTIONS = {
    "FOCUS_FIELD", "SET_DRAFT_VALUE", "NEXT_STEP", "PREVIOUS_STEP",
    "MINIMIZE", "RESTORE",
}
ONBOARDING_STEPS = (
    "identity", "branches", "rooms", "subjects", "staff", "courses",
    "billing", "preview",
)
CEFR_LEVELS = ("A1", "A2", "B1", "B2", "C1", "C2")
IELTS_COMPONENTS = ("listening", "reading", "writing", "speaking")
BUSINESS_TIMEZONE = ZoneInfo("Asia/Tashkent")


def business_today() -> date:
    return datetime.now(BUSINESS_TIMEZONE).date()


class DraftStart(BaseModel):
    relationship: Literal[
        "owner", "founder", "director", "administrator", "teacher"
    ]
    ownership_type: Literal["public", "private"]
    operator_model: Literal["center", "independent_tutor"] = "center"
    setup_mode: Literal["manual", "guided", "assistant"] = "guided"


class DraftPatch(BaseModel):
    step: str = Field(min_length=1, max_length=80)
    payload: dict[str, Any] = Field(default_factory=dict)
    expected_version: int = Field(ge=1)


class HumanConfirm(BaseModel):
    expected_version: int | None = Field(default=None, ge=1)
    confirmation: bool


class VerificationDecision(BaseModel):
    decision: Literal["verified", "rejected"]
    note: str | None = Field(default=None, max_length=2000)
    confirmation: bool = False


class BranchCreate(BaseModel):
    context_id: int
    name: str = Field(min_length=2, max_length=180)
    address: str | None = Field(default=None, max_length=500)
    phone: str | None = Field(default=None, max_length=80)
    work_start: str | None = None
    work_end: str | None = None
    work_days: list[int] = Field(
        default_factory=lambda: [1, 2, 3, 4, 5, 6],
        min_length=1,
        max_length=7,
    )


class RoomCreate(BaseModel):
    context_id: int
    branch_id: int
    name: str = Field(min_length=1, max_length=120)
    room_type: Literal[
        "classroom", "laboratory", "computer", "language", "meeting",
        "online", "other",
    ] = "classroom"
    capacity: int | None = Field(default=None, ge=1, le=5000)
    metadata: dict[str, Any] = Field(default_factory=dict)


class SubjectCreate(BaseModel):
    context_id: int
    code: str = Field(min_length=1, max_length=40)
    name: str = Field(min_length=2, max_length=160)
    supports_latex: bool = False
    formula_metadata: dict[str, Any] = Field(default_factory=dict)


class StaffAssign(BaseModel):
    context_id: int
    user_id: int
    role_key: str
    branch_id: int | None = None
    subject_ids: list[int] = Field(default_factory=list, max_length=100)
    employment_type: Literal[
        "full_time", "part_time", "contract", "hourly"
    ] | None = None
    weekly_capacity_hours: int | None = Field(default=None, ge=1, le=80)
    confirmation: bool = False


class StaffStatusChange(BaseModel):
    context_id: int
    status: Literal["suspended", "ended"]
    confirmation: bool = False


class AvailabilityRow(BaseModel):
    weekday: int = Field(ge=1, le=7)
    starts_at: str
    ends_at: str
    availability: Literal["available", "preferred", "unavailable"]
    effective_from: date | None = None
    effective_to: date | None = None
    note: str | None = Field(default=None, max_length=500)


class AvailabilityPut(BaseModel):
    context_id: int
    rows: list[AvailabilityRow] = Field(default_factory=list, max_length=200)


class CourseCreate(BaseModel):
    context_id: int
    branch_id: int | None = None
    subject_id: int
    teacher_user_id: int | None = None
    name: str = Field(min_length=2, max_length=180)
    course_type: Literal[
        "group", "individual", "intensive", "club", "exam_prep"
    ]
    delivery_mode: Literal["offline", "online_live", "hybrid"] = "offline"
    target_framework: Literal[
        "general", "custom", "cefr", "ielts", "national_exam", "school", "other"
    ] = "custom"
    cefr_level: str | None = None
    level_from: str | None = None
    level_to: str | None = None
    target_score: Decimal | None = None
    target_components: dict[str, Decimal] = Field(default_factory=dict)
    ielts_targets: dict[str, Decimal] | None = None
    ielts_test_type: Literal["academic", "general"] | None = None
    level_label: str | None = Field(default=None, max_length=80)
    monthly_price: Decimal | None = Field(default=None, ge=0)
    sessions_per_week: int | None = Field(default=None, ge=1, le=7)
    weekdays: list[int] = Field(default_factory=list, max_length=7)
    starts_at: str | None = None
    start_date: date | None = None
    end_date: date | None = None
    capacity: int = Field(default=20, ge=1, le=5000)
    duration_minutes: int = Field(default=90, ge=15, le=480)
    status: Literal["draft", "active"] = "draft"
    metadata: dict[str, Any] = Field(default_factory=dict)


class CourseActivation(BaseModel):
    context_id: int
    teacher_user_id: int | None = Field(default=None, ge=1)
    confirmation: bool = False


class EnrollmentCreate(BaseModel):
    context_id: int
    course_id: int
    student_user_id: int | None = None
    requested_status: Literal["pending", "waitlisted", "active"] = "pending"
    entry_status: str | None = Field(default=None, max_length=30)
    note: str | None = Field(default=None, max_length=1000)
    notes: str | None = Field(default=None, max_length=1000)
    start_date: date | None = None
    confirmation: bool = False


class EnrollmentDecision(BaseModel):
    status: Literal[
        "active", "waitlisted", "rejected", "paused", "withdrawn"
    ] | None = None
    decision: Literal["approve", "waitlist", "reject", "pause", "withdraw"] | None = None
    confirmation: bool


class ParentLinkCreate(BaseModel):
    context_id: int
    parent_user_id: int
    student_user_id: int
    branch_id: int | None = Field(default=None, ge=1)
    confirmation: bool = False


class ParentLinkRevoke(BaseModel):
    context_id: int
    confirmation: bool = False


class ScheduleCreate(BaseModel):
    context_id: int
    course_id: int
    teacher_user_id: int | None = None
    room_id: int | None = None
    schedule_kind: Literal["weekly", "dated"] = "dated"
    weekday: int | None = Field(default=None, ge=1, le=7)
    lesson_date: date | None = None
    effective_from: date | None = None
    effective_to: date | None = None
    starts_at: str
    ends_at: str | None = None
    duration_minutes: int | None = Field(default=None, ge=15, le=480)
    topic: str | None = Field(default=None, max_length=240)
    status: Literal["draft", "published"] = "draft"
    confirmation: bool = False


class AttendanceMark(BaseModel):
    context_id: int
    course_id: int
    student_user_id: int
    lesson_date: date | None = None
    attendance_date: date | None = None
    schedule_slot_id: int | None = None
    status: Literal[
        "present", "absent", "late", "excused", "absent_excused",
        "absent_unexcused", "sick"
    ]
    note: str | None = Field(default=None, max_length=500)


class GradeCreate(BaseModel):
    context_id: int
    course_id: int
    student_user_id: int
    grade_type: Literal[
        "daily", "homework", "quiz", "mock_exam", "exam", "project", "final"
    ] = "daily"
    grade_date: date | None = None
    assessment_name: str | None = Field(default=None, max_length=200)
    score: Decimal = Field(ge=0)
    max_score: Decimal = Field(gt=0)
    note: str | None = Field(default=None, max_length=1000)
    idempotency_key: str | None = Field(default=None, min_length=12, max_length=120)


class LessonPlanCreate(BaseModel):
    context_id: int
    course_id: int
    lesson_date: date | None = None
    title: str = Field(min_length=2, max_length=240)
    objective: str | None = Field(default=None, max_length=5000)
    explanation: str | None = Field(default=None, max_length=20000)
    activities: str | None = Field(default=None, max_length=20000)
    formula_latex: str | None = Field(default=None, max_length=10000)
    objectives: list[Any] = Field(default_factory=list, max_length=100)
    stages: list[Any] = Field(default_factory=list, max_length=100)
    content_text: str | None = Field(default=None, max_length=50_000)
    content_latex: str | None = Field(default=None, max_length=20_000)
    source_refs: list[Any] = Field(default_factory=list, max_length=200)
    duration_minutes: int = Field(default=45, ge=15, le=480)


class HomeworkCreate(BaseModel):
    context_id: int
    course_id: int
    lesson_plan_id: int | None = None
    title: str = Field(min_length=2, max_length=240)
    instructions: str | None = Field(default=None, max_length=50_000)
    homework_text: str | None = Field(default=None, max_length=20000)
    content_latex: str | None = Field(default=None, max_length=20_000)
    formula_latex: str | None = Field(default=None, max_length=10000)
    resource_refs: list[Any] = Field(default_factory=list, max_length=200)
    due_at: datetime | None = None
    due_date: date | None = None
    max_score: Decimal = Field(default=100, gt=0)
    status: Literal["draft", "published"] = "draft"
    confirmation: bool = False


class HomeworkSubmit(BaseModel):
    context_id: int
    answer_text: str | None = Field(default=None, max_length=50000)
    answer_latex: str | None = Field(default=None, max_length=20000)
    attachment_refs: list[Any] = Field(default_factory=list, max_length=50)


class HomeworkGrade(BaseModel):
    context_id: int
    score: Decimal = Field(ge=0)
    feedback: str | None = Field(default=None, max_length=10000)
    confirmation: bool = False


class AssessmentItemCreate(BaseModel):
    question_ref: str = Field(min_length=1, max_length=200)
    question_source: str = Field(default="generated_tests", max_length=80)
    points: Decimal = Field(default=1, gt=0)
    section_key: str | None = Field(default=None, max_length=80)
    metadata: dict[str, Any] = Field(default_factory=dict, max_length=100)


class AssessmentCreate(BaseModel):
    context_id: int
    course_id: int
    assessment_type: Literal[
        "diagnostic", "placement", "quiz", "midterm", "mock_exam", "exam",
        "final", "cefr_mock", "ielts_mock", "other",
    ]
    title: str = Field(min_length=2, max_length=240)
    instructions: str | None = Field(default=None, max_length=50_000)
    framework: Literal[
        "custom", "cefr", "ielts", "national_exam", "school", "other"
    ] = "custom"
    duration_minutes: int | None = Field(default=None, ge=1, le=600)
    max_attempts: int = Field(default=1, ge=1, le=100)
    opens_at: datetime | str | None = None
    closes_at: datetime | str | None = None
    items: list[AssessmentItemCreate] = Field(default_factory=list, max_length=500)
    settings: dict[str, Any] = Field(default_factory=dict, max_length=100)
    total_score: Decimal | None = Field(default=None, ge=0)
    formula_latex: str | None = Field(default=None, max_length=10000)


class AttemptSubmit(BaseModel):
    context_id: int | None = Field(default=None, ge=1)
    answers: list[dict[str, Any]] = Field(default_factory=list, max_length=1000)
    idempotency_key: str = Field(min_length=12, max_length=120)


class AttemptDraft(BaseModel):
    context_id: int | None = Field(default=None, ge=1)
    answers: list[dict[str, Any]] = Field(default_factory=list, max_length=1000)


class AttemptScore(BaseModel):
    context_id: int
    score: Decimal = Field(ge=0)
    max_score: Decimal = Field(gt=0)
    component_scores: dict[str, Decimal] = Field(
        default_factory=dict, max_length=20
    )
    idempotency_key: str = Field(min_length=12, max_length=120)
    confirmation: bool


class BillingPlanCreate(BaseModel):
    context_id: int
    course_id: int | None = None
    name: str = Field(min_length=2, max_length=180)
    amount: Decimal = Field(gt=0)
    billing_cycle: Literal[
        "one_time", "per_lesson", "weekly", "monthly", "course"
    ] = "monthly"
    currency: str = Field(default="UZS", min_length=3, max_length=8)
    billing_day: int | None = Field(default=5, ge=1, le=28)
    confirmation: bool = False


class DiscountCreate(BaseModel):
    context_id: int
    name: str = Field(min_length=2, max_length=180)
    discount_type: Literal["fixed", "percent"]
    value: Decimal = Field(gt=0)
    student_user_id: int | None = None
    course_id: int | None = None
    starts_on: date | None = None
    ends_on: date | None = None
    confirmation: bool = False


class InvoiceCreate(BaseModel):
    context_id: int
    plan_id: int
    enrollment_id: int
    period_start: date
    period_end: date
    due_date: date
    discount_amount: Decimal = Field(default=0, ge=0)
    confirmation: bool = False


class PaymentCreate(BaseModel):
    context_id: int
    invoice_id: int
    amount: Decimal = Field(gt=0)
    payment_method: Literal[
        "cash", "card", "bank_transfer", "online", "other"
    ] = "cash"
    idempotency_key: str = Field(min_length=12, max_length=120)
    reference: str | None = Field(default=None, max_length=300)
    paid_at: date | datetime | None = None
    confirmation: bool = False


class WorkLogCreate(BaseModel):
    context_id: int
    teacher_user_id: int
    course_id: int
    schedule_slot_id: int | None = None
    work_date: date
    minutes_worked: int = Field(ge=1, le=1440)
    pay_unit: Literal["lesson", "hour", "fixed"] = "lesson"
    rate: Decimal | None = Field(default=None, ge=0)
    note: str | None = Field(default=None, max_length=1000)
    idempotency_key: str | None = Field(default=None, min_length=12, max_length=120)


class WorkLogDecision(BaseModel):
    context_id: int
    status: Literal["approved", "rejected"]
    confirmation: bool = False


class AssistantStart(BaseModel):
    workflow_key: str = Field(default="center_onboarding", max_length=100)
    context_id: int | None = None
    draft_id: int | None = None
    avatar_enabled: bool = True
    speech_enabled: bool = True
    avatar_variant: Literal["female", "male", "neutral"] = "female"


class AssistantAction(BaseModel):
    action_id: str = Field(min_length=1, max_length=60)
    ui_anchor: str | None = Field(default=None, max_length=120)
    payload: dict[str, Any] = Field(default_factory=dict)


def create_learning_center_router(jwt_check: Callable[[str], int]) -> APIRouter:
    router = APIRouter(prefix="/api/markaz-v2", tags=["O'quv markazi v2"])

    def authenticated_user(
        authorization: str | None = Header(default=None),
    ) -> int:
        if not authorization or not authorization.lower().startswith("bearer "):
            raise HTTPException(status_code=401, detail="Bearer sessiya tokeni topilmadi")
        token = authorization[7:].strip()
        if not token:
            raise HTTPException(status_code=401, detail="Sessiya tokeni bo'sh")
        return int(jwt_check(token))

    @contextmanager
    def database(*, readonly: bool = False) -> Iterator[tuple[Any, Any]]:
        try:
            with db_session(readonly=readonly) as result:
                yield result
        except DatabaseBusyError as exc:
            raise HTTPException(status_code=503, detail=str(exc)) from exc
        except psycopg2.errors.UndefinedTable as exc:
            raise HTTPException(
                status_code=503,
                detail=(
                    "Markaz bazasi o'rnatilmagan: 001, 007 va 008 SQL "
                    "fayllarini tartib bilan bajaring."
                ),
            ) from exc
        except psycopg2.errors.ExclusionViolation as exc:
            raise HTTPException(
                status_code=409,
                detail="O'qituvchi, guruh yoki xona vaqti boshqa dars bilan to'qnashdi.",
            ) from exc
        except psycopg2.errors.UniqueViolation as exc:
            raise HTTPException(
                status_code=409,
                detail="Bu ma'lumot avval saqlangan; takror yozuv yaratilmadi.",
            ) from exc
        except psycopg2.errors.ForeignKeyViolation as exc:
            raise HTTPException(
                status_code=422,
                detail="Bog'langan foydalanuvchi yoki resurs topilmadi.",
            ) from exc
        except (
            psycopg2.errors.InvalidTextRepresentation,
            psycopg2.errors.NotNullViolation,
            psycopg2.DataError,
        ) as exc:
            raise HTTPException(
                status_code=422,
                detail="Yuborilgan ma'lumot formati noto'g'ri.",
            ) from exc
        except psycopg2.errors.CheckViolation as exc:
            message = str(exc)
            if "capacity" in message:
                detail = "Kurs sig'imi to'lgan; o'quvchini kutish ro'yxatiga qo'ying."
                code = 409
            else:
                detail = "Yuborilgan ma'lumot markaz qoidalariga mos emas."
                code = 422
            raise HTTPException(status_code=code, detail=detail) from exc

    def ensure_schema(cur: Any) -> None:
        cur.execute(
            """SELECT version FROM app_schema_migrations
               WHERE version IN (
                 '007_learning_center_core','008_learning_center_operations'
               )"""
        )
        if len(cur.fetchall()) != 2:
            raise HTTPException(
                status_code=503,
                detail="database/007 va database/008 markaz migratsiyalari bajarilmagan.",
            )

    def system_admin(cur: Any, user_id: int) -> bool:
        cur.execute("SELECT 1 FROM admin_akkaunt WHERE uid=%s", (user_id,))
        return cur.fetchone() is not None

    def require_human(confirmation: bool, message: str) -> None:
        if confirmation is not True:
            raise HTTPException(status_code=409, detail=message)

    def parse_time(value: str, field: str) -> time:
        if not re.fullmatch(r"(?:[01]\d|2[0-3]):[0-5]\d", value.strip()):
            raise HTTPException(status_code=422, detail=f"{field} HH:MM shaklida bo'lsin")
        return time.fromisoformat(value.strip())

    def normalize_work_days(value: Any, field: str = "work_days") -> list[int]:
        if value is None:
            return [1, 2, 3, 4, 5, 6]
        if not isinstance(value, list):
            raise HTTPException(
                status_code=422, detail=f"{field} ro'yxat bo'lishi kerak"
            )
        try:
            days = [int(item) for item in value]
        except (TypeError, ValueError) as exc:
            raise HTTPException(
                status_code=422, detail=f"{field} faqat 1–7 kun raqamlarini oladi"
            ) from exc
        if not days or len(days) > 7 or len(set(days)) != len(days):
            raise HTTPException(
                status_code=422,
                detail=f"{field} takrorlanmagan 1–7 kunlardan iborat bo'lsin",
            )
        if any(day < 1 or day > 7 for day in days):
            raise HTTPException(
                status_code=422, detail=f"{field} kunlari 1 dan 7 gacha bo'lsin"
            )
        return sorted(days)

    def require_encoded_size(
        value: Any, *, maximum: int, label: str,
    ) -> None:
        encoded = json.dumps(
            value, ensure_ascii=False, default=str, separators=(",", ":")
        ).encode("utf-8")
        if len(encoded) > maximum:
            raise HTTPException(
                status_code=413,
                detail=f"{label} hajmi {maximum} baytdan oshmasin.",
            )

    def roles_for_permission(permission: str) -> set[str]:
        return {
            role
            for role, permissions in ROLE_PERMISSIONS.items()
            if permission in permissions
        }

    def require_permission(
        cur: Any, context_id: int, user_id: int, permission: str,
        *, branch_id: int | None = None, require_global: bool = False,
    ) -> set[str]:
        allowed = roles_for_permission(permission)
        if not allowed:
            raise RuntimeError(f"Noma'lum permission: {permission}")
        return require_roles(
            cur, context_id, user_id, allowed, branch_id=branch_id,
            require_global=require_global or branch_id is None,
        )

    def require_role_hierarchy(
        cur: Any, *, context_id: int, actor_user_id: int,
        actor_roles: set[str], target_user_id: int, target_role: str,
    ) -> None:
        if system_admin(cur, actor_user_id):
            return
        if actor_user_id == target_user_id:
            raise HTTPException(
                status_code=403,
                detail="O'zingizga rol bera yoki o'z vakolatingizni to'xtata olmaysiz",
            )
        allowed_targets: set[str] = set()
        for actor_role in actor_roles:
            allowed_targets.update(ROLE_GRANT_MATRIX.get(actor_role, set()))
        if target_role not in allowed_targets:
            raise HTTPException(
                status_code=403,
                detail="Siz bu darajadagi rolni boshqara olmaysiz",
            )

    def has_permission(
        cur: Any, context_id: int, user_id: int, permission: str,
        *, branch_id: int | None = None,
    ) -> bool:
        if system_admin(cur, user_id):
            return True
        allowed = roles_for_permission(permission)
        cur.execute(
            """SELECT 1 FROM center_role_assignments
               WHERE context_id=%s AND user_id=%s AND status='active'
                 AND starts_at<=NOW() AND (ends_at IS NULL OR ends_at>NOW())
                 AND role_key=ANY(%s)
                 AND (
                   (%s IS NULL AND branch_id IS NULL)
                   OR (
                     %s IS NOT NULL
                     AND (branch_id IS NULL OR branch_id=%s)
                   )
                 )
               LIMIT 1""",
            (
                context_id, user_id, list(allowed),
                branch_id, branch_id, branch_id,
            ),
        )
        return cur.fetchone() is not None

    def course_resource(
        cur: Any, context_id: int, course_id: int, *, lock: bool = False,
    ) -> dict[str, Any]:
        cur.execute(
            f"""SELECT c.*,g.name group_name
                FROM center_courses c
                JOIN course_groups g ON g.id=c.group_id
                WHERE c.id=%s AND c.context_id=%s
                {"FOR UPDATE" if lock else ""}""",
            (course_id, context_id),
        )
        course = cur.fetchone()
        if not course:
            raise HTTPException(status_code=404, detail="Kurs topilmadi")
        return dict(course)

    def lock_course_queue(cur: Any, context_id: int, course_id: int) -> None:
        cur.execute(
            """SELECT pg_advisory_xact_lock(
                 74129,hashtext('center-course-'||%s::TEXT||'-'||%s::TEXT)
               )""",
            (context_id, course_id),
        )
        cur.execute(
            """SELECT 1 FROM center_courses
               WHERE id=%s AND context_id=%s FOR UPDATE""",
            (course_id, context_id),
        )
        if not cur.fetchone():
            raise HTTPException(status_code=404, detail="Kurs topilmadi")

    def ensure_teacher_available(
        cur: Any, *, context_id: int, teacher_user_id: int,
        schedule_kind: str, weekday: int, starts_at: time, ends_at: time,
        lesson_date: date | None = None, effective_from: date | None = None,
        effective_to: date | None = None,
    ) -> None:
        if schedule_kind == "dated":
            cur.execute(
                """SELECT 1 FROM center_teacher_availability
                   WHERE context_id=%s AND teacher_user_id=%s
                     AND weekday=%s AND availability='unavailable'
                     AND starts_at<%s AND %s<ends_at
                     AND %s BETWEEN
                       COALESCE(effective_from,DATE '0001-01-01')
                       AND COALESCE(effective_to,DATE '9999-12-31')
                   LIMIT 1""",
                (
                    context_id, teacher_user_id, weekday, ends_at, starts_at,
                    lesson_date,
                ),
            )
        else:
            cur.execute(
                """SELECT 1 FROM center_teacher_availability
                   WHERE context_id=%s AND teacher_user_id=%s
                     AND weekday=%s AND availability='unavailable'
                     AND starts_at<%s AND %s<ends_at
                     AND COALESCE(effective_from,DATE '0001-01-01')
                           <=COALESCE(%s,DATE '9999-12-31')
                     AND COALESCE(%s,DATE '0001-01-01')
                           <=COALESCE(effective_to,DATE '9999-12-31')
                   LIMIT 1""",
                (
                    context_id, teacher_user_id, weekday, ends_at, starts_at,
                    effective_to, effective_from,
                ),
            )
        if cur.fetchone():
            raise HTTPException(
                status_code=409,
                detail="O'qituvchi tanlangan davrda bu vaqtda mavjud emas.",
            )

    def require_branch_workday(
        cur: Any, *, context_id: int, branch_id: int | None, weekday: int,
        starts_at: time | None = None, ends_at: time | None = None,
    ) -> None:
        if branch_id is not None:
            cur.execute(
                """SELECT work_days,work_start,work_end FROM center_branches
                   WHERE id=%s AND context_id=%s AND active=TRUE""",
                (branch_id, context_id),
            )
            branch = cur.fetchone()
            if not branch:
                raise HTTPException(status_code=404, detail="Filial topilmadi")
            work_days = list(branch["work_days"] or [1, 2, 3, 4, 5, 6])
            if (
                starts_at is not None
                and branch["work_start"] is not None
                and starts_at < branch["work_start"]
            ):
                raise HTTPException(
                    status_code=409,
                    detail="Dars filialning ish boshlanish vaqtidan oldin",
                )
            if (
                ends_at is not None
                and branch["work_end"] is not None
                and ends_at > branch["work_end"]
            ):
                raise HTTPException(
                    status_code=409,
                    detail="Dars filialning ish tugash vaqtidan keyin",
                )
        else:
            cur.execute(
                "SELECT work_days FROM center_profiles WHERE context_id=%s",
                (context_id,),
            )
            profile = cur.fetchone()
            work_days = list(
                (profile and profile["work_days"]) or [1, 2, 3, 4, 5, 6]
            )
        if weekday not in work_days:
            raise HTTPException(
                status_code=409,
                detail="Tanlangan kun filialning ish kunlari ro'yxatida yo'q",
            )

    def ensure_teacher_weekly_capacity(
        cur: Any, *, context_id: int, teacher_user_id: int,
        branch_id: int | None, starts_at: time, ends_at: time,
    ) -> None:
        cur.execute(
            """SELECT pg_advisory_xact_lock(
                 74127,hashtext('center-teacher-'||%s::TEXT)
               )""",
            (teacher_user_id,),
        )
        duration_minutes = (
            ends_at.hour * 60 + ends_at.minute
            - starts_at.hour * 60 - starts_at.minute
        )
        cur.execute(
            """SELECT branch_id,weekly_capacity_hours
               FROM center_role_assignments
               WHERE context_id=%s AND user_id=%s AND role_key='teacher'
                 AND status='active' AND starts_at<=NOW()
                 AND (ends_at IS NULL OR ends_at>NOW())
                 AND weekly_capacity_hours IS NOT NULL
                 AND (branch_id IS NULL OR branch_id IS NOT DISTINCT FROM %s)""",
            (context_id, teacher_user_id, branch_id),
        )
        limits = cur.fetchall()
        for limit_row in limits:
            scoped_branch = limit_row["branch_id"]
            cur.execute(
                """SELECT COALESCE(SUM(
                     EXTRACT(EPOCH FROM (ends_at-starts_at))/60
                   ),0) used_minutes
                   FROM center_schedule_slots
                   WHERE context_id=%s AND teacher_user_id=%s
                     AND schedule_kind='weekly' AND status<>'cancelled'
                     AND (%s IS NULL OR branch_id=%s)""",
                (
                    context_id, teacher_user_id, scoped_branch, scoped_branch,
                ),
            )
            used = Decimal(str(cur.fetchone()["used_minutes"] or 0))
            allowed = Decimal(str(limit_row["weekly_capacity_hours"])) * 60
            if used + duration_minutes > allowed:
                scope_label = "markaz" if scoped_branch is None else "filial"
                raise HTTPException(
                    status_code=409,
                    detail=(
                        f"O'qituvchining {scope_label} bo'yicha haftalik "
                        "yuklama limiti oshib ketadi"
                    ),
                )

    def audit(
        cur: Any, context_id: int | None, actor: int, action: str,
        target_type: str | None = None, target_id: int | None = None,
        payload: dict[str, Any] | None = None,
    ) -> None:
        cur.execute(
            """INSERT INTO center_audit_log(
                 context_id,actor_user_id,action_key,target_type,target_id,payload
               ) VALUES(%s,%s,%s,%s,%s,%s::jsonb)""",
            (
                context_id, actor, action, target_type, target_id,
                json.dumps(payload or {}, ensure_ascii=False, default=str),
            ),
        )

    def active_roles(
        cur: Any, context_id: int, user_id: int,
    ) -> tuple[set[str], list[dict[str, Any]]]:
        cur.execute(
            """SELECT c.active,p.onboarding_status,p.verification_status
               FROM learning_contexts c
               JOIN center_profiles p ON p.context_id=c.id
               WHERE c.id=%s AND c.context_type='learning_center'""",
            (context_id,),
        )
        state = cur.fetchone()
        if (
            not state or not state["active"]
            or state["onboarding_status"] in {"suspended", "archived"}
            or state["verification_status"] == "rejected"
        ):
            raise HTTPException(status_code=403, detail="Bu o'quv markazi faol emas.")
        cur.execute(
            """SELECT role_key,branch_id FROM center_role_assignments
               WHERE context_id=%s AND user_id=%s AND status='active'
                 AND starts_at<=NOW() AND (ends_at IS NULL OR ends_at>NOW())""",
            (context_id, user_id),
        )
        assignments = [dict(row) for row in cur.fetchall()]
        return {row["role_key"] for row in assignments}, assignments

    def permissions_for(roles: set[str]) -> set[str]:
        if "system_admin" in roles:
            return set(ALL_PERMISSIONS)
        result: set[str] = set()
        for role in roles:
            result.update(ROLE_PERMISSIONS.get(role, set()))
        return result

    def require_roles(
        cur: Any, context_id: int, user_id: int, allowed: set[str],
        *, branch_id: int | None = None, require_global: bool = False,
    ) -> set[str]:
        if system_admin(cur, user_id):
            return {"system_admin"}
        roles, assignments = active_roles(cur, context_id, user_id)
        matching = [row for row in assignments if row["role_key"] in allowed]
        if not matching:
            raise HTTPException(
                status_code=403,
                detail="Bu amal uchun markazdagi vakolatingiz yetarli emas.",
            )
        if branch_id is not None:
            matching = [
                row for row in matching
                if row["branch_id"] is None or int(row["branch_id"]) == branch_id
            ]
            if not matching:
                raise HTTPException(
                    status_code=403,
                    detail="Siz boshqa filial ma'lumotini o'zgartira olmaysiz.",
                )
        if require_global and not any(row["branch_id"] is None for row in matching):
            raise HTTPException(
                status_code=403,
                detail="Bu amal uchun markaz miqyosidagi vakolat kerak.",
            )
        return {row["role_key"] for row in matching}

    def branch_scope(
        cur: Any, context_id: int, user_id: int, allowed: set[str],
    ) -> tuple[set[str], bool, set[int]]:
        if system_admin(cur, user_id):
            return {"system_admin"}, True, set()
        roles, assignments = active_roles(cur, context_id, user_id)
        matching = [row for row in assignments if row["role_key"] in allowed]
        if not matching:
            raise HTTPException(status_code=403, detail="Markaz vakolati yetarli emas")
        is_global = any(row["branch_id"] is None for row in matching)
        branches = {
            int(row["branch_id"]) for row in matching if row["branch_id"] is not None
        }
        return {row["role_key"] for row in matching}, is_global, branches

    def require_operational(cur: Any, context_id: int) -> None:
        cur.execute(
            """SELECT onboarding_status,verification_status
               FROM center_profiles WHERE context_id=%s""",
            (context_id,),
        )
        row = cur.fetchone()
        if not row or row["onboarding_status"] != "active":
            raise HTTPException(
                status_code=409,
                detail="Markaz tekshiruvdan o'tib faol bo'lgandan keyin bu amal bajariladi.",
            )

    def course_access(
        cur: Any, context_id: int, course_id: int, user_id: int,
        *, write: bool = False,
    ) -> dict[str, Any]:
        cur.execute(
            """SELECT c.*,g.name group_name
               FROM center_courses c JOIN course_groups g ON g.id=c.group_id
               WHERE c.id=%s AND c.context_id=%s""",
            (course_id, context_id),
        )
        course = cur.fetchone()
        if not course:
            raise HTTPException(status_code=404, detail="Kurs topilmadi")
        if system_admin(cur, user_id):
            return dict(course)
        roles, assignments = active_roles(cur, context_id, user_id)
        def scoped_assignment_ok(role_keys: set[str]) -> bool:
            return any(
                row["role_key"] in role_keys
                and (
                    row["branch_id"] is None
                    or (
                        course["branch_id"] is not None
                        and int(row["branch_id"]) == int(course["branch_id"])
                    )
                )
                for row in assignments
            )

        if roles & MANAGER_ROLES and scoped_assignment_ok(MANAGER_ROLES):
            return dict(course)
        if "methodist" in roles and scoped_assignment_ok({"methodist"}):
            return dict(course)
        if (
            "teacher" in roles
            and int(course["teacher_user_id"] or 0) == user_id
            and scoped_assignment_ok({"teacher"})
        ):
            return dict(course)
        if not write and "student" in roles:
            cur.execute(
                """SELECT 1 FROM center_enrollments
                   WHERE context_id=%s AND course_id=%s
                     AND student_user_id=%s AND status IN ('active','completed')
                     AND (
                       starts_on IS NULL
                       OR starts_on<=
                         (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                     )""",
                (context_id, course_id, user_id),
            )
            if cur.fetchone():
                return dict(course)
        if not write and "parent" in roles:
            cur.execute(
                """SELECT 1 FROM center_parent_links p
                   JOIN center_enrollments e
                     ON e.context_id=p.context_id
                    AND e.student_user_id=p.student_user_id
                   WHERE p.context_id=%s AND p.parent_user_id=%s
                     AND p.status='active' AND e.course_id=%s
                     AND e.status IN ('active','completed')
                     AND (
                       e.starts_on IS NULL
                       OR e.starts_on<=
                         (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                     )""",
                (context_id, user_id, course_id),
            )
            if cur.fetchone():
                return dict(course)
        raise HTTPException(status_code=403, detail="Bu kurs uchun ruxsat yo'q")

    def active_enrollment(
        cur: Any, context_id: int, course_id: int, student_user_id: int,
    ) -> dict[str, Any]:
        cur.execute(
            """SELECT * FROM center_enrollments
               WHERE context_id=%s AND course_id=%s AND student_user_id=%s
                 AND status='active'
                 AND (
                   starts_on IS NULL
                   OR starts_on<=
                     (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                 )""",
            (context_id, course_id, student_user_id),
        )
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=409, detail="O'quvchi bu kursda faol emas")
        return dict(row)

    def require_teacher_scope(
        cur: Any, context_id: int, actor_user_id: int, teacher_user_id: int,
    ) -> set[str]:
        if actor_user_id == teacher_user_id:
            return require_roles(
                cur, context_id, actor_user_id, {"teacher", "methodist"} | MANAGER_ROLES
            )
        roles, global_scope, actor_branches = branch_scope(
            cur, context_id, actor_user_id, MANAGER_ROLES
        )
        if global_scope:
            return roles
        cur.execute(
            """SELECT branch_id FROM center_role_assignments
               WHERE context_id=%s AND user_id=%s AND role_key='teacher'
                 AND status='active' AND starts_at<=NOW()
                 AND (ends_at IS NULL OR ends_at>NOW())
                 AND branch_id=ANY(%s)""",
            (context_id, teacher_user_id, list(actor_branches)),
        )
        if not cur.fetchone():
            raise HTTPException(status_code=403, detail="Boshqa filial o'qituvchisi yopiq")
        return roles

    def generic_member_role(role_key: str) -> str:
        return (
            "student" if role_key == "student"
            else "parent_observer" if role_key == "parent"
            else "director" if role_key == "director"
            else "manager" if role_key in MANAGER_ROLES
            else "teacher" if role_key in {"teacher", "methodist"}
            else "assistant"
        )

    def sync_generic_staff_membership(
        cur: Any, *, context_id: int, user_id: int, generic_role: str,
    ) -> None:
        cur.execute(
            """SELECT role_key FROM center_role_assignments
               WHERE context_id=%s AND user_id=%s AND status='active'
                 AND starts_at<=NOW() AND (ends_at IS NULL OR ends_at>NOW())""",
            (context_id, user_id),
        )
        remains_active = any(
            generic_member_role(str(row["role_key"])) == generic_role
            for row in cur.fetchall()
        )
        cur.execute(
            """UPDATE context_memberships SET status=%s,
                 ended_at=CASE WHEN %s THEN NULL ELSE NOW() END,
                 updated_at=NOW()
               WHERE context_id=%s AND user_id=%s AND group_id IS NULL
                 AND member_role=%s AND source='center_v2'""",
            (
                "active" if remains_active else "suspended",
                remains_active, context_id, user_id, generic_role,
            ),
        )

    def upsert_role(
        cur: Any, *, context_id: int, user_id: int, role_key: str,
        branch_id: int | None, status: str, approved_by: int | None,
        employment_type: str | None = None,
        weekly_capacity_hours: int | None = None,
    ) -> int:
        cur.execute(
            """INSERT INTO center_role_assignments(
                 context_id,branch_id,user_id,role_key,status,approved_by_user_id,
                 permissions,employment_type,weekly_capacity_hours
               ) VALUES(
                 %s,%s,%s,%s,%s,%s,'{"source":"center_v2"}',%s,%s
               )
               ON CONFLICT(
                 context_id,(COALESCE(branch_id,0)),user_id,role_key
               ) DO UPDATE SET status=EXCLUDED.status,
                 approved_by_user_id=EXCLUDED.approved_by_user_id,
                 employment_type=COALESCE(
                   EXCLUDED.employment_type,
                   center_role_assignments.employment_type
                 ),
                 weekly_capacity_hours=COALESCE(
                   EXCLUDED.weekly_capacity_hours,
                   center_role_assignments.weekly_capacity_hours
                 ),
                 ends_at=NULL,updated_at=NOW()
               RETURNING id""",
            (
                context_id, branch_id, user_id, role_key, status, approved_by,
                employment_type, weekly_capacity_hours,
            ),
        )
        assignment_id = int(cur.fetchone()["id"])
        generic = generic_member_role(role_key)
        cur.execute(
            """INSERT INTO context_memberships(
                 context_id,user_id,member_role,status,source,
                 approved_by_user_id,metadata
               ) VALUES(%s,%s,%s,%s,'center_v2',%s,%s::jsonb)
               ON CONFLICT(
                 context_id,(COALESCE(group_id,0)),user_id,member_role
               ) DO UPDATE SET status=EXCLUDED.status,
                 approved_by_user_id=EXCLUDED.approved_by_user_id,
                 ended_at=NULL,updated_at=NOW(),metadata=EXCLUDED.metadata""",
            (
                context_id, user_id, generic, status, approved_by,
                json.dumps(
                    {"center_role": role_key, "branch_id": branch_id},
                    ensure_ascii=False,
                ),
            ),
        )
        return assignment_id

    def sync_student_access(
        cur: Any, *, context_id: int, student_user_id: int,
        branch_id: int | None, approved_by: int | None,
    ) -> None:
        cur.execute(
            """SELECT EXISTS(
                 SELECT 1 FROM center_enrollments e
                 JOIN center_courses c
                   ON c.context_id=e.context_id AND c.id=e.course_id
                 WHERE e.context_id=%s AND e.student_user_id=%s
                   AND c.branch_id IS NOT DISTINCT FROM %s
                   AND e.status IN ('active','completed')
               ) branch_active""",
            (context_id, student_user_id, branch_id),
        )
        branch_active = bool(cur.fetchone()["branch_active"])
        upsert_role(
            cur, context_id=context_id, user_id=student_user_id,
            role_key="student", branch_id=branch_id,
            status="active" if branch_active else "pending",
            approved_by=approved_by,
        )
        cur.execute(
            """SELECT EXISTS(
                 SELECT 1 FROM center_enrollments
                 WHERE context_id=%s AND student_user_id=%s
                   AND status IN ('active','completed')
               ) context_active""",
            (context_id, student_user_id),
        )
        context_status = "active" if cur.fetchone()["context_active"] else "pending"
        cur.execute(
            """UPDATE context_memberships SET status=%s,
                 ended_at=CASE WHEN %s='active' THEN NULL ELSE ended_at END,
                 updated_at=NOW()
               WHERE context_id=%s AND user_id=%s
                 AND member_role='student' AND source='center_v2'""",
            (
                context_status, context_status, context_id, student_user_id,
            ),
        )

    def claim_request(
        cur: Any, *, context_id: int, actor: int, key: str, action: str,
    ) -> int | None:
        cur.execute(
            """INSERT INTO center_request_keys(
                 context_id,request_key,actor_user_id,action_key
               ) VALUES(%s,%s,%s,%s)
               ON CONFLICT DO NOTHING RETURNING target_id""",
            (context_id, key, actor, action),
        )
        inserted = cur.fetchone()
        if inserted is not None:
            return None
        cur.execute(
            """SELECT actor_user_id,action_key,target_id
               FROM center_request_keys
               WHERE context_id=%s AND request_key=%s""",
            (context_id, key),
        )
        existing = cur.fetchone()
        if (
            not existing or int(existing["actor_user_id"]) != actor
            or existing["action_key"] != action
        ):
            raise HTTPException(
                status_code=409, detail="Idempotency kaliti boshqa amal uchun ishlatilgan."
            )
        return int(existing["target_id"]) if existing["target_id"] is not None else 0

    def validate_request_key(key: str | None) -> str | None:
        if key is None:
            return None
        normalized = key.strip()
        if not 12 <= len(normalized) <= 120:
            raise HTTPException(
                status_code=422,
                detail="Idempotency-Key 12–120 belgidan iborat bo'lsin",
            )
        return normalized

    def finish_request(
        cur: Any, context_id: int, key: str, target_id: int,
    ) -> None:
        cur.execute(
            """UPDATE center_request_keys SET target_id=%s
               WHERE context_id=%s AND request_key=%s""",
            (target_id, context_id, key),
        )

    @router.get("/health")
    def health() -> dict[str, str]:
        return {
            "status": "ready",
            "module": "learning_center",
            "version": "learning-center-v2",
            "schema": "007_learning_center_core+008_learning_center_operations",
        }

    @router.get("/meta")
    def meta(_: int = Depends(authenticated_user)) -> dict[str, Any]:
        return {
            "roles": [
                {"key": key, "label": value} for key, value in ROLE_LABELS.items()
            ],
            "permissions": sorted(ALL_PERMISSIONS),
            "ownership_types": ["public", "private"],
            "operator_models": ["center", "independent_tutor"],
            "course_types": [
                "group", "individual", "intensive", "club", "exam_prep",
            ],
            "delivery_modes": ["offline", "online_live", "hybrid"],
            "target_frameworks": [
                "custom", "cefr", "ielts", "national_exam", "school", "other",
            ],
            "cefr_levels": list(CEFR_LEVELS),
            "ielts": {
                "minimum": 0,
                "maximum": 9,
                "step": 0.5,
                "components": list(IELTS_COMPONENTS),
                "cefr_exact_equivalence_claimed": False,
            },
            "content_capabilities": {
                "subjects_are_database_driven": True,
                "latex_supported": True,
                "formula_metadata_supported": True,
            },
            "assistant_limits": {
                "can_confirm_privileged_actions": False,
                "can_publish": False,
                "can_record_payment": False,
                "can_assign_roles": False,
                "can_score_attempts": False,
                "can_bypass_permissions": False,
            },
            "scheduler_scope": {
                "cross_context_teacher_collision": "learning_center_v2_only",
                "school_and_university_cross_collision": False,
            },
        }

    @router.post("/onboarding/drafts")
    def start_draft(
        request: DraftStart, user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        if (
            request.operator_model == "independent_tutor"
            and request.ownership_type != "private"
        ):
            raise HTTPException(
                status_code=422,
                detail="Mustaqil repetitor faqat xususiy operator sifatida ochiladi.",
            )
        if (
            request.operator_model == "independent_tutor"
            and request.relationship not in {"owner", "teacher"}
        ):
            raise HTTPException(
                status_code=422,
                detail="Mustaqil repetitor o'zini mulkdor yoki o'qituvchi sifatida ochadi.",
            )
        if (
            request.ownership_type == "public"
            and request.relationship in {"owner", "founder"}
        ):
            raise HTTPException(
                status_code=422, detail="Davlat markazida mulkdor roli bo'lmaydi."
            )
        if (
            request.operator_model == "center"
            and request.relationship not in {"owner", "founder", "director"}
        ):
            raise HTTPException(
                status_code=422,
                detail=(
                    "Oddiy markazni o'qituvchi roli bilan boshqaruvchisiz "
                    "ochib bo'lmaydi. Direktor, administrator, mulkdor yoki "
                    "ta'sischini tanlang; administrator keyin tayinlanadi. "
                    "Yolg'iz repetitor uchun mustaqil "
                    "repetitor turini tanlang."
                ),
            )
        if (
            request.ownership_type == "public"
            and request.relationship != "director"
        ):
            raise HTTPException(
                status_code=422,
                detail="Davlat markazini mas'ul direktor ariza qiladi.",
            )
        with database() as (_, cur):
            ensure_schema(cur)
            cur.execute(
                """INSERT INTO center_setup_drafts(
                     creator_user_id,relationship,ownership_type,operator_model,
                     setup_mode
                   ) VALUES(%s,%s,%s,%s,%s)
                   RETURNING id,current_step,status,version,expires_at""",
                (
                    user_id, request.relationship, request.ownership_type,
                    request.operator_model, request.setup_mode,
                ),
            )
            draft = cur.fetchone()
        return {"draft": draft, "steps": list(ONBOARDING_STEPS)}

    @router.patch("/onboarding/drafts/{draft_id}")
    def patch_draft(
        draft_id: int, request: DraftPatch,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        if request.step not in ONBOARDING_STEPS:
            raise HTTPException(status_code=422, detail="Noma'lum sozlash bosqichi")
        encoded = json.dumps(request.payload, ensure_ascii=False, default=str)
        if len(encoded.encode()) > 150_000:
            raise HTTPException(status_code=413, detail="Qoralama juda katta")
        with database() as (_, cur):
            ensure_schema(cur)
            cur.execute(
                """SELECT * FROM center_setup_drafts
                   WHERE id=%s AND creator_user_id=%s AND status='draft'
                     AND expires_at>NOW() FOR UPDATE""",
                (draft_id, user_id),
            )
            draft = cur.fetchone()
            if not draft or int(draft["version"]) != request.expected_version:
                raise HTTPException(
                    status_code=409,
                    detail="Qoralama versiyasi o'zgargan yoki muddati tugagan.",
                )
            payload = dict(draft["payload"] or {})
            payload[request.step] = request.payload
            ownership_type = str(draft["ownership_type"])
            operator_model = str(draft["operator_model"])
            relationship = str(draft["relationship"])
            if request.step == "identity":
                center_type = str(request.payload.get("center_type") or "")
                ownership_type = str(
                    request.payload.get("ownership_type")
                    or (
                        "public" if center_type == "public_center"
                        else "private"
                    )
                )
                operator_model = str(
                    request.payload.get("operator_model")
                    or (
                        "independent_tutor"
                        if center_type == "independent_tutor"
                        else "center"
                    )
                )
                if ownership_type not in {"public", "private"}:
                    raise HTTPException(status_code=422, detail="Markaz mulk turi noto'g'ri")
                if operator_model not in {"center", "independent_tutor"}:
                    raise HTTPException(status_code=422, detail="Markaz operator turi noto'g'ri")
                if operator_model == "independent_tutor" and ownership_type != "private":
                    raise HTTPException(
                        status_code=422,
                        detail="Mustaqil repetitor faqat xususiy bo'ladi",
                    )
            if request.step == "staff":
                relationship = str(request.payload.get("relationship") or relationship)
                if relationship not in {
                    "owner", "founder", "director", "administrator", "teacher"
                }:
                    raise HTTPException(status_code=422, detail="Markazdagi rolingiz noto'g'ri")
            if (
                operator_model == "center"
                and relationship not in {"owner", "founder", "director"}
            ):
                raise HTTPException(
                    status_code=422,
                    detail="Oddiy markaz uchun mulkdor, ta'sischi yoki direktor kerak",
                )
            if (
                ownership_type == "public"
                and relationship != "director"
            ):
                raise HTTPException(
                    status_code=422,
                    detail="Davlat markazi uchun mas'ul direktor kerak",
                )
            if (
                operator_model == "independent_tutor"
                and relationship not in {"owner", "teacher"}
            ):
                raise HTTPException(
                    status_code=422,
                    detail="Mustaqil repetitor uchun mulkdor yoki o'qituvchi rolini tanlang",
                )
            cur.execute(
                """UPDATE center_setup_drafts SET
                     payload=%s::jsonb,current_step=%s,
                     ownership_type=%s,operator_model=%s,relationship=%s,
                     version=version+1,
                     updated_at=NOW()
                   WHERE id=%s
                   RETURNING id,current_step,payload,status,version,
                     ownership_type,operator_model,relationship,updated_at""",
                (
                    json.dumps(payload, ensure_ascii=False, default=str),
                    request.step, ownership_type, operator_model, relationship,
                    draft_id,
                ),
            )
            updated = cur.fetchone()
        return {"draft": updated}

    def step_rows(payload: dict[str, Any], key: str) -> list[Any]:
        value = payload.get(key) or []
        if isinstance(value, dict):
            value = value.get(key) or value.get("items") or []
        return list(value) if isinstance(value, list) else []

    def draft_preview_payload(draft: dict[str, Any]) -> dict[str, Any]:
        payload = dict(draft.get("payload") or {})
        identity = dict(payload.get("identity") or {})
        problems: list[dict[str, str]] = []
        if len(str(identity.get("name") or "").strip()) < 2:
            problems.append({"field": "identity.name", "message": "Markaz nomini kiriting"})
        if (
            draft["ownership_type"] == "public"
            and draft["relationship"] in {"owner", "founder"}
        ):
            problems.append(
                {
                    "field": "staff.relationship",
                    "message": "Davlat markazida mulkdor yoki ta'sischi roli bo'lmaydi",
                }
            )
        if (
            draft["operator_model"] == "independent_tutor"
            and draft["ownership_type"] != "private"
        ):
            problems.append(
                {
                    "field": "identity.operator_model",
                    "message": "Mustaqil repetitor xususiy operator bo'ladi",
                }
            )
        if (
            draft["operator_model"] == "independent_tutor"
            and draft["relationship"] not in {"owner", "teacher"}
        ):
            problems.append(
                {
                    "field": "staff.relationship",
                    "message": "Mustaqil repetitor mulkdor yoki o'qituvchi bo'ladi",
                }
            )
        if (
            draft["operator_model"] == "center"
            and draft["relationship"] not in {"owner", "founder", "director"}
        ):
            problems.append(
                {
                    "field": "staff.relationship",
                    "message": "Oddiy markazga mulkdor, ta'sischi yoki direktor kerak",
                }
            )
        if (
            draft["ownership_type"] == "public"
            and draft["relationship"] != "director"
        ):
            problems.append(
                {
                    "field": "staff.relationship",
                    "message": "Davlat markazini mas'ul direktor ariza qiladi",
                }
            )
        branches = step_rows(payload, "branches")
        if draft["operator_model"] == "center" and not branches:
            problems.append(
                {"field": "branches", "message": "Kamida bitta filial kiriting"}
            )
        for index, branch in enumerate(branches):
            field = f"branches.{index}"
            if not isinstance(branch, dict):
                problems.append(
                    {"field": field, "message": "Filial qatori noto'g'ri"}
                )
                continue
            if len(str(branch.get("name") or "").strip()) < 2:
                problems.append(
                    {"field": f"{field}.name", "message": "Filial nomini kiriting"}
                )
            try:
                normalize_work_days(
                    branch.get("work_days"), f"{field}.work_days"
                )
                work_start = branch.get("work_start")
                work_end = branch.get("work_end")
                parsed_start = (
                    parse_time(str(work_start), f"{field}.work_start")
                    if work_start else None
                )
                parsed_end = (
                    parse_time(str(work_end), f"{field}.work_end")
                    if work_end else None
                )
                if parsed_start and parsed_end and parsed_start >= parsed_end:
                    raise ValueError("work range")
            except (HTTPException, TypeError, ValueError):
                problems.append(
                    {
                        "field": field,
                        "message": "Filial ish kunlari yoki vaqti noto'g'ri",
                    }
                )
        subjects = step_rows(payload, "subjects")
        if not subjects:
            problems.append(
                {"field": "subjects", "message": "Kamida bitta fan kiriting"}
            )
        subject_names = {
            str(
                item if isinstance(item, str)
                else item.get("name") if isinstance(item, dict) else ""
            ).strip().casefold()
            for item in subjects
        }
        for index, room in enumerate(step_rows(payload, "rooms")):
            field = f"rooms.{index}"
            if not isinstance(room, dict):
                problems.append({"field": field, "message": "Xona qatori noto'g'ri"})
                continue
            try:
                branch_index = int(room.get("branch_index") or 0)
                capacity = int(room.get("capacity") or 1)
                if (
                    branch_index < 0 or branch_index >= max(len(branches), 1)
                    or capacity < 1 or capacity > 5000
                ):
                    raise ValueError("room range")
            except (TypeError, ValueError):
                problems.append(
                    {
                        "field": field,
                        "message": "Xona filiali yoki sig'imi noto'g'ri",
                    }
                )
        for index, course in enumerate(step_rows(payload, "courses")):
            field = f"courses.{index}"
            if not isinstance(course, dict):
                problems.append({"field": field, "message": "Kurs qatori noto'g'ri"})
                continue
            if len(str(course.get("name") or "").strip()) < 2:
                problems.append(
                    {"field": f"{field}.name", "message": "Kurs nomini kiriting"}
                )
            subject_name = str(course.get("subject_name") or "").strip().casefold()
            if subject_name not in subject_names:
                problems.append(
                    {"field": f"{field}.subject_name", "message": "Kurs fanini tanlang"}
                )
            try:
                capacity = int(course.get("capacity") or 20)
                duration = int(course.get("duration_minutes") or 90)
                price = Decimal(str(course.get("monthly_price") or 0))
                weekdays = [int(day) for day in (course.get("weekdays") or [])]
                teacher = course.get("teacher_user_id")
                if teacher is not None:
                    int(teacher)
                if (
                    capacity < 1 or capacity > 5000
                    or duration < 15 or duration > 480
                    or price < 0
                    or len(weekdays) > 7
                    or len(set(weekdays)) != len(weekdays)
                    or any(day < 1 or day > 7 for day in weekdays)
                ):
                    raise ValueError("course range")
            except (TypeError, ValueError, ArithmeticError):
                problems.append(
                    {
                        "field": field,
                        "message": "Kurs sig'imi, narxi, vaqti yoki kunlari noto'g'ri",
                    }
                )
        billing = payload.get("billing")
        if isinstance(billing, dict) and billing.get("enabled"):
            try:
                due_day = int(billing.get("due_day") or 10)
                if due_day < 1 or due_day > 28:
                    raise ValueError("due day")
            except (TypeError, ValueError):
                problems.append(
                    {
                        "field": "billing.due_day",
                        "message": "To'lov kuni 1 dan 28 gacha bo'lsin",
                    }
                )
        summary = {
            "branch_count": len(branches),
            "room_count": len(step_rows(payload, "rooms")),
            "subject_count": len(subjects),
            "course_count": len(step_rows(payload, "courses")),
        }
        warnings: list[str] = []
        if not step_rows(payload, "courses"):
            warnings.append("Kursni hozir yaratmasangiz, keyin Kurslar bo'limidan qo'shasiz.")
        return {
            "identity": identity,
            "summary": summary,
            "ownership_type": draft["ownership_type"],
            "operator_model": draft["operator_model"],
            "warnings": warnings,
            "problems": problems,
            "ready": not problems,
        }

    @router.get("/onboarding/drafts/{draft_id}/preview")
    def preview_draft(
        draft_id: int, user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            cur.execute(
                """SELECT * FROM center_setup_drafts
                   WHERE id=%s AND creator_user_id=%s AND status='draft'
                     AND expires_at>NOW()""",
                (draft_id, user_id),
            )
            draft = cur.fetchone()
            if not draft:
                raise HTTPException(status_code=404, detail="Faol qoralama topilmadi")
        return {"preview": draft_preview_payload(dict(draft)), "version": draft["version"]}

    @router.post("/onboarding/drafts/{draft_id}/commit")
    def commit_draft(
        draft_id: int, request: HumanConfirm,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(request.confirmation, "Markazni yaratish uchun inson tasdig'i kerak.")
        if request.expected_version is None:
            raise HTTPException(status_code=422, detail="expected_version kerak")
        with database() as (_, cur):
            ensure_schema(cur)
            cur.execute(
                """SELECT * FROM center_setup_drafts
                   WHERE id=%s AND creator_user_id=%s AND status='draft'
                     AND expires_at>NOW() FOR UPDATE""",
                (draft_id, user_id),
            )
            draft = cur.fetchone()
            if not draft or int(draft["version"]) != request.expected_version:
                raise HTTPException(status_code=409, detail="Qoralama versiyasi o'zgargan")
            preview = draft_preview_payload(dict(draft))
            if not preview["ready"]:
                raise HTTPException(status_code=422, detail=preview)
            payload = dict(draft["payload"] or {})
            identity = dict(payload["identity"])
            public = draft["ownership_type"] == "public"
            cur.execute(
                """INSERT INTO learning_contexts(
                     context_type,name,owner_user_id,region,district,active,metadata
                   ) VALUES(
                     'learning_center',%s,%s,%s,%s,TRUE,%s::jsonb
                   ) RETURNING id,name,region,district""",
                (
                    str(identity["name"]).strip(), user_id,
                    identity.get("region"), identity.get("district"),
                    json.dumps(
                        {
                            "source": "center_v2",
                            "operator_model": draft["operator_model"],
                        },
                        ensure_ascii=False,
                    ),
                ),
            )
            center = cur.fetchone()
            context_id = int(center["id"])
            cur.execute(
                """INSERT INTO center_profiles(
                     context_id,ownership_type,operator_model,onboarding_status,
                     verification_status,timezone,settings
                   ) VALUES(%s,%s,%s,%s,%s,%s,%s::jsonb)""",
                (
                    context_id, draft["ownership_type"], draft["operator_model"],
                    "pending_verification" if public else "active",
                    "pending" if public else "unverified",
                    identity.get("timezone") or "Asia/Tashkent",
                    json.dumps(
                        {"phone": identity.get("phone"), "setup_version": 2},
                        ensure_ascii=False,
                    ),
                ),
            )
            branch_ids: list[int] = []
            branches = step_rows(payload, "branches")
            if not branches and draft["operator_model"] == "independent_tutor":
                branches = [{"name": "Asosiy", "address": identity.get("address")}]
            for branch in branches[:100]:
                if not isinstance(branch, dict):
                    raise HTTPException(
                        status_code=422, detail="branches qatori obyekt bo'lishi kerak"
                    )
                name = str(branch.get("name") or "").strip()
                if len(name) < 2:
                    raise HTTPException(
                        status_code=422, detail="Har bir filial nomini kiriting"
                    )
                work_days = normalize_work_days(
                    branch.get("work_days"), "branches.work_days"
                )
                cur.execute(
                    """INSERT INTO center_branches(
                         context_id,name,address,phone,work_start,work_end,work_days
                       ) VALUES(%s,%s,%s,%s,%s,%s,%s) RETURNING id""",
                    (
                        context_id, name, branch.get("address"),
                        branch.get("phone"), branch.get("work_start"),
                        branch.get("work_end"), work_days,
                    ),
                )
                branch_ids.append(int(cur.fetchone()["id"]))
            subject_ids: dict[str, int] = {}
            for raw_subject in step_rows(payload, "subjects")[:200]:
                subject = (
                    {"name": raw_subject}
                    if isinstance(raw_subject, str) else dict(raw_subject)
                )
                code = re.sub(
                    r"[^A-Z0-9_]+", "_",
                    str(subject.get("code") or subject.get("name") or "").upper(),
                ).strip("_")[:40]
                name = str(subject.get("name") or "").strip()
                if not code or len(name) < 2:
                    continue
                cur.execute(
                    """INSERT INTO center_subjects(
                         context_id,code,name,supports_latex,formula_metadata
                       ) VALUES(%s,%s,%s,%s,%s::jsonb)
                       ON CONFLICT(context_id,code) DO UPDATE
                         SET name=EXCLUDED.name,updated_at=NOW()
                       RETURNING id""",
                    (
                        context_id, code, name,
                        bool(subject.get("supports_latex", False)),
                        json.dumps(subject.get("formula_metadata") or {}),
                    ),
                )
                subject_ids[name.casefold()] = int(cur.fetchone()["id"])
            room_ids: list[int] = []
            for raw_room in step_rows(payload, "rooms")[:500]:
                if not isinstance(raw_room, dict):
                    continue
                branch_index = int(raw_room.get("branch_index") or 0)
                if not branch_ids or branch_index < 0 or branch_index >= len(branch_ids):
                    continue
                room_name = str(raw_room.get("name") or "").strip()
                if not room_name:
                    continue
                cur.execute(
                    """INSERT INTO center_rooms(
                         context_id,branch_id,name,room_type,capacity,metadata
                       ) VALUES(%s,%s,%s,%s,%s,%s::jsonb)
                       ON CONFLICT(context_id,branch_id,name) DO UPDATE SET
                         room_type=EXCLUDED.room_type,capacity=EXCLUDED.capacity,
                         updated_at=NOW()
                       RETURNING id""",
                    (
                        context_id, branch_ids[branch_index], room_name,
                        raw_room.get("room_type") or "classroom",
                        raw_room.get("capacity"),
                        json.dumps({"source": "onboarding"}),
                    ),
                )
                room_ids.append(int(cur.fetchone()["id"]))
            role = str(draft["relationship"])
            if draft["operator_model"] == "independent_tutor" and role == "teacher":
                upsert_role(
                    cur, context_id=context_id, user_id=user_id, role_key="owner",
                    branch_id=None, status="active", approved_by=user_id,
                )
                upsert_role(
                    cur, context_id=context_id, user_id=user_id, role_key="teacher",
                    branch_id=None, status="active", approved_by=user_id,
                )
            else:
                upsert_role(
                    cur, context_id=context_id, user_id=user_id, role_key=role,
                    branch_id=None, status="active", approved_by=user_id,
                )
            course_ids: list[int] = []
            billing_course_prices: list[tuple[int, Decimal]] = []
            for raw_course in step_rows(payload, "courses")[:100]:
                if not isinstance(raw_course, dict):
                    continue
                course_name = str(raw_course.get("name") or "").strip()
                subject_name = str(raw_course.get("subject_name") or "").strip()
                subject_id = subject_ids.get(subject_name.casefold())
                if not course_name or subject_id is None:
                    continue
                delivery = str(raw_course.get("delivery_mode") or "offline")
                branch_id = None if delivery == "online_live" else (
                    branch_ids[0] if branch_ids else None
                )
                if delivery != "online_live" and branch_id is None:
                    continue
                course_type = str(
                    raw_course.get("course_type")
                    or raw_course.get("format_key")
                    or "group"
                )
                if course_type not in {
                    "group", "individual", "intensive", "club", "exam_prep"
                }:
                    course_type = "group"
                capacity = 1 if course_type == "individual" else max(
                    1, min(5000, int(raw_course.get("capacity") or 20))
                )
                framework = str(raw_course.get("target_framework") or "custom")
                if framework == "general":
                    framework = "custom"
                level_from = level_to = None
                target_score = None
                target_components: dict[str, Any] = {}
                if framework == "cefr":
                    level = str(
                        raw_course.get("cefr_level")
                        or raw_course.get("level_from") or "A1"
                    ).upper()
                    if level not in CEFR_LEVELS:
                        level = "A1"
                    level_from = level_to = level
                elif framework == "ielts":
                    targets = dict(
                        raw_course.get("ielts_targets")
                        or raw_course.get("target_components") or {}
                    )
                    target_score = Decimal(str(targets.get("overall") or 5))
                    target_components = {
                        key: targets.get(key, target_score) for key in IELTS_COMPONENTS
                    }
                ielts_test_type = str(
                    raw_course.get("ielts_test_type")
                    or raw_course.get("ielts_type")
                    or "academic"
                )
                if ielts_test_type not in {"academic", "general"}:
                    ielts_test_type = "academic"
                teacher_id = (
                    user_id if draft["operator_model"] == "independent_tutor"
                    else raw_course.get("teacher_user_id")
                )
                if teacher_id is not None:
                    teacher_id = int(teacher_id)
                    upsert_role(
                        cur, context_id=context_id, user_id=teacher_id,
                        role_key="teacher", branch_id=branch_id, status="active",
                        approved_by=user_id,
                    )
                    cur.execute(
                        """INSERT INTO center_staff_subjects(
                             context_id,teacher_user_id,subject_id
                           ) VALUES(%s,%s,%s)
                           ON CONFLICT(context_id,teacher_user_id,subject_id)
                           DO UPDATE SET active=TRUE""",
                        (context_id, teacher_id, subject_id),
                    )
                cur.execute(
                    """INSERT INTO course_groups(
                         context_id,group_type,delivery_mode,name,subject,
                         teacher_user_id,metadata
                       ) VALUES(%s,'center_course',%s,%s,%s,%s,%s::jsonb)
                       RETURNING id""",
                    (
                        context_id, delivery, course_name, subject_name, teacher_id,
                        json.dumps({"source": "center_onboarding"}),
                    ),
                )
                group_id = int(cur.fetchone()["id"])
                cur.execute(
                    """INSERT INTO center_courses(
                         context_id,group_id,branch_id,subject_id,teacher_user_id,
                         name,course_type,delivery_mode,target_framework,
                         level_from,level_to,target_score,target_components,
                         capacity,duration_minutes,status,metadata
                       ) VALUES(
                         %s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s::jsonb,
                         %s,%s,%s,%s::jsonb
                       ) RETURNING id""",
                    (
                        context_id, group_id, branch_id, subject_id, teacher_id,
                        course_name, course_type, delivery, framework,
                        level_from, level_to, target_score,
                        json.dumps(target_components, default=str),
                        capacity, int(raw_course.get("duration_minutes") or 90),
                        "active" if teacher_id and not public else "draft",
                        json.dumps(
                            {
                                "sessions_per_week": raw_course.get("sessions_per_week"),
                                "schedule_hint": {
                                    "weekdays": raw_course.get("weekdays") or [],
                                    "starts_at": raw_course.get("starts_at"),
                                    "start_date": raw_course.get("start_date"),
                                    "end_date": raw_course.get("end_date"),
                                    "deferred_until_teacher_assignment": (
                                        teacher_id is None
                                    ),
                                },
                                "level_label": raw_course.get("level_label"),
                                "ielts_test_type": (
                                    ielts_test_type if framework == "ielts" else None
                                ),
                            },
                            ensure_ascii=False,
                        ),
                    ),
                )
                course_id = int(cur.fetchone()["id"])
                course_ids.append(course_id)
                if teacher_id:
                    for weekday in list(raw_course.get("weekdays") or [])[:7]:
                        start = parse_time(
                            str(raw_course.get("starts_at") or "17:00"), "starts_at"
                        )
                        duration = int(raw_course.get("duration_minutes") or 90)
                        total_minutes = (
                            start.hour * 60 + start.minute + duration
                        )
                        if total_minutes >= 24 * 60:
                            raise HTTPException(
                                status_code=422,
                                detail="Dars yarim tundan o'tmasin",
                            )
                        end = time(total_minutes // 60, total_minutes % 60)
                        require_branch_workday(
                            cur, context_id=context_id, branch_id=branch_id,
                            weekday=int(weekday), starts_at=start, ends_at=end,
                        )
                        ensure_teacher_available(
                            cur, context_id=context_id,
                            teacher_user_id=teacher_id,
                            schedule_kind="weekly", weekday=int(weekday),
                            starts_at=start, ends_at=end,
                        )
                        ensure_teacher_weekly_capacity(
                            cur, context_id=context_id,
                            teacher_user_id=int(teacher_id),
                            branch_id=branch_id, starts_at=start, ends_at=end,
                        )
                        cur.execute(
                            """INSERT INTO center_schedule_slots(
                                 context_id,course_id,group_id,branch_id,
                                 teacher_user_id,schedule_kind,weekday,
                                 starts_at,ends_at,status,created_by_user_id,
                                 metadata
                               ) VALUES(
                                 %s,%s,%s,%s,%s,'weekly',%s,%s,%s,'draft',%s,
                                 '{"source":"onboarding"}'
                               )""",
                            (
                                context_id, course_id, group_id, branch_id,
                                teacher_id, int(weekday), start, end, user_id,
                            ),
                        )
                price = Decimal(str(raw_course.get("monthly_price") or 0))
                if price > 0:
                    billing_course_prices.append((course_id, price))
            billing = dict(payload.get("billing") or {})
            if billing.get("enabled"):
                billing_cycle = str(
                    billing.get("billing_period") or "monthly"
                )
                if billing_cycle == "lesson":
                    billing_cycle = "per_lesson"
                if billing_cycle not in {
                    "one_time", "per_lesson", "weekly", "monthly", "course"
                }:
                    raise HTTPException(
                        status_code=422, detail="To'lov davri noto'g'ri"
                    )
                for course_id, price in billing_course_prices:
                    cur.execute(
                        """INSERT INTO center_billing_plans(
                             context_id,course_id,name,amount,billing_cycle,
                             currency,billing_day,created_by_user_id
                           ) VALUES(%s,%s,%s,%s,%s,%s,%s,%s)""",
                        (
                            context_id, course_id, "Oylik to'lov", price,
                            billing_cycle,
                            billing.get("currency") or "UZS",
                            int(billing.get("due_day") or 10), user_id,
                        ),
                    )
            cur.execute(
                """UPDATE center_setup_drafts SET
                     status='confirmed',confirmed_context_id=%s,updated_at=NOW()
                   WHERE id=%s""",
                (context_id, draft_id),
            )
            audit(cur, context_id, user_id, "center.create", "center", context_id)
        return {
            "context_id": context_id,
            "status": "pending_verification" if public else "active",
            "verification_status": "pending" if public else "unverified",
            "workspace": {
                **dict(center),
                "ownership_type": draft["ownership_type"],
                "operator_model": draft["operator_model"],
                "branch_ids": branch_ids,
                "room_ids": room_ids,
                "course_ids": course_ids,
            },
        }

    @router.get("/admin/verifications")
    def list_verifications(
        status: Literal["pending", "verified", "rejected"] = Query(default="pending"),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=50, ge=1, le=200),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            if not system_admin(cur, user_id):
                raise HTTPException(
                    status_code=403,
                    detail="Davlat markazini faqat tizim administratori tekshiradi",
                )
            cur.execute(
                """SELECT c.id context_id,c.name,c.region,c.district,c.owner_user_id,
                          p.ownership_type,p.operator_model,p.onboarding_status,
                          p.verification_status,p.created_at,p.updated_at
                   FROM center_profiles p
                   JOIN learning_contexts c ON c.id=p.context_id
                   WHERE p.ownership_type='public'
                     AND p.verification_status=%s AND c.id>%s
                   ORDER BY c.id LIMIT %s""",
                (status, after_id or 0, limit + 1),
            )
            rows = [dict(row) for row in cur.fetchall()]
        return {
            "items": rows[:limit],
            "next_cursor": (
                rows[limit - 1]["context_id"] if len(rows) > limit else None
            ),
        }

    @router.post("/admin/verifications/{context_id}/decision")
    def decide_verification(
        context_id: int,
        request: VerificationDecision,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(
            request.confirmation,
            "Davlat markazi tekshiruv qarori uchun inson tasdig'i kerak.",
        )
        with database() as (_, cur):
            ensure_schema(cur)
            if not system_admin(cur, user_id):
                raise HTTPException(
                    status_code=403,
                    detail="Markaz o'zini o'zi tasdiqlay olmaydi",
                )
            cur.execute(
                """SELECT p.*,c.active,c.name
                   FROM center_profiles p
                   JOIN learning_contexts c ON c.id=p.context_id
                   WHERE p.context_id=%s AND p.ownership_type='public'
                   FOR UPDATE OF p,c""",
                (context_id,),
            )
            profile = cur.fetchone()
            if not profile:
                raise HTTPException(
                    status_code=404, detail="Davlat markazi arizasi topilmadi"
                )
            if profile["verification_status"] != "pending":
                raise HTTPException(
                    status_code=409,
                    detail="Bu tekshiruv arizasi bo'yicha qaror avval berilgan",
                )
            approved = request.decision == "verified"
            onboarding_status = "active" if approved else "suspended"
            cur.execute(
                """UPDATE center_profiles SET verification_status=%s,
                     onboarding_status=%s,
                     settings=settings||%s::jsonb,updated_at=NOW()
                   WHERE context_id=%s
                   RETURNING context_id,ownership_type,operator_model,
                     onboarding_status,verification_status,updated_at""",
                (
                    request.decision, onboarding_status,
                    json.dumps(
                        {
                            "verification_decided_by": user_id,
                            "verification_note": request.note,
                            "verification_decided_at": datetime.now(
                                timezone.utc
                            ).isoformat(),
                        },
                        ensure_ascii=False,
                    ),
                    context_id,
                ),
            )
            updated = cur.fetchone()
            cur.execute(
                """UPDATE learning_contexts SET active=%s,updated_at=NOW()
                   WHERE id=%s""",
                (approved, context_id),
            )
            audit(
                cur, context_id, user_id, "center.verification.decision",
                "center", context_id,
                {"decision": request.decision, "note": request.note},
            )
        return {"item": updated, "active": approved}

    @router.get("/workspaces")
    def workspaces(
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=30, ge=1, le=100),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            if system_admin(cur, user_id):
                cur.execute(
                    """SELECT c.id context_id,c.name,c.region,c.district,
                              p.ownership_type,p.operator_model,
                              p.onboarding_status,p.verification_status,
                              p.legacy_center_id,
                              ARRAY['system_admin']::TEXT[] roles
                       FROM learning_contexts c
                       JOIN center_profiles p ON p.context_id=c.id
                       WHERE c.id>%s ORDER BY c.id LIMIT %s""",
                    (after_id or 0, limit + 1),
                )
            else:
                cur.execute(
                    """SELECT c.id context_id,c.name,c.region,c.district,
                              p.ownership_type,p.operator_model,
                              p.onboarding_status,p.verification_status,
                              p.legacy_center_id,
                              array_agg(DISTINCT r.role_key) roles
                       FROM center_role_assignments r
                       JOIN learning_contexts c ON c.id=r.context_id
                       JOIN center_profiles p ON p.context_id=c.id
                       WHERE r.user_id=%s AND r.status='active'
                         AND r.starts_at<=NOW()
                         AND (r.ends_at IS NULL OR r.ends_at>NOW())
                         AND c.id>%s
                       GROUP BY c.id,c.name,c.region,c.district,p.ownership_type,
                                p.operator_model,p.onboarding_status,
                                p.verification_status,p.legacy_center_id
                       ORDER BY c.id LIMIT %s""",
                    (user_id, after_id or 0, limit + 1),
                )
            rows = [dict(row) for row in cur.fetchall()]
        for row in rows:
            row["permissions"] = sorted(permissions_for(set(row["roles"])))
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["context_id"] if len(rows) > limit else None,
        }

    @router.get("/dashboard")
    def dashboard(
        context_id: int = Query(ge=1),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            roles = require_roles(cur, context_id, user_id, VIEW_ROLES)
            cur.execute(
                """SELECT c.id context_id,c.name,c.region,c.district,
                          p.ownership_type,p.operator_model,p.onboarding_status,
                          p.verification_status,p.timezone,p.default_currency,
                          p.work_days,p.settings
                   FROM learning_contexts c
                   JOIN center_profiles p ON p.context_id=c.id
                   WHERE c.id=%s""",
                (context_id,),
            )
            center = dict(cur.fetchone())
            broad_dashboard_roles = MANAGER_ROLES | {
                "receptionist", "accountant", "methodist",
            }
            dashboard_global_scope = False
            dashboard_branch_ids: set[int] = set()
            if roles & broad_dashboard_roles or "system_admin" in roles:
                _, global_scope, branch_ids = branch_scope(
                    cur, context_id, user_id, broad_dashboard_roles
                )
                dashboard_global_scope = global_scope
                dashboard_branch_ids = branch_ids
                branch_filter = None if global_scope else list(branch_ids)
                cur.execute(
                    """SELECT
                        (SELECT COUNT(*) FROM center_branches b
                         WHERE b.context_id=%s AND b.active
                           AND (%s IS NULL OR b.id=ANY(%s))) branches,
                        (SELECT COUNT(*) FROM center_rooms r
                         WHERE r.context_id=%s AND r.active
                           AND (%s IS NULL OR r.branch_id=ANY(%s))) rooms,
                        (SELECT COUNT(*) FROM center_courses c
                         WHERE c.context_id=%s AND c.status='active'
                           AND (%s IS NULL OR c.branch_id=ANY(%s))) active_courses,
                        (SELECT COUNT(*) FROM center_enrollments e
                         JOIN center_courses c ON c.id=e.course_id
                         WHERE e.context_id=%s AND e.status='active'
                           AND (%s IS NULL OR c.branch_id=ANY(%s))) active_students,
                        (SELECT COUNT(DISTINCT r.user_id)
                         FROM center_role_assignments r
                         WHERE r.context_id=%s AND r.status='active'
                           AND r.starts_at<=NOW()
                           AND (r.ends_at IS NULL OR r.ends_at>NOW())
                           AND r.role_key NOT IN ('student','parent')
                           AND (%s IS NULL OR r.branch_id=ANY(%s))) staff,
                        (SELECT COALESCE(SUM(i.amount-i.paid_amount),0)
                         FROM center_invoices i
                         JOIN center_courses c ON c.id=i.course_id
                         WHERE i.context_id=%s
                           AND i.status IN ('unpaid','partial')
                           AND (%s IS NULL OR c.branch_id=ANY(%s))) debt""",
                    (
                        context_id, branch_filter, branch_filter,
                        context_id, branch_filter, branch_filter,
                        context_id, branch_filter, branch_filter,
                        context_id, branch_filter, branch_filter,
                        context_id, branch_filter, branch_filter,
                        context_id, branch_filter, branch_filter,
                    ),
                )
            elif "teacher" in roles:
                cur.execute(
                    """SELECT
                        0::BIGINT branches,0::BIGINT rooms,
                        (SELECT COUNT(*) FROM center_courses c
                         WHERE c.context_id=%s AND c.status='active'
                           AND c.teacher_user_id=%s
                           AND EXISTS(
                             SELECT 1 FROM center_role_assignments tr
                             WHERE tr.context_id=c.context_id
                               AND tr.user_id=%s
                               AND tr.role_key='teacher'
                               AND tr.status='active'
                               AND tr.starts_at<=NOW()
                               AND (tr.ends_at IS NULL OR tr.ends_at>NOW())
                               AND (
                                 tr.branch_id IS NULL
                                 OR (
                                   c.branch_id IS NOT NULL
                                   AND tr.branch_id=c.branch_id
                                 )
                               )
                           )) active_courses,
                        (SELECT COUNT(DISTINCT e.student_user_id)
                         FROM center_enrollments e
                         JOIN center_courses c
                           ON c.context_id=e.context_id AND c.id=e.course_id
                         WHERE e.context_id=%s AND e.status='active'
                           AND c.teacher_user_id=%s
                           AND EXISTS(
                             SELECT 1 FROM center_role_assignments tr
                             WHERE tr.context_id=c.context_id
                               AND tr.user_id=%s
                               AND tr.role_key='teacher'
                               AND tr.status='active'
                               AND tr.starts_at<=NOW()
                               AND (tr.ends_at IS NULL OR tr.ends_at>NOW())
                               AND (
                                 tr.branch_id IS NULL
                                 OR (
                                   c.branch_id IS NOT NULL
                                   AND tr.branch_id=c.branch_id
                                 )
                               )
                           )) active_students,
                        1::BIGINT staff,0::NUMERIC debt""",
                    (
                        context_id, user_id, user_id,
                        context_id, user_id, user_id,
                    ),
                )
            else:
                cur.execute(
                    """SELECT
                        0::BIGINT branches,0::BIGINT rooms,
                        (SELECT COUNT(DISTINCT e.course_id)
                         FROM center_enrollments e
                         WHERE e.context_id=%s AND e.status IN ('active','completed')
                           AND (
                             e.starts_on IS NULL
                             OR e.starts_on<=
                               (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                           )
                           AND (
                             e.student_user_id=%s OR EXISTS(
                               SELECT 1 FROM center_parent_links p
                               WHERE p.context_id=e.context_id
                                 AND p.student_user_id=e.student_user_id
                                 AND p.parent_user_id=%s AND p.status='active'
                             )
                           )) active_courses,
                        (SELECT COUNT(DISTINCT e.student_user_id)
                         FROM center_enrollments e
                         WHERE e.context_id=%s AND e.status='active'
                           AND (
                             e.starts_on IS NULL
                             OR e.starts_on<=
                               (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                           )
                           AND (
                             e.student_user_id=%s OR EXISTS(
                               SELECT 1 FROM center_parent_links p
                               WHERE p.context_id=e.context_id
                                 AND p.student_user_id=e.student_user_id
                                 AND p.parent_user_id=%s AND p.status='active'
                             )
                           )) active_students,
                        0::BIGINT staff,
                        (SELECT COALESCE(SUM(i.amount-i.paid_amount),0)
                         FROM center_invoices i
                         WHERE i.context_id=%s
                           AND i.status IN ('unpaid','partial')
                           AND (
                             i.student_user_id=%s OR EXISTS(
                               SELECT 1 FROM center_parent_links p
                               WHERE p.context_id=i.context_id
                                 AND p.student_user_id=i.student_user_id
                                 AND p.parent_user_id=%s AND p.status='active'
                             )
                           )) debt""",
                    (
                        context_id, user_id, user_id,
                        context_id, user_id, user_id,
                        context_id, user_id, user_id,
                    ),
                )
            counts = dict(cur.fetchone())
            can_view_staff_debt = bool(
                "system_admin" in roles or roles & FINANCE_VIEW_STAFF_ROLES
            )
            is_pure_learner_view = bool(
                roles & {"student", "parent"} and not roles & STAFF_ROLES
            )
            if not (can_view_staff_debt or is_pure_learner_view):
                counts["debt"] = Decimal("0")
            can_manage_parent_links = "system_admin" in roles
            if not can_manage_parent_links:
                cur.execute(
                    """SELECT EXISTS(
                         SELECT 1 FROM center_role_assignments r
                         WHERE r.context_id=%s AND r.user_id=%s
                           AND r.status='active' AND r.starts_at<=NOW()
                           AND (r.ends_at IS NULL OR r.ends_at>NOW())
                           AND r.branch_id IS NULL
                           AND r.role_key=ANY(%s)
                       ) allowed""",
                    (
                        context_id, user_id,
                        list(roles_for_permission("enrollments.manage")),
                    ),
                )
                can_manage_parent_links = bool(cur.fetchone()["allowed"])
            linked_children: list[dict[str, Any]] = []
            if "parent" in roles:
                cur.execute(
                    """SELECT p.student_user_id,u.full_name student_name
                       FROM center_parent_links p
                       JOIN users u ON u.user_id=p.student_user_id
                       WHERE p.context_id=%s AND p.parent_user_id=%s
                         AND p.status='active'
                       ORDER BY u.full_name,p.student_user_id""",
                    (context_id, user_id),
                )
                linked_children = [dict(row) for row in cur.fetchall()]
        permission_set = permissions_for(roles)
        menus = ["overview", "courses", "groups", "analytics"]
        if (
            "system_admin" in roles
            or roles & (SCHEDULE_ROLES | {"student", "parent"})
        ):
            menus.append("schedule")
        if (
            "system_admin" in roles
            or roles & (MANAGER_ROLES | {"methodist", "teacher", "student", "parent"})
        ):
            menus.extend(["lessons", "attendance", "assessments"])
        if (
            "system_admin" in roles
            or roles & (ENROLLMENT_ROLES | {"methodist", "teacher", "accountant"})
        ):
            menus.insert(3, "students")
        if "billing.view" in permission_set:
            menus.append("payments")
        if "staff.manage" in permission_set or "teacher" in roles:
            menus.append("staff")
        if "branches.manage" in permission_set:
            menus.append("settings")
        return {
            "center": center,
            "current_user_id": user_id,
            "roles": sorted(roles),
            "permissions": sorted(permission_set),
            "capabilities": {
                "can_manage_parent_links": can_manage_parent_links,
                "branch_scope": {
                    "global": dashboard_global_scope,
                    "branch_ids": sorted(dashboard_branch_ids),
                },
            },
            "linked_children": linked_children,
            "menus": menus,
            "counts": counts,
        }

    @router.get("/users/search")
    def search_users(
        context_id: int = Query(ge=1),
        q: str = Query(min_length=3, max_length=100),
        limit: int = Query(default=20, ge=1, le=50),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        literal = q.strip().replace("\\", "\\\\").replace("%", "\\%").replace("_", "\\_")
        account_id = int(q.strip()) if q.strip().isdigit() else None
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            require_roles(cur, context_id, user_id, MANAGER_ROLES | {"receptionist"})
            cur.execute(
                """SELECT u.user_id,u.full_name,
                          '#'||u.user_id::TEXT account_identifier,
                          EXISTS(
                            SELECT 1 FROM center_role_assignments r
                            WHERE r.context_id=%s AND r.user_id=u.user_id
                              AND r.status='active' AND r.starts_at<=NOW()
                              AND (r.ends_at IS NULL OR r.ends_at>NOW())
                          ) already_in_center
                   FROM users u
                   WHERE (
                     u.full_name ILIKE %s ESCAPE '\\'
                     OR (%s IS NOT NULL AND u.user_id=%s)
                   )
                   ORDER BY u.full_name,u.user_id LIMIT %s""",
                (
                    context_id, f"%{literal}%", account_id, account_id, limit,
                ),
            )
            rows = cur.fetchall()
        return {
            "items": rows,
            "next_cursor": None,
            "search_basis": "full_name_or_account_id",
            "phone_or_email_available": False,
        }

    @router.get("/branches")
    def list_branches(
        context_id: int = Query(ge=1),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=50, ge=1, le=200),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            _, global_scope, branch_ids = branch_scope(
                cur, context_id, user_id, STAFF_ROLES
            )
            branch_filter = None if global_scope else list(branch_ids)
            cur.execute(
                """SELECT id,name,address,phone,work_start,work_end,work_days,active
                   FROM center_branches
                   WHERE context_id=%s AND id>%s
                     AND (%s IS NULL OR id=ANY(%s))
                   ORDER BY id LIMIT %s""",
                (
                    context_id, after_id or 0, branch_filter, branch_filter,
                    limit + 1,
                ),
            )
            rows = cur.fetchall()
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
        }

    @router.post("/branches")
    def create_branch(
        request: BranchCreate, user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        start = parse_time(request.work_start, "work_start") if request.work_start else None
        end = parse_time(request.work_end, "work_end") if request.work_end else None
        if start and end and start >= end:
            raise HTTPException(status_code=422, detail="Filial tugash vaqti kechroq bo'lsin")
        work_days = normalize_work_days(request.work_days)
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            require_roles(
                cur, request.context_id, user_id, MANAGER_ROLES,
                require_global=True,
            )
            cur.execute(
                """INSERT INTO center_branches(
                     context_id,name,address,phone,work_start,work_end,work_days
                   ) VALUES(%s,%s,%s,%s,%s,%s,%s)
                   RETURNING id,name,address,phone,work_start,work_end,
                     work_days,active""",
                (
                    request.context_id, request.name.strip(), request.address,
                    request.phone, start, end, work_days,
                ),
            )
            branch = cur.fetchone()
            audit(cur, request.context_id, user_id, "branch.create", "branch", branch["id"])
        return {"item": branch}

    @router.get("/rooms")
    def list_rooms(
        context_id: int = Query(ge=1),
        branch_id: int | None = Query(default=None, ge=1),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=100, ge=1, le=300),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            _, global_scope, branch_ids = branch_scope(
                cur, context_id, user_id, STAFF_ROLES
            )
            if branch_id is not None and not (
                global_scope or branch_id in branch_ids
            ):
                raise HTTPException(status_code=403, detail="Boshqa filial xonalari yopiq")
            branch_filter = None if global_scope else list(branch_ids)
            cur.execute(
                """SELECT r.id,r.branch_id,b.name branch_name,r.name,r.room_type,
                          r.capacity,r.active,r.metadata
                   FROM center_rooms r
                   JOIN center_branches b ON b.id=r.branch_id
                   WHERE r.context_id=%s AND r.id>%s
                     AND (%s IS NULL OR r.branch_id=%s)
                     AND (%s IS NULL OR r.branch_id=ANY(%s))
                   ORDER BY r.id LIMIT %s""",
                (
                    context_id, after_id or 0, branch_id, branch_id,
                    branch_filter, branch_filter, limit + 1,
                ),
            )
            rows = cur.fetchall()
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
        }

    @router.post("/rooms")
    def create_room(
        request: RoomCreate, user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_encoded_size(
            request.metadata, maximum=50_000, label="Xona metadata"
        )
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            require_roles(
                cur, request.context_id, user_id, MANAGER_ROLES,
                branch_id=request.branch_id,
            )
            cur.execute(
                """INSERT INTO center_rooms(
                     context_id,branch_id,name,room_type,capacity,metadata
                   ) VALUES(%s,%s,%s,%s,%s,%s::jsonb)
                   RETURNING id,branch_id,name,room_type,capacity,active,metadata""",
                (
                    request.context_id, request.branch_id, request.name.strip(),
                    request.room_type, request.capacity,
                    json.dumps(request.metadata, ensure_ascii=False),
                ),
            )
            room = cur.fetchone()
            audit(cur, request.context_id, user_id, "room.create", "room", room["id"])
        return {"item": room}

    @router.get("/subjects")
    def list_subjects(
        context_id: int = Query(ge=1),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=100, ge=1, le=300),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            require_roles(cur, context_id, user_id, VIEW_ROLES)
            cur.execute(
                """SELECT id,code,name,supports_latex,formula_metadata,active
                   FROM center_subjects
                   WHERE context_id=%s AND id>%s
                   ORDER BY id LIMIT %s""",
                (context_id, after_id or 0, limit + 1),
            )
            rows = cur.fetchall()
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
        }

    @router.post("/subjects")
    def create_subject(
        request: SubjectCreate, user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_encoded_size(
            request.formula_metadata, maximum=50_000,
            label="Fan formula metadata",
        )
        code = re.sub(r"[^A-Z0-9_]+", "_", request.code.upper()).strip("_")
        if not code:
            raise HTTPException(status_code=422, detail="Fan kodi bo'sh")
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            require_roles(cur, request.context_id, user_id, MANAGER_ROLES | {"methodist"})
            cur.execute(
                """INSERT INTO center_subjects(
                     context_id,code,name,supports_latex,formula_metadata
                   ) VALUES(%s,%s,%s,%s,%s::jsonb)
                   RETURNING id,code,name,supports_latex,formula_metadata,active""",
                (
                    request.context_id, code, request.name.strip(),
                    request.supports_latex,
                    json.dumps(request.formula_metadata, ensure_ascii=False),
                ),
            )
            subject = cur.fetchone()
            audit(
                cur, request.context_id, user_id, "subject.create", "subject",
                subject["id"],
            )
        return {"item": subject}

    @router.get("/staff")
    def list_staff(
        context_id: int = Query(ge=1),
        role: str | None = Query(default=None, max_length=40),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=100, ge=1, le=300),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            if system_admin(cur, user_id):
                global_scope = True
                privileged_branch_ids: list[int] = []
                global_availability_manager = True
            else:
                _, assignments = active_roles(cur, context_id, user_id)
                academic_assignments = [
                    item
                    for item in assignments
                    if item["role_key"] in ACADEMIC_ROLES
                ]
                if not academic_assignments:
                    raise HTTPException(
                        status_code=403,
                        detail="Markaz vakolati yetarli emas",
                    )
                privileged_assignments = [
                    item
                    for item in academic_assignments
                    if item["role_key"] in (MANAGER_ROLES | {"methodist"})
                ]
                global_scope = any(
                    item["branch_id"] is None
                    for item in privileged_assignments
                )
                privileged_branch_ids = sorted(
                    {
                        int(item["branch_id"])
                        for item in privileged_assignments
                        if item["branch_id"] is not None
                    }
                )
                global_availability_manager = any(
                    item["branch_id"] is None
                    and item["role_key"] in MANAGER_ROLES
                    for item in assignments
                )
            cur.execute(
                """SELECT r.id,r.user_id,u.full_name,r.branch_id,
                          b.name branch_name,r.role_key,r.status,
                          r.employment_type,r.weekly_capacity_hours,
                          COALESCE(
                            array_agg(DISTINCT ss.subject_id)
                              FILTER(WHERE ss.subject_id IS NOT NULL),
                            ARRAY[]::BIGINT[]
                          ) subject_ids,
                          COALESCE(
                            array_agg(DISTINCT s.name)
                              FILTER(WHERE s.name IS NOT NULL),
                            ARRAY[]::TEXT[]
                          ) subject_names
                   FROM center_role_assignments r
                   JOIN users u ON u.user_id=r.user_id
                   LEFT JOIN center_branches b
                     ON b.context_id=r.context_id AND b.id=r.branch_id
                   LEFT JOIN center_staff_subjects ss
                     ON ss.context_id=r.context_id
                    AND ss.teacher_user_id=r.user_id AND ss.active
                   LEFT JOIN center_subjects s
                     ON s.context_id=ss.context_id AND s.id=ss.subject_id
                   WHERE r.context_id=%s AND r.id>%s
                     AND r.role_key NOT IN ('student','parent')
                     AND (%s OR r.user_id=%s OR r.branch_id=ANY(%s))
                     AND (%s IS NULL OR r.role_key=%s)
                   GROUP BY r.id,r.user_id,u.full_name,r.branch_id,b.name,r.role_key,
                            r.status,r.employment_type,r.weekly_capacity_hours
                   ORDER BY r.id LIMIT %s""",
                (
                    context_id, after_id or 0, global_scope, user_id,
                    privileged_branch_ids, role, role, limit + 1,
                ),
            )
            rows = [dict(row) for row in cur.fetchall()]
        for row in rows:
            row["can_edit_availability"] = bool(
                global_availability_manager
                or (
                    int(row["user_id"]) == user_id
                    and row["role_key"] == "teacher"
                    and row["status"] == "active"
                )
            )
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
        }

    @router.post("/staff")
    def assign_staff(
        request: StaffAssign, user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(request.confirmation, "Xodim rolini berish uchun inson tasdig'i kerak.")
        role = request.role_key.strip()
        if role not in set(ROLE_LABELS) - {"student", "parent"}:
            raise HTTPException(status_code=422, detail="Xodim roli noto'g'ri")
        if role in {"owner", "founder"} and request.branch_id is not None:
            raise HTTPException(
                status_code=422,
                detail="Mulkdor va ta'sischi roli faqat markaz miqyosida beriladi",
            )
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            actor_roles = require_roles(
                cur, request.context_id, user_id, MANAGER_ROLES,
                branch_id=request.branch_id,
                require_global=request.branch_id is None,
            )
            require_role_hierarchy(
                cur, context_id=request.context_id, actor_user_id=user_id,
                actor_roles=actor_roles, target_user_id=request.user_id,
                target_role=role,
            )
            cur.execute("SELECT 1 FROM users WHERE user_id=%s", (request.user_id,))
            if not cur.fetchone():
                raise HTTPException(status_code=404, detail="Foydalanuvchi topilmadi")
            assignment_id = upsert_role(
                cur, context_id=request.context_id, user_id=request.user_id,
                role_key=role, branch_id=request.branch_id, status="active",
                approved_by=user_id,
                employment_type=request.employment_type,
                weekly_capacity_hours=request.weekly_capacity_hours,
            )
            if role in {"teacher", "methodist"}:
                for subject_id in sorted(set(request.subject_ids)):
                    cur.execute(
                        """INSERT INTO center_staff_subjects(
                             context_id,teacher_user_id,subject_id
                           ) VALUES(%s,%s,%s)
                           ON CONFLICT(context_id,teacher_user_id,subject_id)
                           DO UPDATE SET active=TRUE""",
                        (request.context_id, request.user_id, subject_id),
                    )
            audit(
                cur, request.context_id, user_id, "staff.assign",
                "role_assignment", assignment_id,
                {
                    "role": role,
                    "target_user_id": request.user_id,
                    "employment_type": request.employment_type,
                    "weekly_capacity_hours": request.weekly_capacity_hours,
                },
            )
        return {
            "assignment_id": assignment_id,
            "status": "active",
            "employment_type": request.employment_type,
            "weekly_capacity_hours": request.weekly_capacity_hours,
        }

    @router.post("/staff/{assignment_id}/status")
    def change_staff_status(
        assignment_id: int,
        request: StaffStatusChange,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(
            request.confirmation,
            "Xodim vakolatini to'xtatish uchun inson tasdig'i kerak.",
        )
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            cur.execute(
                """SELECT * FROM center_role_assignments
                   WHERE id=%s AND context_id=%s FOR UPDATE""",
                (assignment_id, request.context_id),
            )
            assignment = cur.fetchone()
            if not assignment:
                raise HTTPException(status_code=404, detail="Xodim roli topilmadi")
            actor_roles = require_roles(
                cur, request.context_id, user_id, MANAGER_ROLES,
                branch_id=assignment["branch_id"],
                require_global=assignment["branch_id"] is None,
            )
            require_role_hierarchy(
                cur, context_id=request.context_id, actor_user_id=user_id,
                actor_roles=actor_roles,
                target_user_id=int(assignment["user_id"]),
                target_role=str(assignment["role_key"]),
            )
            if assignment["role_key"] in {"owner", "founder", "director"}:
                cur.execute(
                    """SELECT COUNT(*) remaining
                       FROM center_role_assignments
                       WHERE context_id=%s AND status='active'
                         AND role_key IN ('owner','founder','director')
                         AND branch_id IS NULL AND starts_at<=NOW()
                         AND (ends_at IS NULL OR ends_at>NOW())
                         AND id<>%s""",
                    (request.context_id, assignment_id),
                )
                if int(cur.fetchone()["remaining"]) < 1:
                    raise HTTPException(
                        status_code=409,
                        detail="Markazda kamida bitta yuqori rahbar qolishi kerak",
                    )
            if assignment["role_key"] == "teacher":
                cur.execute(
                    """SELECT COUNT(*) active_courses FROM center_courses
                       WHERE context_id=%s AND teacher_user_id=%s
                         AND status='active'""",
                    (request.context_id, assignment["user_id"]),
                )
                if int(cur.fetchone()["active_courses"]) > 0:
                    raise HTTPException(
                        status_code=409,
                        detail="Avval o'qituvchining faol kurslarini qayta biriktiring",
                    )
            cur.execute(
                """UPDATE center_role_assignments SET status=%s,
                     ends_at=CASE WHEN %s='ended' THEN NOW() ELSE ends_at END,
                     updated_at=NOW()
                   WHERE id=%s RETURNING *""",
                (request.status, request.status, assignment_id),
            )
            updated = cur.fetchone()
            sync_generic_staff_membership(
                cur, context_id=request.context_id,
                user_id=int(assignment["user_id"]),
                generic_role=generic_member_role(str(assignment["role_key"])),
            )
            audit(
                cur, request.context_id, user_id, "staff.status.change",
                "role_assignment", assignment_id,
                {
                    "status": request.status,
                    "target_user_id": assignment["user_id"],
                    "role_key": assignment["role_key"],
                },
            )
        return {"item": updated}

    @router.get("/staff/{teacher_user_id}/availability")
    def get_availability(
        teacher_user_id: int,
        context_id: int = Query(ge=1),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            require_teacher_scope(cur, context_id, user_id, teacher_user_id)
            cur.execute(
                """SELECT id,weekday,starts_at,ends_at,availability,
                          effective_from,effective_to,note
                   FROM center_teacher_availability
                   WHERE context_id=%s AND teacher_user_id=%s
                   ORDER BY weekday,starts_at,id""",
                (context_id, teacher_user_id),
            )
            rows = cur.fetchall()
        return {"items": rows, "next_cursor": None}

    @router.put("/staff/{teacher_user_id}/availability")
    def put_availability(
        teacher_user_id: int, request: AvailabilityPut,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        parsed: list[tuple[Any, ...]] = []
        for row in request.rows:
            start, end = parse_time(row.starts_at, "starts_at"), parse_time(row.ends_at, "ends_at")
            if start >= end:
                raise HTTPException(status_code=422, detail="Mavjudlik tugash vaqti noto'g'ri")
            if row.effective_from and row.effective_to and row.effective_to < row.effective_from:
                raise HTTPException(status_code=422, detail="Mavjudlik sanalari noto'g'ri")
            parsed.append(
                (
                    request.context_id, teacher_user_id, row.weekday, start, end,
                    row.availability, row.effective_from, row.effective_to, row.note,
                )
            )
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            if user_id == teacher_user_id:
                require_roles(
                    cur, request.context_id, user_id, {"teacher", "methodist"}
                )
            else:
                require_roles(
                    cur, request.context_id, user_id, MANAGER_ROLES,
                    require_global=True,
                )
            cur.execute(
                """SELECT 1 FROM center_role_assignments
                   WHERE context_id=%s AND user_id=%s AND role_key='teacher'
                     AND status='active' AND starts_at<=NOW()
                     AND (ends_at IS NULL OR ends_at>NOW())""",
                (request.context_id, teacher_user_id),
            )
            if not cur.fetchone():
                raise HTTPException(status_code=404, detail="Faol o'qituvchi topilmadi")
            cur.execute(
                """DELETE FROM center_teacher_availability
                   WHERE context_id=%s AND teacher_user_id=%s""",
                (request.context_id, teacher_user_id),
            )
            for values in parsed:
                cur.execute(
                    """INSERT INTO center_teacher_availability(
                         context_id,teacher_user_id,weekday,starts_at,ends_at,
                         availability,effective_from,effective_to,note
                       ) VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                    values,
                )
            audit(
                cur, request.context_id, user_id, "teacher.availability.replace",
                "user", teacher_user_id, {"rows": len(parsed)},
            )
        return {"teacher_user_id": teacher_user_id, "saved": len(parsed)}

    def normalize_course_target(
        request: CourseCreate,
    ) -> tuple[str, str | None, str | None, Decimal | None, dict[str, Decimal]]:
        framework = "custom" if request.target_framework == "general" else request.target_framework
        level_from, level_to = request.level_from, request.level_to
        target_score = request.target_score
        components = dict(request.target_components)
        if framework == "cefr":
            level_from = (request.cefr_level or level_from or "").upper()
            level_to = (level_to or level_from).upper()
            if level_from not in CEFR_LEVELS or level_to not in CEFR_LEVELS:
                raise HTTPException(status_code=422, detail="CEFR darajasi A1–C2 bo'lsin")
            if CEFR_LEVELS.index(level_from) > CEFR_LEVELS.index(level_to):
                raise HTTPException(status_code=422, detail="CEFR boshlang'ich darajasi yuqori")
        elif framework == "ielts":
            targets = dict(request.ielts_targets or {})
            if not components:
                components = {
                    key: Decimal(str(targets.get(key))) for key in IELTS_COMPONENTS
                    if targets.get(key) is not None
                }
            target_score = target_score or (
                Decimal(str(targets["overall"])) if targets.get("overall") is not None else None
            )
            all_scores = [target_score, *(components.get(key) for key in IELTS_COMPONENTS)]
            if any(score is None for score in all_scores):
                raise HTTPException(
                    status_code=422,
                    detail="IELTS overall va 4 komponent maqsadini kiriting.",
                )
            if any(
                score < 0 or score > 9 or score * 2 != (score * 2).to_integral_value()
                for score in all_scores if score is not None
            ):
                raise HTTPException(
                    status_code=422, detail="IELTS maqsadi 0–9 va 0.5 qadamda bo'lsin"
                )
        else:
            level_from = request.level_label or level_from
            level_to = level_to or level_from
            target_score = None
            components = {}
        return framework, level_from, level_to, target_score, components

    def decorate_course_schedule(
        rows: list[Any],
    ) -> list[dict[str, Any]]:
        decorated = [dict(row) for row in rows]
        for row in decorated:
            weekdays = [
                str(day) for day in (row.get("weekdays") or [])
                if day is not None
            ]
            starts_at = row.get("starts_at")
            time_label = (
                starts_at.strftime("%H:%M")
                if hasattr(starts_at, "strftime") else None
            )
            parts = []
            if weekdays:
                parts.append(", ".join(weekdays))
            if time_label:
                parts.append(time_label)
            row["schedule_label"] = " · ".join(parts) or None
        return decorated

    @router.get("/courses/catalog")
    def course_catalog(
        context_id: int = Query(ge=1),
        branch_id: int | None = Query(default=None, ge=1),
        subject_id: int | None = Query(default=None, ge=1),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=50, ge=1, le=100),
        _: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            require_operational(cur, context_id)
            cur.execute(
                """SELECT c.id,c.branch_id,b.name branch_name,c.subject_id,
                          s.name subject_name,c.name,c.course_type,
                          c.delivery_mode,c.target_framework,c.level_from,
                          c.level_to,c.capacity,c.duration_minutes,
                          c.metadata->>'ielts_test_type' ielts_test_type,
                          u.full_name teacher_name,
                          COALESCE(enrollment.enrolled_count,0) enrolled_count,
                          GREATEST(
                            c.capacity-COALESCE(enrollment.enrolled_count,0),0
                          ) available_seats,
                          ARRAY(
                            SELECT sl.weekday
                            FROM center_schedule_slots sl
                            WHERE sl.context_id=c.context_id
                              AND sl.course_id=c.id
                              AND sl.schedule_kind='weekly'
                              AND sl.status<>'cancelled'
                            ORDER BY sl.weekday,sl.starts_at,sl.id
                          ) weekdays,
                          (SELECT MIN(sl.starts_at)
                           FROM center_schedule_slots sl
                           WHERE sl.context_id=c.context_id
                             AND sl.course_id=c.id
                             AND sl.status<>'cancelled') starts_at
                   FROM center_courses c
                   JOIN center_subjects s ON s.id=c.subject_id
                   LEFT JOIN center_branches b ON b.id=c.branch_id
                   LEFT JOIN users u ON u.user_id=c.teacher_user_id
                   LEFT JOIN LATERAL (
                     SELECT COUNT(*) enrolled_count
                     FROM center_enrollments e
                     WHERE e.context_id=c.context_id AND e.course_id=c.id
                       AND e.status='active'
                   ) enrollment ON TRUE
                   WHERE c.context_id=%s AND c.status='active' AND c.id>%s
                     AND (%s IS NULL OR c.branch_id=%s)
                     AND (%s IS NULL OR c.subject_id=%s)
                   ORDER BY c.id LIMIT %s""",
                (
                    context_id, after_id or 0, branch_id, branch_id,
                    subject_id, subject_id, limit + 1,
                ),
            )
            rows = decorate_course_schedule(cur.fetchall())
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
        }

    @router.get("/courses")
    def list_courses(
        context_id: int = Query(ge=1),
        status: str | None = Query(default=None, max_length=30),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=50, ge=1, le=200),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            roles = require_roles(cur, context_id, user_id, VIEW_ROLES)
            cur.execute(
                """SELECT c.id,c.group_id,c.branch_id,b.name branch_name,
                          c.subject_id,s.name subject_name,c.teacher_user_id,
                          u.full_name teacher_name,c.name,c.course_type,
                          c.delivery_mode,c.target_framework,c.level_from,c.level_to,
                          c.target_score,c.target_components,c.capacity,
                          c.duration_minutes,c.status,c.metadata,
                          c.metadata->>'ielts_test_type' ielts_test_type,
                          ARRAY(
                            SELECT sl.weekday
                            FROM center_schedule_slots sl
                            WHERE sl.context_id=c.context_id
                              AND sl.course_id=c.id
                              AND sl.schedule_kind='weekly'
                              AND sl.status<>'cancelled'
                            ORDER BY sl.weekday,sl.starts_at,sl.id
                          ) weekdays,
                          (SELECT MIN(sl.starts_at)
                           FROM center_schedule_slots sl
                           WHERE sl.context_id=c.context_id
                             AND sl.course_id=c.id
                             AND sl.status<>'cancelled') starts_at,
                          (SELECT COUNT(*) FROM center_enrollments e
                           WHERE e.context_id=c.context_id AND e.course_id=c.id
                             AND e.status='active') enrolled_count
                   FROM center_courses c
                   JOIN center_subjects s ON s.id=c.subject_id
                   LEFT JOIN center_branches b ON b.id=c.branch_id
                   LEFT JOIN users u ON u.user_id=c.teacher_user_id
                   WHERE c.context_id=%s AND c.id>%s
                     AND (%s IS NULL OR c.status=%s)
                     AND (
                       %s OR (
                         c.teacher_user_id=%s AND EXISTS(
                           SELECT 1 FROM center_role_assignments teacher_role
                           WHERE teacher_role.context_id=c.context_id
                             AND teacher_role.user_id=%s
                             AND teacher_role.role_key='teacher'
                             AND teacher_role.status='active'
                             AND teacher_role.starts_at<=NOW()
                             AND (
                               teacher_role.ends_at IS NULL
                               OR teacher_role.ends_at>NOW()
                             )
                             AND (
                               teacher_role.branch_id IS NULL
                               OR teacher_role.branch_id IS NOT DISTINCT FROM c.branch_id
                             )
                         )
                       )
                       OR EXISTS(
                         SELECT 1 FROM center_enrollments e
                         WHERE e.context_id=c.context_id AND e.course_id=c.id
                           AND e.student_user_id=%s
                           AND e.status IN ('active','completed')
                           AND (
                             e.starts_on IS NULL
                             OR e.starts_on<=
                               (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                           )
                       )
                       OR EXISTS(
                         SELECT 1 FROM center_parent_links p
                         JOIN center_enrollments e
                           ON e.context_id=p.context_id
                          AND e.student_user_id=p.student_user_id
                         WHERE p.context_id=c.context_id
                           AND p.parent_user_id=%s AND p.status='active'
                           AND e.course_id=c.id
                           AND e.status IN ('active','completed')
                           AND (
                             e.starts_on IS NULL
                             OR e.starts_on<=
                               (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                           )
                       )
                       OR EXISTS(
                         SELECT 1 FROM center_role_assignments r
                         WHERE r.context_id=c.context_id AND r.user_id=%s
                           AND r.status='active' AND r.starts_at<=NOW()
                           AND (r.ends_at IS NULL OR r.ends_at>NOW())
                           AND r.role_key=ANY(%s)
                           AND (
                             r.branch_id IS NULL
                             OR (c.branch_id IS NOT NULL AND r.branch_id=c.branch_id)
                           )
                       )
                     )
                   ORDER BY c.id LIMIT %s""",
                (
                    context_id, after_id or 0, status, status,
                    "system_admin" in roles, user_id, user_id,
                    user_id, user_id, user_id,
                    list(MANAGER_ROLES | {"methodist", "receptionist", "accountant"}),
                    limit + 1,
                ),
            )
            rows = decorate_course_schedule(cur.fetchall())
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
        }

    @router.post("/courses")
    def create_course(
        request: CourseCreate,
        idempotency_key: str | None = Header(
            default=None, alias="Idempotency-Key"
        ),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        idempotency_key = validate_request_key(idempotency_key)
        require_encoded_size(
            {
                "metadata": request.metadata,
                "target_components": request.target_components,
                "ielts_targets": request.ielts_targets,
            },
            maximum=100_000,
            label="Kurs qo'shimcha ma'lumotlari",
        )
        framework, level_from, level_to, target_score, components = (
            normalize_course_target(request)
        )
        if request.delivery_mode != "online_live" and request.branch_id is None:
            raise HTTPException(status_code=422, detail="Oflayn/gibrid kurs uchun filial kerak")
        if request.course_type == "individual" and request.capacity != 1:
            raise HTTPException(status_code=422, detail="Individual kurs sig'imi 1 bo'ladi")
        if request.start_date and request.end_date and request.end_date < request.start_date:
            raise HTTPException(status_code=422, detail="Kurs sanalari noto'g'ri")
        if request.weekdays and any(day < 1 or day > 7 for day in request.weekdays):
            raise HTTPException(status_code=422, detail="Hafta kuni 1–7 bo'lsin")
        start_time = parse_time(request.starts_at, "starts_at") if request.starts_at else None
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            scoped_roles = require_roles(
                cur, request.context_id, user_id, ACADEMIC_ROLES,
                branch_id=request.branch_id,
                require_global=request.branch_id is None,
            )
            cur.execute(
                """SELECT name FROM center_subjects
                   WHERE id=%s AND context_id=%s AND active=TRUE""",
                (request.subject_id, request.context_id),
            )
            subject = cur.fetchone()
            if not subject:
                raise HTTPException(status_code=404, detail="Fan topilmadi")
            teacher_id = request.teacher_user_id
            may_assign_other_teacher = bool(
                scoped_roles & (MANAGER_ROLES | {"methodist"})
            )
            if not may_assign_other_teacher:
                if teacher_id not in (None, user_id):
                    raise HTTPException(status_code=403, detail="Faqat o'zingizga kurs ochasiz")
                teacher_id = user_id
            if teacher_id is not None:
                cur.execute(
                    """SELECT 1 FROM center_role_assignments
                       WHERE context_id=%s AND user_id=%s AND role_key='teacher'
                         AND status='active' AND starts_at<=NOW()
                         AND (ends_at IS NULL OR ends_at>NOW())
                         AND (
                           branch_id IS NULL OR branch_id IS NOT DISTINCT FROM %s
                         )""",
                    (request.context_id, teacher_id, request.branch_id),
                )
                if not cur.fetchone():
                    raise HTTPException(status_code=409, detail="O'qituvchi bu filialga biriktirilmagan")
                cur.execute(
                    """SELECT 1 FROM center_staff_subjects
                       WHERE context_id=%s AND teacher_user_id=%s
                         AND subject_id=%s AND active=TRUE""",
                    (request.context_id, teacher_id, request.subject_id),
                )
                if not cur.fetchone():
                    raise HTTPException(
                        status_code=409, detail="O'qituvchining fan malakasi biriktirilmagan"
                    )
            scheduled_end: time | None = None
            if teacher_id and start_time and request.weekdays:
                total_minutes = (
                    start_time.hour * 60
                    + start_time.minute
                    + request.duration_minutes
                )
                if total_minutes >= 24 * 60:
                    raise HTTPException(
                        status_code=422, detail="Dars yarim tundan o'tmasin"
                    )
                scheduled_end = time(total_minutes // 60, total_minutes % 60)
                for weekday in sorted(set(request.weekdays)):
                    require_branch_workday(
                        cur, context_id=request.context_id,
                        branch_id=request.branch_id, weekday=weekday,
                        starts_at=start_time, ends_at=scheduled_end,
                    )
                    ensure_teacher_available(
                        cur, context_id=request.context_id,
                        teacher_user_id=teacher_id,
                        schedule_kind="weekly", weekday=weekday,
                        starts_at=start_time, ends_at=scheduled_end,
                        effective_from=request.start_date,
                        effective_to=request.end_date,
                    )
            if idempotency_key:
                existing_target = claim_request(
                    cur, context_id=request.context_id, actor=user_id,
                    key=idempotency_key, action="course.create",
                )
                if existing_target:
                    cur.execute(
                        """SELECT * FROM center_courses
                           WHERE id=%s AND context_id=%s""",
                        (existing_target, request.context_id),
                    )
                    return {
                        "item": cur.fetchone(),
                        "idempotent_replay": True,
                    }
            cur.execute(
                """INSERT INTO course_groups(
                     context_id,group_type,delivery_mode,name,subject,
                     teacher_user_id,metadata
                   ) VALUES(%s,'center_course',%s,%s,%s,%s,%s::jsonb)
                   RETURNING id""",
                (
                    request.context_id, request.delivery_mode, request.name.strip(),
                    subject["name"], teacher_id,
                    json.dumps({"source": "center_v2"}),
                ),
            )
            group_id = int(cur.fetchone()["id"])
            cur.execute(
                """INSERT INTO center_courses(
                     context_id,group_id,branch_id,subject_id,teacher_user_id,name,
                     course_type,delivery_mode,target_framework,level_from,level_to,
                     target_score,target_components,capacity,duration_minutes,status,
                     metadata
                   ) VALUES(
                     %s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s::jsonb,%s,%s,%s,
                     %s::jsonb
                   ) RETURNING *""",
                (
                    request.context_id, group_id, request.branch_id,
                    request.subject_id, teacher_id, request.name.strip(),
                    request.course_type, request.delivery_mode, framework,
                    level_from, level_to, target_score,
                    json.dumps(components, default=str),
                    request.capacity, request.duration_minutes, "draft",
                    json.dumps(
                        {
                            **request.metadata,
                            "level_label": request.level_label,
                            "monthly_price_hint": str(request.monthly_price or 0),
                            "sessions_per_week": request.sessions_per_week,
                            "start_date": request.start_date,
                            "end_date": request.end_date,
                            "ielts_test_type": (
                                request.ielts_test_type
                                if framework == "ielts" else None
                            ),
                        },
                        ensure_ascii=False,
                        default=str,
                    ),
                ),
            )
            course = cur.fetchone()
            if teacher_id and start_time and request.weekdays:
                assert scheduled_end is not None
                for weekday in sorted(set(request.weekdays)):
                    ensure_teacher_weekly_capacity(
                        cur, context_id=request.context_id,
                        teacher_user_id=int(teacher_id),
                        branch_id=request.branch_id,
                        starts_at=start_time, ends_at=scheduled_end,
                    )
                    cur.execute(
                        """INSERT INTO center_schedule_slots(
                             context_id,course_id,group_id,branch_id,
                             teacher_user_id,schedule_kind,weekday,effective_from,
                             effective_to,starts_at,ends_at,status,created_by_user_id,
                             metadata
                           ) VALUES(
                             %s,%s,%s,%s,%s,'weekly',%s,%s,%s,%s,%s,'draft',%s,
                             '{"source":"course_create"}'
                           )""",
                        (
                            request.context_id, course["id"], group_id,
                            request.branch_id, teacher_id, weekday,
                            request.start_date, request.end_date,
                            start_time, scheduled_end,
                            user_id,
                        ),
                    )
            if idempotency_key:
                finish_request(
                    cur, request.context_id, idempotency_key, course["id"]
                )
            audit(
                cur, request.context_id, user_id, "course.create", "course",
                course["id"],
            )
        return {"item": course}

    @router.post("/courses/{course_id}/activate")
    def activate_course(
        course_id: int,
        request: CourseActivation,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(
            request.confirmation,
            "Kursni faollashtirish uchun inson tasdig'i kerak.",
        )
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            course = course_resource(
                cur, request.context_id, course_id, lock=True
            )
            scoped_roles = require_permission(
                cur, request.context_id, user_id, "courses.manage",
                branch_id=course["branch_id"],
            )
            teacher_id = request.teacher_user_id or course["teacher_user_id"]
            if teacher_id is None:
                cur.execute(
                    """SELECT ARRAY_AGG(DISTINCT teacher_user_id)
                              FILTER(WHERE teacher_user_id IS NOT NULL)
                              scheduled_teachers
                       FROM center_schedule_slots
                       WHERE context_id=%s AND course_id=%s
                         AND status<>'cancelled'""",
                    (request.context_id, course_id),
                )
                scheduled_teachers = list(
                    cur.fetchone()["scheduled_teachers"] or []
                )
                if len(scheduled_teachers) == 1:
                    teacher_id = int(scheduled_teachers[0])
                elif not scheduled_teachers:
                    raise HTTPException(
                        status_code=409,
                        detail="Avval kurs jadvaliga o'qituvchi biriktiring",
                    )
                else:
                    raise HTTPException(
                        status_code=409,
                        detail=(
                            "Jadvalda bir nechta o'qituvchi bor; "
                            "kurs o'qituvchisini aniq tanlang"
                        ),
                    )
            if not (
                "system_admin" in scoped_roles
                or scoped_roles & (MANAGER_ROLES | {"methodist"})
            ):
                if (
                    int(teacher_id) != user_id
                    or int(course["teacher_user_id"] or 0) != user_id
                ):
                    raise HTTPException(
                        status_code=403,
                        detail="Faqat avval sizga biriktirilgan kursingizni faollashtirasiz",
                    )
            cur.execute(
                """SELECT 1 FROM center_role_assignments
                   WHERE context_id=%s AND user_id=%s AND role_key='teacher'
                     AND status='active' AND starts_at<=NOW()
                     AND (ends_at IS NULL OR ends_at>NOW())
                     AND (
                       branch_id IS NULL OR branch_id IS NOT DISTINCT FROM %s
                     )""",
                (request.context_id, teacher_id, course["branch_id"]),
            )
            if not cur.fetchone():
                raise HTTPException(
                    status_code=409,
                    detail="O'qituvchi kurs filialiga biriktirilmagan",
                )
            cur.execute(
                """SELECT 1 FROM center_staff_subjects
                   WHERE context_id=%s AND teacher_user_id=%s
                     AND subject_id=%s AND active=TRUE""",
                (request.context_id, teacher_id, course["subject_id"]),
            )
            if not cur.fetchone():
                raise HTTPException(
                    status_code=409,
                    detail="O'qituvchining fan malakasi biriktirilmagan",
                )
            cur.execute(
                """SELECT COUNT(*) slot_count,
                          COUNT(*) FILTER(WHERE teacher_user_id=%s) teacher_slots
                   FROM center_schedule_slots
                   WHERE context_id=%s AND course_id=%s
                     AND status<>'cancelled'""",
                (teacher_id, request.context_id, course_id),
            )
            schedule = cur.fetchone()
            if int(schedule["slot_count"]) < 1:
                raise HTTPException(
                    status_code=409,
                    detail="Kursni faollashtirishdan oldin dars vaqtini kiriting",
                )
            if int(schedule["teacher_slots"]) != int(schedule["slot_count"]):
                raise HTTPException(
                    status_code=409,
                    detail="Jadvaldagi o'qituvchi kurs o'qituvchisiga mos emas",
                )
            cur.execute(
                """UPDATE center_courses SET teacher_user_id=%s,status='active',
                     updated_at=NOW()
                   WHERE id=%s RETURNING *""",
                (teacher_id, course_id),
            )
            updated = cur.fetchone()
            cur.execute(
                """UPDATE course_groups SET teacher_user_id=%s,active=TRUE,
                     updated_at=NOW() WHERE id=%s""",
                (teacher_id, course["group_id"]),
            )
            cur.execute(
                """UPDATE center_schedule_slots SET status='published',
                     updated_at=NOW()
                   WHERE context_id=%s AND course_id=%s AND status='draft'""",
                (request.context_id, course_id),
            )
            audit(
                cur, request.context_id, user_id, "course.activate",
                "course", course_id, {"teacher_user_id": teacher_id},
            )
        return {"item": updated}

    @router.get("/enrollments")
    def list_enrollments(
        context_id: int = Query(ge=1),
        course_id: int | None = Query(default=None, ge=1),
        status: str | None = Query(default=None, max_length=30),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=100, ge=1, le=300),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            roles = require_roles(cur, context_id, user_id, VIEW_ROLES)
            if course_id is not None:
                filtered_course = course_resource(cur, context_id, course_id)
                can_staff_read = (
                    has_permission(
                        cur, context_id, user_id, "enrollments.manage",
                        branch_id=filtered_course["branch_id"],
                    )
                    or has_permission(
                        cur, context_id, user_id, "billing.manage",
                        branch_id=filtered_course["branch_id"],
                    )
                )
                if not can_staff_read:
                    course_access(
                        cur, context_id, course_id, user_id, write=False
                    )
            cur.execute(
                """SELECT e.id,e.course_id,c.name course_name,e.student_user_id,
                          u.full_name student_name,e.status,e.waitlist_position,
                          e.starts_on,e.starts_on start_date,
                          e.enrolled_at,e.ended_at,e.note
                   FROM center_enrollments e
                   JOIN center_courses c ON c.id=e.course_id
                   JOIN users u ON u.user_id=e.student_user_id
                   WHERE e.context_id=%s AND e.id>%s
                     AND (%s IS NULL OR e.course_id=%s)
                     AND (%s IS NULL OR e.status=%s)
                     AND (
                       %s OR e.student_user_id=%s
                       OR (
                         c.teacher_user_id=%s AND EXISTS(
                           SELECT 1 FROM center_role_assignments teacher_role
                           WHERE teacher_role.context_id=e.context_id
                             AND teacher_role.user_id=%s
                             AND teacher_role.role_key='teacher'
                             AND teacher_role.status='active'
                             AND teacher_role.starts_at<=NOW()
                             AND (
                               teacher_role.ends_at IS NULL
                               OR teacher_role.ends_at>NOW()
                             )
                             AND (
                               teacher_role.branch_id IS NULL
                               OR teacher_role.branch_id IS NOT DISTINCT FROM c.branch_id
                             )
                         )
                       )
                       OR EXISTS(
                         SELECT 1 FROM center_parent_links p
                         WHERE p.context_id=e.context_id
                           AND p.parent_user_id=%s
                           AND p.student_user_id=e.student_user_id
                           AND p.status='active'
                       )
                       OR EXISTS(
                         SELECT 1 FROM center_role_assignments r
                         WHERE r.context_id=e.context_id AND r.user_id=%s
                           AND r.status='active' AND r.starts_at<=NOW()
                           AND (r.ends_at IS NULL OR r.ends_at>NOW())
                           AND r.role_key=ANY(%s)
                           AND (
                             r.branch_id IS NULL
                             OR (c.branch_id IS NOT NULL AND r.branch_id=c.branch_id)
                           )
                       )
                     )
                   ORDER BY e.id LIMIT %s""",
                (
                    context_id, after_id or 0, course_id, course_id, status, status,
                    "system_admin" in roles, user_id, user_id, user_id,
                    user_id, user_id,
                    list(
                        ENROLLMENT_ROLES
                        | {"methodist"}
                        | FINANCE_WRITE_ROLES
                    ),
                    limit + 1,
                ),
            )
            rows = cur.fetchall()
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
        }

    @router.post("/enrollments")
    def create_enrollment(
        request: EnrollmentCreate,
        idempotency_key: str | None = Header(
            default=None, alias="Idempotency-Key"
        ),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        idempotency_key = validate_request_key(idempotency_key)
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            lock_course_queue(cur, request.context_id, request.course_id)
            course = course_resource(
                cur, request.context_id, request.course_id
            )
            if course["status"] != "active":
                raise HTTPException(
                    status_code=409,
                    detail="Faqat faol kursga ariza berish mumkin",
                )
            if idempotency_key:
                existing_target = claim_request(
                    cur, context_id=request.context_id, actor=user_id,
                    key=idempotency_key, action="enrollment.create",
                )
                if existing_target:
                    cur.execute(
                        """SELECT * FROM center_enrollments
                           WHERE id=%s AND context_id=%s""",
                        (existing_target, request.context_id),
                    )
                    return {
                        "item": cur.fetchone(),
                        "idempotent_replay": True,
                    }
            student_id = request.student_user_id or user_id
            manager = has_permission(
                cur, request.context_id, user_id, "enrollments.manage",
                branch_id=course["branch_id"],
            )
            if not manager and student_id != user_id:
                raise HTTPException(status_code=403, detail="Faqat o'zingiz ariza berasiz")
            status_aliases = {
                "active": "active", "pending": "pending", "waitlist": "waitlisted",
                "waitlisted": "waitlisted",
            }
            requested = status_aliases.get(
                str(request.entry_status or request.requested_status), "pending"
            )
            if requested == "active":
                if not manager:
                    raise HTTPException(status_code=403, detail="Faollashtirish uchun qabulxona kerak")
                require_human(
                    request.confirmation, "O'quvchini faol qabul qilish uchun tasdiq kerak."
                )
            elif requested not in {"pending", "waitlisted"}:
                requested = "pending"
            if not manager and requested != "pending":
                raise HTTPException(
                    status_code=403,
                    detail="Mustaqil ariza avval qabulxona tekshiruviga tushadi",
                )
            cur.execute(
                """SELECT status,waitlist_position FROM center_enrollments
                   WHERE context_id=%s AND course_id=%s
                     AND student_user_id=%s FOR UPDATE""",
                (request.context_id, request.course_id, student_id),
            )
            existing_enrollment = cur.fetchone()
            if (
                existing_enrollment
                and existing_enrollment["status"] in {"active", "completed"}
                and requested not in {"active"}
            ):
                raise HTTPException(
                    status_code=409,
                    detail="O'quvchi bu kursda allaqachon faol",
                )
            cur.execute(
                """SELECT COUNT(*) occupied FROM center_enrollments
                   WHERE context_id=%s AND course_id=%s AND status='active'""",
                (request.context_id, request.course_id),
            )
            occupied = int(cur.fetchone()["occupied"])
            if requested == "active" and occupied >= int(course["capacity"]):
                requested = "waitlisted"
            wait_position = None
            if requested == "waitlisted":
                if (
                    existing_enrollment
                    and existing_enrollment["status"] == "waitlisted"
                    and existing_enrollment["waitlist_position"] is not None
                ):
                    wait_position = int(
                        existing_enrollment["waitlist_position"]
                    )
                else:
                    cur.execute(
                        """SELECT COALESCE(MAX(waitlist_position),0)+1 position
                           FROM center_enrollments
                           WHERE context_id=%s AND course_id=%s
                             AND status='waitlisted'""",
                        (request.context_id, request.course_id),
                    )
                    wait_position = int(cur.fetchone()["position"])
            cur.execute(
                """INSERT INTO center_enrollments(
                     context_id,course_id,student_user_id,status,waitlist_position,
                     starts_on,enrolled_at,approved_by_user_id,note
                   ) VALUES(
                     %s,%s,%s,%s,%s,%s,
                     CASE WHEN %s='active' THEN NOW() END,%s,%s
                   )
                   ON CONFLICT(context_id,course_id,student_user_id)
                   DO UPDATE SET status=EXCLUDED.status,
                     waitlist_position=EXCLUDED.waitlist_position,
                     enrolled_at=COALESCE(center_enrollments.enrolled_at,EXCLUDED.enrolled_at),
                     approved_by_user_id=EXCLUDED.approved_by_user_id,
                     starts_on=COALESCE(
                       EXCLUDED.starts_on,center_enrollments.starts_on
                     ),
                     note=EXCLUDED.note,updated_at=NOW()
                   RETURNING *""",
                (
                    request.context_id, request.course_id, student_id, requested,
                    wait_position, request.start_date, requested,
                    user_id if manager else None,
                    request.note or request.notes,
                ),
            )
            enrollment = cur.fetchone()
            if idempotency_key:
                finish_request(
                    cur, request.context_id, idempotency_key, enrollment["id"]
                )
            sync_student_access(
                cur, context_id=request.context_id,
                student_user_id=student_id, branch_id=course["branch_id"],
                approved_by=user_id if manager else None,
            )
            audit(
                cur, request.context_id, user_id, "enrollment.create",
                "enrollment", enrollment["id"], {"status": requested},
            )
        return {"item": enrollment}

    @router.post("/enrollments/{enrollment_id}/decision")
    def decide_enrollment(
        enrollment_id: int, request: EnrollmentDecision,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(request.confirmation, "Qabul qarori uchun inson tasdig'i kerak.")
        status = request.status or {
            "approve": "active", "waitlist": "waitlisted", "reject": "rejected",
            "pause": "paused", "withdraw": "withdrawn",
        }.get(request.decision or "")
        if status is None:
            raise HTTPException(status_code=422, detail="Qabul qarorini tanlang")
        with database() as (_, cur):
            ensure_schema(cur)
            cur.execute(
                """SELECT e.*,c.branch_id,c.capacity FROM center_enrollments e
                   JOIN center_courses c ON c.id=e.course_id
                   WHERE e.id=%s""",
                (enrollment_id,),
            )
            enrollment = cur.fetchone()
            if not enrollment:
                raise HTTPException(status_code=404, detail="Ariza topilmadi")
            require_operational(cur, enrollment["context_id"])
            require_roles(
                cur, enrollment["context_id"], user_id, ENROLLMENT_ROLES,
                branch_id=enrollment["branch_id"],
                require_global=enrollment["branch_id"] is None,
            )
            lock_course_queue(
                cur, enrollment["context_id"], enrollment["course_id"]
            )
            cur.execute(
                """SELECT e.*,c.branch_id,c.capacity
                   FROM center_enrollments e
                   JOIN center_courses c ON c.id=e.course_id
                   WHERE e.id=%s AND e.context_id=%s FOR UPDATE""",
                (enrollment_id, enrollment["context_id"]),
            )
            enrollment = cur.fetchone()
            if not enrollment:
                raise HTTPException(status_code=404, detail="Ariza topilmadi")
            if status == "active":
                cur.execute(
                    """SELECT COUNT(*) occupied FROM center_enrollments
                       WHERE context_id=%s AND course_id=%s
                         AND status='active' AND id<>%s""",
                    (
                        enrollment["context_id"], enrollment["course_id"],
                        enrollment_id,
                    ),
                )
                if int(cur.fetchone()["occupied"]) >= int(enrollment["capacity"]):
                    status = "waitlisted"
            wait_position = None
            if status == "waitlisted":
                cur.execute(
                    """SELECT COALESCE(MAX(waitlist_position),0)+1 position
                       FROM center_enrollments
                       WHERE context_id=%s AND course_id=%s
                         AND status='waitlisted' AND id<>%s""",
                    (
                        enrollment["context_id"], enrollment["course_id"],
                        enrollment_id,
                    ),
                )
                wait_position = int(cur.fetchone()["position"])
            cur.execute(
                """UPDATE center_enrollments SET status=%s,waitlist_position=%s,
                     enrolled_at=CASE WHEN %s='active'
                       THEN COALESCE(enrolled_at,NOW()) ELSE enrolled_at END,
                     ended_at=CASE WHEN %s IN ('withdrawn','rejected')
                       THEN NOW() ELSE NULL END,
                     approved_by_user_id=%s,updated_at=NOW()
                   WHERE id=%s RETURNING *""",
                (
                    status, wait_position, status, status,
                    user_id, enrollment_id,
                ),
            )
            updated = cur.fetchone()
            sync_student_access(
                cur, context_id=enrollment["context_id"],
                student_user_id=enrollment["student_user_id"],
                branch_id=enrollment["branch_id"], approved_by=user_id,
            )
            audit(
                cur, enrollment["context_id"], user_id, "enrollment.decide",
                "enrollment", enrollment_id, {"status": status},
            )
        return {"item": updated}

    @router.get("/parent-links")
    def list_parent_links(
        context_id: int = Query(ge=1),
        student_user_id: int | None = Query(default=None, ge=1),
        status: Literal[
            "pending", "active", "revoked", "rejected"
        ] | None = Query(default=None),
        limit: int = Query(default=100, ge=1, le=300),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            require_permission(
                cur, context_id, user_id, "enrollments.manage",
                require_global=True,
            )
            cur.execute(
                """SELECT p.context_id,p.parent_user_id,
                          parent.full_name parent_name,p.student_user_id,
                          student.full_name student_name,p.status,
                          p.approved_by_user_id,p.created_at,p.updated_at
                   FROM center_parent_links p
                   JOIN users parent ON parent.user_id=p.parent_user_id
                   JOIN users student ON student.user_id=p.student_user_id
                   WHERE p.context_id=%s
                     AND (%s IS NULL OR p.student_user_id=%s)
                     AND (%s IS NULL OR p.status=%s)
                   ORDER BY p.student_user_id,p.parent_user_id
                   LIMIT %s""",
                (
                    context_id, student_user_id, student_user_id,
                    status, status, limit,
                ),
            )
            rows = cur.fetchall()
        return {"items": rows, "next_cursor": None}

    @router.post("/parent-links")
    def create_parent_link(
        request: ParentLinkCreate, user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(request.confirmation, "Ota-ona bog'lanishi uchun tasdiq kerak.")
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            cur.execute(
                """SELECT 1 FROM center_enrollments e
                   JOIN center_courses c ON c.id=e.course_id
                   WHERE e.context_id=%s AND e.student_user_id=%s
                     AND e.status='active'
                     AND (%s IS NULL OR c.branch_id=%s)
                   LIMIT 1""",
                (
                    request.context_id, request.student_user_id,
                    request.branch_id, request.branch_id,
                ),
            )
            if not cur.fetchone():
                raise HTTPException(status_code=409, detail="Faol o'quvchi topilmadi")
            require_permission(
                cur, request.context_id, user_id, "enrollments.manage",
                require_global=True,
            )
            cur.execute(
                "SELECT 1 FROM users WHERE user_id=%s",
                (request.parent_user_id,),
            )
            if not cur.fetchone():
                raise HTTPException(
                    status_code=404, detail="Ota-ona foydalanuvchisi topilmadi"
                )
            cur.execute(
                """INSERT INTO center_parent_links(
                     context_id,parent_user_id,student_user_id,status,
                     approved_by_user_id
                   ) VALUES(%s,%s,%s,'active',%s)
                   ON CONFLICT(context_id,parent_user_id,student_user_id)
                   DO UPDATE SET status='active',approved_by_user_id=EXCLUDED.approved_by_user_id,
                     updated_at=NOW()
                   RETURNING *""",
                (
                    request.context_id, request.parent_user_id,
                    request.student_user_id, user_id,
                ),
            )
            link = cur.fetchone()
            upsert_role(
                cur, context_id=request.context_id, user_id=request.parent_user_id,
                role_key="parent", branch_id=None, status="active",
                approved_by=user_id,
            )
            audit(
                cur, request.context_id, user_id, "parent.link", "user",
                request.parent_user_id,
                {"student_user_id": request.student_user_id},
            )
        return {"item": link}

    @router.post(
        "/parent-links/{parent_user_id}/{student_user_id}/revoke"
    )
    def revoke_parent_link(
        parent_user_id: int,
        student_user_id: int,
        request: ParentLinkRevoke,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(
            request.confirmation,
            "Ota-ona bog'lanishini bekor qilish uchun inson tasdig'i kerak.",
        )
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            require_permission(
                cur, request.context_id, user_id, "enrollments.manage",
                require_global=True,
            )
            cur.execute(
                """UPDATE center_parent_links SET status='revoked',
                     updated_at=NOW()
                   WHERE context_id=%s AND parent_user_id=%s
                     AND student_user_id=%s AND status='active'
                   RETURNING *""",
                (request.context_id, parent_user_id, student_user_id),
            )
            link = cur.fetchone()
            if not link:
                raise HTTPException(
                    status_code=404, detail="Faol ota-ona bog'lanishi topilmadi"
                )
            cur.execute(
                """SELECT 1 FROM center_parent_links
                   WHERE context_id=%s AND parent_user_id=%s
                     AND status='active' LIMIT 1""",
                (request.context_id, parent_user_id),
            )
            if not cur.fetchone():
                cur.execute(
                    """UPDATE center_role_assignments SET status='suspended',
                         updated_at=NOW()
                       WHERE context_id=%s AND user_id=%s
                         AND role_key='parent' AND status='active'""",
                    (request.context_id, parent_user_id),
                )
                sync_generic_staff_membership(
                    cur, context_id=request.context_id,
                    user_id=parent_user_id,
                    generic_role="parent_observer",
                )
            audit(
                cur, request.context_id, user_id, "parent.link.revoke",
                "user", parent_user_id,
                {"student_user_id": student_user_id},
            )
        return {"item": link, "status": "revoked"}

    @router.get("/schedule")
    def list_schedule(
        context_id: int = Query(ge=1),
        from_date: date | None = Query(default=None),
        course_id: int | None = Query(default=None, ge=1),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=100, ge=1, le=300),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        start_date = from_date or business_today()
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            roles = require_roles(cur, context_id, user_id, VIEW_ROLES)
            if course_id:
                filtered_course = course_resource(cur, context_id, course_id)
                if not has_permission(
                    cur, context_id, user_id, "schedule.manage",
                    branch_id=filtered_course["branch_id"],
                ):
                    course_access(cur, context_id, course_id, user_id)
            cur.execute(
                """SELECT sl.id,sl.course_id,c.name course_name,sl.group_id,
                          sl.branch_id,b.name branch_name,sl.room_id,r.name room_name,
                          sl.teacher_user_id,u.full_name teacher_name,
                          sl.schedule_kind,sl.weekday,
                          occurrence.lesson_date,
                          sl.effective_from,sl.effective_to,sl.starts_at,sl.ends_at,
                          sl.status,sl.metadata->>'topic' topic
                   FROM center_schedule_slots sl
                   JOIN center_courses c ON c.id=sl.course_id
                   LEFT JOIN center_branches b ON b.id=sl.branch_id
                   LEFT JOIN center_rooms r ON r.id=sl.room_id
                   JOIN users u ON u.user_id=sl.teacher_user_id
                   CROSS JOIN LATERAL (
                     SELECT CASE
                       WHEN sl.schedule_kind='dated' THEN sl.lesson_date
                       ELSE
                         GREATEST(
                           %s::DATE,
                           COALESCE(sl.effective_from,%s::DATE)
                         )
                         + (
                           (
                             sl.weekday-EXTRACT(
                               ISODOW FROM GREATEST(
                                 %s::DATE,
                                 COALESCE(sl.effective_from,%s::DATE)
                               )
                             )::INT+7
                           )%%7
                         )
                     END AS lesson_date
                   ) occurrence
                   WHERE sl.context_id=%s AND sl.id>%s
                     AND sl.status<>'cancelled'
                     AND (%s IS NULL OR sl.course_id=%s)
                     AND occurrence.lesson_date>=%s
                     AND occurrence.lesson_date<=
                       COALESCE(sl.effective_to,DATE '9999-12-31')
                     AND (
                       sl.status='published' OR %s OR EXISTS(
                         SELECT 1 FROM center_role_assignments draft_role
                         WHERE draft_role.context_id=sl.context_id
                           AND draft_role.user_id=%s
                           AND draft_role.status='active'
                           AND draft_role.starts_at<=NOW()
                           AND (
                             draft_role.ends_at IS NULL
                             OR draft_role.ends_at>NOW()
                           )
                           AND draft_role.role_key=ANY(%s)
                           AND (
                             draft_role.branch_id IS NULL
                             OR draft_role.branch_id IS NOT DISTINCT FROM sl.branch_id
                           )
                       )
                     )
                     AND (
                       %s OR (
                         sl.teacher_user_id=%s AND EXISTS(
                           SELECT 1 FROM center_role_assignments teacher_role
                           WHERE teacher_role.context_id=sl.context_id
                             AND teacher_role.user_id=%s
                             AND teacher_role.role_key='teacher'
                             AND teacher_role.status='active'
                             AND teacher_role.starts_at<=NOW()
                             AND (
                               teacher_role.ends_at IS NULL
                               OR teacher_role.ends_at>NOW()
                             )
                             AND (
                               teacher_role.branch_id IS NULL
                               OR teacher_role.branch_id IS NOT DISTINCT FROM sl.branch_id
                             )
                         )
                       )
                       OR EXISTS(
                         SELECT 1 FROM center_enrollments e
                         WHERE e.context_id=sl.context_id AND e.course_id=sl.course_id
                           AND e.student_user_id=%s
                           AND e.status IN ('active','completed')
                           AND (
                             e.starts_on IS NULL
                             OR e.starts_on<=
                               (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                           )
                       )
                       OR EXISTS(
                         SELECT 1 FROM center_parent_links p
                         JOIN center_enrollments e
                           ON e.context_id=p.context_id
                          AND e.student_user_id=p.student_user_id
                         WHERE p.context_id=sl.context_id
                           AND p.parent_user_id=%s AND p.status='active'
                           AND e.course_id=sl.course_id
                           AND e.status IN ('active','completed')
                           AND (
                             e.starts_on IS NULL
                             OR e.starts_on<=
                               (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                           )
                       )
                       OR EXISTS(
                         SELECT 1 FROM center_role_assignments ra
                         WHERE ra.context_id=sl.context_id AND ra.user_id=%s
                           AND ra.status='active' AND ra.starts_at<=NOW()
                           AND (ra.ends_at IS NULL OR ra.ends_at>NOW())
                           AND ra.role_key=ANY(%s)
                           AND (
                             ra.branch_id IS NULL OR (
                               sl.branch_id IS NOT NULL AND ra.branch_id=sl.branch_id
                             )
                           )
                       )
                     )
                   ORDER BY sl.id LIMIT %s""",
                (
                    start_date, start_date, start_date, start_date,
                    context_id, after_id or 0, course_id, course_id, start_date,
                    "system_admin" in roles, user_id, list(SCHEDULE_ROLES),
                    "system_admin" in roles, user_id, user_id,
                    user_id, user_id, user_id,
                    list(SCHEDULE_ROLES),
                    limit + 1,
                ),
            )
            rows = cur.fetchall()
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
            "weekly_rows_show_next_occurrence": True,
        }

    @router.post("/schedule")
    def create_schedule(
        request: ScheduleCreate,
        idempotency_key: str | None = Header(
            default=None, alias="Idempotency-Key"
        ),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        idempotency_key = validate_request_key(idempotency_key)
        start = parse_time(request.starts_at, "starts_at")
        if request.ends_at:
            end = parse_time(request.ends_at, "ends_at")
        elif request.duration_minutes:
            total = start.hour * 60 + start.minute + request.duration_minutes
            if total >= 24 * 60:
                raise HTTPException(status_code=422, detail="Dars yarim tundan o'tmasin")
            end = time(total // 60, total % 60)
        else:
            raise HTTPException(status_code=422, detail="ends_at yoki duration_minutes kerak")
        if end <= start:
            raise HTTPException(status_code=422, detail="Dars tugash vaqti kechroq bo'lsin")
        if request.schedule_kind == "weekly":
            if request.weekday is None or request.lesson_date is not None:
                raise HTTPException(status_code=422, detail="Haftalik darsga weekday kerak")
        else:
            if request.lesson_date is None or request.weekday is not None:
                raise HTTPException(status_code=422, detail="Sanali darsga lesson_date kerak")
        if request.effective_from and request.effective_to and request.effective_to < request.effective_from:
            raise HTTPException(status_code=422, detail="Jadval amal sanalari noto'g'ri")
        if request.status == "published":
            require_human(request.confirmation, "Darsni e'lon qilish uchun tasdiq kerak.")
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            course = course_resource(cur, request.context_id, request.course_id)
            require_permission(
                cur, request.context_id, user_id, "schedule.manage",
                branch_id=course["branch_id"],
            )
            teacher_id = request.teacher_user_id or course["teacher_user_id"]
            if teacher_id is None:
                raise HTTPException(status_code=422, detail="O'qituvchini tanlang")
            roles, _ = active_roles(cur, request.context_id, user_id)
            if "teacher" in roles and not roles & (
                MANAGER_ROLES | {"methodist", "receptionist"}
            ):
                if (
                    teacher_id != user_id
                    or int(course["teacher_user_id"] or 0) != user_id
                ):
                    raise HTTPException(
                        status_code=403,
                        detail="Faqat o'zingizga biriktirilgan kursga dars qo'shasiz",
                    )
            cur.execute(
                """SELECT 1 FROM center_staff_subjects
                   WHERE context_id=%s AND teacher_user_id=%s
                     AND subject_id=%s AND active=TRUE""",
                (request.context_id, teacher_id, course["subject_id"]),
            )
            if not cur.fetchone():
                raise HTTPException(status_code=409, detail="O'qituvchi bu fanga biriktirilmagan")
            if idempotency_key:
                existing_target = claim_request(
                    cur, context_id=request.context_id, actor=user_id,
                    key=idempotency_key, action="schedule.create",
                )
                if existing_target:
                    cur.execute(
                        """SELECT * FROM center_schedule_slots
                           WHERE id=%s AND context_id=%s""",
                        (existing_target, request.context_id),
                    )
                    return {
                        "item": cur.fetchone(),
                        "idempotent_replay": True,
                    }
            if request.room_id is not None:
                cur.execute(
                    """SELECT capacity,room_type FROM center_rooms
                       WHERE id=%s AND context_id=%s
                         AND branch_id IS NOT DISTINCT FROM %s AND active=TRUE""",
                    (request.room_id, request.context_id, course["branch_id"]),
                )
                room = cur.fetchone()
                if not room:
                    raise HTTPException(status_code=409, detail="Xona kurs filialiga mos emas")
                if (
                    room["room_type"] != "online"
                    and (
                        room["capacity"] is None
                        or int(room["capacity"]) < int(course["capacity"])
                    )
                ):
                    raise HTTPException(
                        status_code=409,
                        detail=(
                            "Xona sig'imi kurs sig'imidan kichik; kattaroq xona "
                            "tanlang yoki kurs sig'imini kamaytiring"
                        ),
                    )
            relevant_day = request.weekday or request.lesson_date.isoweekday()
            require_branch_workday(
                cur, context_id=request.context_id,
                branch_id=course["branch_id"], weekday=relevant_day,
                starts_at=start, ends_at=end,
            )
            ensure_teacher_available(
                cur, context_id=request.context_id,
                teacher_user_id=int(teacher_id),
                schedule_kind=request.schedule_kind,
                weekday=relevant_day, starts_at=start, ends_at=end,
                lesson_date=request.lesson_date,
                effective_from=request.effective_from,
                effective_to=request.effective_to,
            )
            if request.schedule_kind == "weekly":
                ensure_teacher_weekly_capacity(
                    cur, context_id=request.context_id,
                    teacher_user_id=int(teacher_id),
                    branch_id=course["branch_id"],
                    starts_at=start, ends_at=end,
                )
            cur.execute(
                """INSERT INTO center_schedule_slots(
                     context_id,course_id,group_id,branch_id,room_id,
                     teacher_user_id,schedule_kind,weekday,lesson_date,
                     effective_from,effective_to,starts_at,ends_at,status,
                     created_by_user_id,metadata
                   ) VALUES(
                     %s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s::jsonb
                   ) RETURNING *""",
                (
                    request.context_id, request.course_id, course["group_id"],
                    course["branch_id"], request.room_id, teacher_id,
                    request.schedule_kind, request.weekday, request.lesson_date,
                    request.effective_from, request.effective_to, start, end,
                    request.status, user_id,
                    json.dumps({"topic": request.topic}, ensure_ascii=False),
                ),
            )
            slot = cur.fetchone()
            if idempotency_key:
                finish_request(
                    cur, request.context_id, idempotency_key, slot["id"]
                )
            audit(
                cur, request.context_id, user_id, "schedule.create",
                "schedule_slot", slot["id"], {"status": request.status},
            )
        return {"item": slot}

    @router.post("/schedule/{slot_id}/publish")
    def publish_schedule_slot(
        slot_id: int,
        request: HumanConfirm,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(
            request.confirmation,
            "Dars vaqtini e'lon qilish uchun inson tasdig'i kerak.",
        )
        with database() as (_, cur):
            ensure_schema(cur)
            cur.execute(
                """SELECT sl.*,c.branch_id course_branch_id,
                          c.teacher_user_id course_teacher_user_id
                   FROM center_schedule_slots sl
                   JOIN center_courses c ON c.id=sl.course_id
                   WHERE sl.id=%s FOR UPDATE""",
                (slot_id,),
            )
            slot = cur.fetchone()
            if not slot:
                raise HTTPException(
                    status_code=404, detail="Dars jadvali topilmadi"
                )
            require_operational(cur, slot["context_id"])
            roles = require_permission(
                cur, slot["context_id"], user_id, "schedule.manage",
                branch_id=slot["course_branch_id"],
            )
            if not (
                "system_admin" in roles
                or roles & (MANAGER_ROLES | {"methodist", "receptionist"})
            ):
                if (
                    int(slot["teacher_user_id"]) != user_id
                    or int(slot["course_teacher_user_id"] or 0) != user_id
                ):
                    raise HTTPException(
                        status_code=403,
                        detail="Faqat o'zingizga biriktirilgan darsni e'lon qilasiz",
                    )
            cur.execute(
                """UPDATE center_schedule_slots SET status='published',
                     updated_at=NOW()
                   WHERE id=%s RETURNING *""",
                (slot_id,),
            )
            updated = cur.fetchone()
            audit(
                cur, slot["context_id"], user_id, "schedule.publish",
                "schedule_slot", slot_id,
            )
        return {"item": updated}

    @router.get("/attendance")
    def list_attendance(
        context_id: int = Query(ge=1),
        course_id: int | None = Query(default=None, ge=1),
        date_value: date | None = Query(default=None, alias="date"),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=100, ge=1, le=300),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            roles = require_roles(cur, context_id, user_id, VIEW_ROLES)
            if course_id:
                course_access(cur, context_id, course_id, user_id)
            cur.execute(
                """SELECT a.id,a.course_id,c.name course_name,a.student_user_id,
                          u.full_name student_name,a.lesson_date attendance_date,
                          a.status,a.note,a.schedule_slot_id,a.updated_at
                   FROM center_attendance a
                   JOIN center_courses c ON c.id=a.course_id
                   JOIN users u ON u.user_id=a.student_user_id
                   WHERE a.context_id=%s AND a.id>%s
                     AND (%s IS NULL OR a.course_id=%s)
                     AND (%s IS NULL OR a.lesson_date=%s)
                     AND (
                       %s OR a.student_user_id=%s OR (
                         c.teacher_user_id=%s AND EXISTS(
                           SELECT 1 FROM center_role_assignments teacher_role
                           WHERE teacher_role.context_id=a.context_id
                             AND teacher_role.user_id=%s
                             AND teacher_role.role_key='teacher'
                             AND teacher_role.status='active'
                             AND teacher_role.starts_at<=NOW()
                             AND (
                               teacher_role.ends_at IS NULL
                               OR teacher_role.ends_at>NOW()
                             )
                             AND (
                               teacher_role.branch_id IS NULL
                               OR teacher_role.branch_id IS NOT DISTINCT FROM c.branch_id
                             )
                         )
                       )
                       OR EXISTS(
                         SELECT 1 FROM center_parent_links p
                         WHERE p.context_id=a.context_id
                           AND p.parent_user_id=%s
                           AND p.student_user_id=a.student_user_id
                           AND p.status='active'
                       )
                       OR EXISTS(
                         SELECT 1 FROM center_role_assignments r
                         WHERE r.context_id=a.context_id AND r.user_id=%s
                           AND r.status='active' AND r.starts_at<=NOW()
                           AND (r.ends_at IS NULL OR r.ends_at>NOW())
                           AND r.role_key=ANY(%s)
                           AND (
                             r.branch_id IS NULL
                             OR (c.branch_id IS NOT NULL AND r.branch_id=c.branch_id)
                           )
                       )
                     )
                   ORDER BY a.id LIMIT %s""",
                (
                    context_id, after_id or 0, course_id, course_id,
                    date_value, date_value, "system_admin" in roles,
                    user_id, user_id, user_id, user_id, user_id,
                    list(MANAGER_ROLES | {"methodist"}), limit + 1,
                ),
            )
            rows = cur.fetchall()
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
        }

    @router.post("/attendance")
    def mark_attendance(
        request: AttendanceMark, user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        attendance_date = request.attendance_date or request.lesson_date
        if attendance_date is None:
            raise HTTPException(status_code=422, detail="Davomat sanasi kerak")
        status = {
            "absent_excused": "excused",
            "absent_unexcused": "absent",
        }.get(request.status, request.status)
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            course_access(cur, request.context_id, request.course_id, user_id, write=True)
            enrollment = active_enrollment(
                cur, request.context_id, request.course_id, request.student_user_id
            )
            if request.schedule_slot_id is not None:
                cur.execute(
                    """SELECT 1 FROM center_schedule_slots
                       WHERE id=%s AND context_id=%s AND course_id=%s""",
                    (
                        request.schedule_slot_id, request.context_id,
                        request.course_id,
                    ),
                )
                if not cur.fetchone():
                    raise HTTPException(status_code=404, detail="Dars jadvali topilmadi")
            cur.execute(
                """INSERT INTO center_attendance(
                     context_id,course_id,enrollment_id,schedule_slot_id,
                     lesson_date,student_user_id,status,note,marked_by_user_id
                   ) VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s)
                   ON CONFLICT(context_id,course_id,student_user_id,lesson_date)
                   DO UPDATE SET status=EXCLUDED.status,note=EXCLUDED.note,
                     schedule_slot_id=EXCLUDED.schedule_slot_id,
                     marked_by_user_id=EXCLUDED.marked_by_user_id,updated_at=NOW()
                   RETURNING *""",
                (
                    request.context_id, request.course_id, enrollment["id"],
                    request.schedule_slot_id, attendance_date,
                    request.student_user_id, status, request.note, user_id,
                ),
            )
            item = cur.fetchone()
            audit(
                cur, request.context_id, user_id, "attendance.mark",
                "attendance", item["id"],
            )
        return {"item": item}

    @router.get("/grades")
    def list_grades(
        context_id: int = Query(ge=1),
        course_id: int | None = Query(default=None, ge=1),
        date_value: date | None = Query(default=None, alias="date"),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=100, ge=1, le=300),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            roles = require_roles(cur, context_id, user_id, VIEW_ROLES)
            if course_id:
                course_access(cur, context_id, course_id, user_id)
            cur.execute(
                """SELECT g.id,g.course_id,c.name course_name,g.student_user_id,
                          u.full_name student_name,g.grade_type,g.score,g.max_score,
                          g.graded_at,g.graded_at::DATE grade_date,g.note
                   FROM center_grade_entries g
                   JOIN center_courses c ON c.id=g.course_id
                   JOIN users u ON u.user_id=g.student_user_id
                   WHERE g.context_id=%s AND g.id>%s
                     AND (%s IS NULL OR g.course_id=%s)
                     AND (%s IS NULL OR g.graded_at::DATE=%s)
                     AND (
                       %s OR g.student_user_id=%s OR (
                         c.teacher_user_id=%s AND EXISTS(
                           SELECT 1 FROM center_role_assignments teacher_role
                           WHERE teacher_role.context_id=g.context_id
                             AND teacher_role.user_id=%s
                             AND teacher_role.role_key='teacher'
                             AND teacher_role.status='active'
                             AND teacher_role.starts_at<=NOW()
                             AND (
                               teacher_role.ends_at IS NULL
                               OR teacher_role.ends_at>NOW()
                             )
                             AND (
                               teacher_role.branch_id IS NULL
                               OR teacher_role.branch_id IS NOT DISTINCT FROM c.branch_id
                             )
                         )
                       )
                       OR EXISTS(
                         SELECT 1 FROM center_parent_links p
                         WHERE p.context_id=g.context_id
                           AND p.parent_user_id=%s
                           AND p.student_user_id=g.student_user_id
                           AND p.status='active'
                       )
                       OR EXISTS(
                         SELECT 1 FROM center_role_assignments r
                         WHERE r.context_id=g.context_id AND r.user_id=%s
                           AND r.status='active' AND r.starts_at<=NOW()
                           AND (r.ends_at IS NULL OR r.ends_at>NOW())
                           AND r.role_key=ANY(%s)
                           AND (
                             r.branch_id IS NULL
                             OR (c.branch_id IS NOT NULL AND r.branch_id=c.branch_id)
                           )
                       )
                     )
                   ORDER BY g.id LIMIT %s""",
                (
                    context_id, after_id or 0, course_id, course_id,
                    date_value, date_value, "system_admin" in roles,
                    user_id, user_id, user_id, user_id, user_id,
                    list(MANAGER_ROLES | {"methodist"}), limit + 1,
                ),
            )
            rows = cur.fetchall()
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
        }

    @router.post("/grades")
    def create_grade(
        request: GradeCreate,
        idempotency_header: str | None = Header(
            default=None, alias="Idempotency-Key"
        ),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        if request.score > request.max_score:
            raise HTTPException(status_code=422, detail="Ball maksimal balldan oshdi")
        key = request.idempotency_key or idempotency_header
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            course_access(cur, request.context_id, request.course_id, user_id, write=True)
            enrollment = active_enrollment(
                cur, request.context_id, request.course_id, request.student_user_id
            )
            if key:
                existing = claim_request(
                    cur, context_id=request.context_id, actor=user_id,
                    key=key, action="grade.create",
                )
                if existing:
                    cur.execute(
                        "SELECT * FROM center_grade_entries WHERE id=%s",
                        (existing,),
                    )
                    return {"item": cur.fetchone(), "idempotent_replay": True}
            note = request.note
            if request.assessment_name:
                note = f"{request.assessment_name}: {note or ''}".strip()
            graded_at = (
                datetime.combine(
                    request.grade_date, time(12), tzinfo=BUSINESS_TIMEZONE
                )
                if request.grade_date else datetime.now(timezone.utc)
            )
            cur.execute(
                """INSERT INTO center_grade_entries(
                     context_id,course_id,enrollment_id,student_user_id,
                     teacher_user_id,grade_type,score,max_score,graded_at,note,
                     idempotency_key
                   ) VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                   RETURNING *""",
                (
                    request.context_id, request.course_id, enrollment["id"],
                    request.student_user_id, user_id, request.grade_type,
                    request.score, request.max_score, graded_at, note, key,
                ),
            )
            item = cur.fetchone()
            if key:
                finish_request(cur, request.context_id, key, item["id"])
            audit(
                cur, request.context_id, user_id, "grade.create", "grade", item["id"]
            )
        return {"item": item}

    def list_course_documents(
        cur: Any, *, table: str, context_id: int, user_id: int,
        course_id: int | None, after_id: int, limit: int,
    ) -> list[Any]:
        roles = require_roles(cur, context_id, user_id, VIEW_ROLES)
        if course_id:
            course_access(cur, context_id, course_id, user_id)
        if table == "center_lesson_plans":
            fields = (
                "d.id,d.course_id,c.name course_name,d.teacher_user_id,"
                "u.full_name teacher_name,d.lesson_date,d.title,d.objectives,"
                "d.stages,d.content_text,d.content_latex,d.source_refs,"
                "d.duration_minutes,d.status,d.version,d.published_at,d.updated_at"
            )
        else:
            fields = (
                "d.id,d.course_id,c.name course_name,d.teacher_user_id,"
                "u.full_name teacher_name,d.title,d.instructions,d.content_latex,"
                "d.resource_refs,d.due_at,d.max_score,d.status,d.published_at,d.updated_at"
            )
        cur.execute(
            f"""SELECT {fields}
                FROM {table} d
                JOIN center_courses c ON c.id=d.course_id
                JOIN users u ON u.user_id=d.teacher_user_id
                WHERE d.context_id=%s AND d.id>%s
                  AND (%s IS NULL OR d.course_id=%s)
                  AND (
                    %s OR (
                      d.teacher_user_id=%s AND EXISTS(
                        SELECT 1 FROM center_role_assignments teacher_role
                        WHERE teacher_role.context_id=d.context_id
                          AND teacher_role.user_id=%s
                          AND teacher_role.role_key='teacher'
                          AND teacher_role.status='active'
                          AND teacher_role.starts_at<=NOW()
                          AND (
                            teacher_role.ends_at IS NULL
                            OR teacher_role.ends_at>NOW()
                          )
                          AND (
                            teacher_role.branch_id IS NULL
                            OR teacher_role.branch_id IS NOT DISTINCT FROM c.branch_id
                          )
                      )
                    )
                    OR EXISTS(
                      SELECT 1 FROM center_enrollments e
                      WHERE e.context_id=d.context_id AND e.course_id=d.course_id
                        AND e.student_user_id=%s AND e.status IN ('active','completed')
                        AND (
                          e.starts_on IS NULL
                          OR e.starts_on<=
                            (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                        )
                        AND d.status='published'
                    )
                    OR EXISTS(
                      SELECT 1 FROM center_parent_links p
                      JOIN center_enrollments e
                        ON e.context_id=p.context_id
                       AND e.student_user_id=p.student_user_id
                      WHERE p.context_id=d.context_id AND p.parent_user_id=%s
                        AND p.status='active' AND e.course_id=d.course_id
                        AND e.status IN ('active','completed')
                        AND (
                          e.starts_on IS NULL
                          OR e.starts_on<=
                            (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                        )
                        AND d.status='published'
                    )
                    OR EXISTS(
                      SELECT 1 FROM center_role_assignments r
                      WHERE r.context_id=d.context_id AND r.user_id=%s
                        AND r.status='active' AND r.starts_at<=NOW()
                        AND (r.ends_at IS NULL OR r.ends_at>NOW())
                        AND r.role_key=ANY(%s)
                        AND (
                          r.branch_id IS NULL
                          OR (c.branch_id IS NOT NULL AND r.branch_id=c.branch_id)
                        )
                    )
                  )
                ORDER BY d.id LIMIT %s""",
            (
                context_id, after_id, course_id, course_id,
                "system_admin" in roles, user_id, user_id,
                user_id, user_id, user_id,
                list(MANAGER_ROLES | {"methodist"}), limit + 1,
            ),
        )
        return cur.fetchall()

    @router.get("/lesson-plans")
    def list_lesson_plans(
        context_id: int = Query(ge=1),
        course_id: int | None = Query(default=None, ge=1),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=50, ge=1, le=200),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            rows = list_course_documents(
                cur, table="center_lesson_plans", context_id=context_id,
                user_id=user_id, course_id=course_id, after_id=after_id or 0,
                limit=limit,
            )
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
        }

    @router.post("/lesson-plans")
    def create_lesson_plan(
        request: LessonPlanCreate,
        idempotency_key: str | None = Header(
            default=None, alias="Idempotency-Key"
        ),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        idempotency_key = validate_request_key(idempotency_key)
        require_encoded_size(
            {
                "objectives": request.objectives,
                "stages": request.stages,
                "source_refs": request.source_refs,
            },
            maximum=250_000,
            label="Dars ishlanmasi tuzilmasi",
        )
        objectives = request.objectives or (
            [request.objective] if request.objective else []
        )
        stages = request.stages or [
            item for item in (
                {"name": "Tushuntirish", "content": request.explanation}
                if request.explanation else None,
                {"name": "Faoliyat va mashqlar", "content": request.activities}
                if request.activities else None,
            ) if item
        ]
        content_latex = request.content_latex or request.formula_latex
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            course = course_access(
                cur, request.context_id, request.course_id, user_id, write=True
            )
            if idempotency_key:
                existing_target = claim_request(
                    cur, context_id=request.context_id, actor=user_id,
                    key=idempotency_key, action="lesson_plan.create",
                )
                if existing_target:
                    cur.execute(
                        """SELECT * FROM center_lesson_plans
                           WHERE id=%s AND context_id=%s""",
                        (existing_target, request.context_id),
                    )
                    return {
                        "item": cur.fetchone(),
                        "idempotent_replay": True,
                    }
            cur.execute(
                """INSERT INTO center_lesson_plans(
                     context_id,course_id,teacher_user_id,lesson_date,title,
                     objectives,stages,content_text,content_latex,source_refs,
                     duration_minutes
                   ) VALUES(%s,%s,%s,%s,%s,%s::jsonb,%s::jsonb,%s,%s,%s::jsonb,%s)
                   RETURNING *""",
                (
                    request.context_id, request.course_id,
                    course["teacher_user_id"] or user_id, request.lesson_date,
                    request.title.strip(),
                    json.dumps(objectives, ensure_ascii=False),
                    json.dumps(stages, ensure_ascii=False),
                    request.content_text or request.explanation,
                    content_latex,
                    json.dumps(request.source_refs, ensure_ascii=False),
                    request.duration_minutes,
                ),
            )
            item = cur.fetchone()
            if idempotency_key:
                finish_request(
                    cur, request.context_id, idempotency_key, item["id"]
                )
            audit(
                cur, request.context_id, user_id, "lesson_plan.create",
                "lesson_plan", item["id"],
            )
        return {"item": item}

    @router.post("/lesson-plans/{plan_id}/publish")
    def publish_lesson_plan(
        plan_id: int, request: HumanConfirm,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(request.confirmation, "Dars rejasini e'lon qilish uchun tasdiq kerak.")
        with database() as (_, cur):
            ensure_schema(cur)
            cur.execute(
                """SELECT * FROM center_lesson_plans WHERE id=%s FOR UPDATE""",
                (plan_id,),
            )
            plan = cur.fetchone()
            if not plan:
                raise HTTPException(status_code=404, detail="Dars rejasi topilmadi")
            require_operational(cur, plan["context_id"])
            course_access(cur, plan["context_id"], plan["course_id"], user_id, write=True)
            cur.execute(
                """UPDATE center_lesson_plans SET status='published',
                     published_at=NOW(),version=version+1,updated_at=NOW()
                   WHERE id=%s RETURNING *""",
                (plan_id,),
            )
            item = cur.fetchone()
            audit(
                cur, plan["context_id"], user_id, "lesson_plan.publish",
                "lesson_plan", plan_id,
            )
        return {"item": item}

    @router.get("/lesson-plans/{plan_id}/docx")
    def lesson_plan_docx(
        plan_id: int,
        context_id: int = Query(ge=1),
        user_id: int = Depends(authenticated_user),
    ) -> Response:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            cur.execute(
                """SELECT p.*,c.name course_name,s.name subject_name,
                          u.full_name teacher_name,lc.name center_name
                   FROM center_lesson_plans p
                   JOIN center_courses c ON c.id=p.course_id
                   JOIN center_subjects s ON s.id=c.subject_id
                   JOIN users u ON u.user_id=p.teacher_user_id
                   JOIN learning_contexts lc ON lc.id=p.context_id
                   WHERE p.id=%s AND p.context_id=%s""",
                (plan_id, context_id),
            )
            plan = cur.fetchone()
            if not plan:
                raise HTTPException(status_code=404, detail="Dars rejasi topilmadi")
            course_access(cur, context_id, plan["course_id"], user_id)
            roles, _ = active_roles(cur, context_id, user_id)
            if (
                plan["status"] != "published"
                and not ("system_admin" in roles or roles & ACADEMIC_ROLES)
            ):
                raise HTTPException(status_code=403, detail="Qoralama Word fayli yopiq")
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Cm, Pt
        except ImportError as exc:
            raise HTTPException(
                status_code=503, detail="python-docx kutubxonasi o'rnatilmagan"
            ) from exc
        document = Document()
        section = document.sections[0]
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(str(plan["title"]))
        run.bold = True
        run.font.size = Pt(16)
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(
            f"{plan['center_name']} · {plan['course_name']} · {plan['subject_name']}"
        ).italic = True
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in (
            ("O'qituvchi", plan["teacher_name"]),
            ("Sana", plan["lesson_date"] or "Belgilanmagan"),
            ("Davomiyligi", f"{plan['duration_minutes']} daqiqa"),
            ("Holat", plan["status"]),
        ):
            cells = table.add_row().cells
            cells[0].text = str(label)
            cells[1].text = str(value)
        objectives = list(plan["objectives"] or [])
        if objectives:
            document.add_heading("Dars maqsadlari", level=1)
            for objective in objectives:
                text_value = (
                    objective.get("text") or objective.get("title")
                    if isinstance(objective, dict)
                    else objective
                )
                if text_value:
                    document.add_paragraph(str(text_value), style="List Bullet")
        stages = list(plan["stages"] or [])
        if stages:
            document.add_heading("Dars jarayoni", level=1)
            for index, stage in enumerate(stages, start=1):
                if isinstance(stage, dict):
                    document.add_heading(
                        f"{index}. {stage.get('title') or stage.get('name') or 'Bosqich'}",
                        level=2,
                    )
                    stage_text = (
                        stage.get("content")
                        or stage.get("description")
                        or stage.get("activity")
                        or ""
                    )
                    document.add_paragraph(str(stage_text))
                else:
                    document.add_paragraph(str(stage), style="List Number")
        if plan["content_text"]:
            document.add_heading("Tushuntirish", level=1)
            document.add_paragraph(str(plan["content_text"]))
        if plan["content_latex"]:
            document.add_heading("Formula (LaTeX)", level=1)
            paragraph = document.add_paragraph()
            latex_run = paragraph.add_run(str(plan["content_latex"]))
            latex_run.font.name = "Consolas"
        refs = list(plan["source_refs"] or [])
        if refs:
            document.add_heading("Manbalar", level=1)
            for ref in refs:
                document.add_paragraph(str(ref), style="List Bullet")
        stream = io.BytesIO()
        document.save(stream)
        safe_name = re.sub(r"[^A-Za-z0-9_-]+", "_", str(plan["title"]))[:80]
        return Response(
            content=stream.getvalue(),
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            headers={
                "Content-Disposition": f'attachment; filename="{safe_name or "dars_reja"}.docx"'
            },
        )

    @router.get("/homework")
    def list_homework(
        context_id: int = Query(ge=1),
        course_id: int | None = Query(default=None, ge=1),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=50, ge=1, le=200),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            rows = [
                dict(row)
                for row in list_course_documents(
                    cur, table="center_homework", context_id=context_id,
                    user_id=user_id, course_id=course_id,
                    after_id=after_id or 0, limit=limit,
                )
            ]
            homework_ids = [int(row["id"]) for row in rows[:limit]]
            enrolled_courses: set[int] = set()
            own_submissions: dict[int, dict[str, Any]] = {}
            if homework_ids:
                course_ids = sorted(
                    {int(row["course_id"]) for row in rows[:limit]}
                )
                cur.execute(
                    """SELECT course_id FROM center_enrollments
                       WHERE context_id=%s AND student_user_id=%s
                         AND course_id=ANY(%s)
                         AND status IN ('active','completed')
                         AND (
                           starts_on IS NULL
                           OR starts_on<=
                             (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                         )""",
                    (context_id, user_id, course_ids),
                )
                enrolled_courses = {
                    int(row["course_id"]) for row in cur.fetchall()
                }
                cur.execute(
                    """SELECT id,homework_id,status,submitted_at,score,feedback,
                              graded_at,updated_at
                       FROM center_homework_submissions
                       WHERE context_id=%s AND student_user_id=%s
                         AND homework_id=ANY(%s)""",
                    (context_id, user_id, homework_ids),
                )
                own_submissions = {
                    int(row["homework_id"]): dict(row)
                    for row in cur.fetchall()
                }
            now = datetime.now(timezone.utc)
            for row in rows[:limit]:
                own = own_submissions.get(int(row["id"]))
                deadline_open = (
                    row["due_at"] is None or now <= row["due_at"]
                )
                learner_access = int(row["course_id"]) in enrolled_courses
                row["my_submission_id"] = own["id"] if own else None
                row["my_submission_status"] = own["status"] if own else None
                row["my_submission_score"] = own["score"] if own else None
                row["my_submission_feedback"] = (
                    own["feedback"] if own else None
                )
                row["my_submission_submitted_at"] = (
                    own["submitted_at"] if own else None
                )
                row["my_submission_graded_at"] = (
                    own["graded_at"] if own else None
                )
                row["can_submit"] = bool(
                    learner_access
                    and deadline_open
                    and (own is None or own["status"] == "returned")
                )
                row["can_resubmit"] = bool(
                    learner_access and deadline_open and own is not None
                    and own["status"] == "returned"
                )
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
        }

    @router.post("/homework")
    def create_homework(
        request: HomeworkCreate,
        idempotency_key: str | None = Header(
            default=None, alias="Idempotency-Key"
        ),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        idempotency_key = validate_request_key(idempotency_key)
        require_encoded_size(
            request.resource_refs, maximum=100_000,
            label="Uy vazifasi manbalari",
        )
        if request.status == "published":
            require_human(request.confirmation, "Vazifani e'lon qilish uchun tasdiq kerak.")
        due_at = request.due_at
        if due_at is not None and due_at.tzinfo is None:
            due_at = due_at.replace(tzinfo=BUSINESS_TIMEZONE)
        if due_at is None and request.due_date is not None:
            due_at = datetime.combine(
                request.due_date, time(23, 59), tzinfo=BUSINESS_TIMEZONE
            )
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            course = course_access(
                cur, request.context_id, request.course_id, user_id, write=True
            )
            if request.lesson_plan_id is not None:
                cur.execute(
                    """SELECT 1 FROM center_lesson_plans
                       WHERE id=%s AND context_id=%s AND course_id=%s""",
                    (
                        request.lesson_plan_id, request.context_id,
                        request.course_id,
                    ),
                )
                if not cur.fetchone():
                    raise HTTPException(status_code=404, detail="Dars rejasi mos emas")
            if idempotency_key:
                existing_target = claim_request(
                    cur, context_id=request.context_id, actor=user_id,
                    key=idempotency_key, action="homework.create",
                )
                if existing_target:
                    cur.execute(
                        """SELECT * FROM center_homework
                           WHERE id=%s AND context_id=%s""",
                        (existing_target, request.context_id),
                    )
                    return {
                        "item": cur.fetchone(),
                        "idempotent_replay": True,
                    }
            cur.execute(
                """INSERT INTO center_homework(
                     context_id,course_id,lesson_plan_id,teacher_user_id,title,
                     instructions,content_latex,resource_refs,due_at,max_score,
                     status,published_at
                   ) VALUES(
                     %s,%s,%s,%s,%s,%s,%s,%s::jsonb,%s,%s,%s,
                     CASE WHEN %s='published' THEN NOW() END
                   ) RETURNING *""",
                (
                    request.context_id, request.course_id, request.lesson_plan_id,
                    course["teacher_user_id"] or user_id, request.title.strip(),
                    request.instructions or request.homework_text,
                    request.content_latex or request.formula_latex,
                    json.dumps(request.resource_refs, ensure_ascii=False),
                    due_at, request.max_score, request.status, request.status,
                ),
            )
            item = cur.fetchone()
            if idempotency_key:
                finish_request(
                    cur, request.context_id, idempotency_key, item["id"]
                )
            audit(
                cur, request.context_id, user_id, "homework.create",
                "homework", item["id"], {"status": request.status},
            )
        return {"item": item}

    @router.post("/homework/{homework_id}/publish")
    def publish_homework(
        homework_id: int, request: HumanConfirm,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(request.confirmation, "Vazifani e'lon qilish uchun tasdiq kerak.")
        with database() as (_, cur):
            ensure_schema(cur)
            cur.execute("SELECT * FROM center_homework WHERE id=%s FOR UPDATE", (homework_id,))
            homework = cur.fetchone()
            if not homework:
                raise HTTPException(status_code=404, detail="Vazifa topilmadi")
            require_operational(cur, homework["context_id"])
            course_access(
                cur, homework["context_id"], homework["course_id"], user_id, write=True
            )
            cur.execute(
                """UPDATE center_homework SET status='published',
                     published_at=NOW(),updated_at=NOW()
                   WHERE id=%s RETURNING *""",
                (homework_id,),
            )
            item = cur.fetchone()
            audit(
                cur, homework["context_id"], user_id, "homework.publish",
                "homework", homework_id,
            )
        return {"item": item}

    @router.post("/homework/{homework_id}/submissions")
    def submit_homework(
        homework_id: int, request: HomeworkSubmit,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_encoded_size(
            request.attachment_refs, maximum=100_000,
            label="Topshiriq ilovalari",
        )
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            cur.execute(
                """SELECT * FROM center_homework
                   WHERE id=%s AND context_id=%s AND status='published'""",
                (homework_id, request.context_id),
            )
            homework = cur.fetchone()
            if not homework:
                raise HTTPException(status_code=404, detail="E'lon qilingan vazifa topilmadi")
            if (
                homework["due_at"] is not None
                and datetime.now(timezone.utc) > homework["due_at"]
            ):
                raise HTTPException(
                    status_code=409, detail="Uy vazifasini topshirish muddati tugagan"
                )
            active_enrollment(
                cur, request.context_id, homework["course_id"], user_id
            )
            cur.execute(
                """SELECT status FROM center_homework_submissions
                   WHERE context_id=%s AND homework_id=%s
                     AND student_user_id=%s FOR UPDATE""",
                (request.context_id, homework_id, user_id),
            )
            previous = cur.fetchone()
            if previous and previous["status"] != "returned":
                raise HTTPException(
                    status_code=409,
                    detail=(
                        "Topshiriq qayta yuborilishi uchun o'qituvchi uni "
                        "tuzatishga qaytarishi kerak"
                    ),
                )
            cur.execute(
                """INSERT INTO center_homework_submissions(
                     context_id,homework_id,student_user_id,answer_text,
                     answer_latex,attachment_refs,status,submitted_at
                   ) VALUES(%s,%s,%s,%s,%s,%s::jsonb,'submitted',NOW())
                   ON CONFLICT(context_id,homework_id,student_user_id)
                   DO UPDATE SET answer_text=EXCLUDED.answer_text,
                     answer_latex=EXCLUDED.answer_latex,
                     attachment_refs=EXCLUDED.attachment_refs,status='submitted',
                     score=NULL,feedback=NULL,graded_at=NULL,
                     graded_by_user_id=NULL,
                     submitted_at=NOW(),updated_at=NOW()
                   RETURNING *""",
                (
                    request.context_id, homework_id, user_id, request.answer_text,
                    request.answer_latex,
                    json.dumps(request.attachment_refs, ensure_ascii=False),
                ),
            )
            item = cur.fetchone()
            audit(
                cur, request.context_id, user_id, "homework.submit",
                "homework_submission", item["id"],
            )
        return {"item": item}

    @router.get("/homework/{homework_id}/submissions")
    def list_homework_submissions(
        homework_id: int,
        context_id: int = Query(ge=1),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=100, ge=1, le=300),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            cur.execute(
                """SELECT course_id FROM center_homework
                   WHERE id=%s AND context_id=%s""",
                (homework_id, context_id),
            )
            homework = cur.fetchone()
            if not homework:
                raise HTTPException(status_code=404, detail="Vazifa topilmadi")
            course_access(cur, context_id, homework["course_id"], user_id, write=True)
            cur.execute(
                """SELECT s.*,u.full_name student_name
                   FROM center_homework_submissions s
                   JOIN users u ON u.user_id=s.student_user_id
                   WHERE s.context_id=%s AND s.homework_id=%s AND s.id>%s
                   ORDER BY s.id LIMIT %s""",
                (context_id, homework_id, after_id or 0, limit + 1),
            )
            rows = cur.fetchall()
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
        }

    @router.get("/homework/{homework_id}/my-submission")
    def get_my_homework_submission(
        homework_id: int,
        context_id: int = Query(ge=1),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            cur.execute(
                """SELECT id,course_id,due_at,status
                   FROM center_homework
                   WHERE id=%s AND context_id=%s AND status='published'""",
                (homework_id, context_id),
            )
            homework = cur.fetchone()
            if not homework:
                raise HTTPException(
                    status_code=404,
                    detail="E'lon qilingan vazifa topilmadi",
                )
            active_enrollment(
                cur, context_id, homework["course_id"], user_id
            )
            cur.execute(
                """SELECT id,homework_id,status,answer_text,answer_latex,
                          attachment_refs,submitted_at,score,feedback,
                          graded_at,updated_at
                   FROM center_homework_submissions
                   WHERE context_id=%s AND homework_id=%s
                     AND student_user_id=%s""",
                (context_id, homework_id, user_id),
            )
            item = cur.fetchone()
            deadline_open = (
                homework["due_at"] is None
                or datetime.now(timezone.utc) <= homework["due_at"]
            )
        return {
            "item": item,
            "scope": "self",
            "due_at": homework["due_at"],
            "can_submit": bool(
                deadline_open
                and (item is None or item["status"] == "returned")
            ),
            "can_resubmit": bool(
                deadline_open and item is not None
                and item["status"] == "returned"
            ),
        }

    @router.post("/homework-submissions/{submission_id}/grade")
    def grade_homework_submission(
        submission_id: int, request: HomeworkGrade,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(request.confirmation, "Vazifa bahosi uchun inson tasdig'i kerak.")
        with database() as (_, cur):
            ensure_schema(cur)
            cur.execute(
                """SELECT s.*,h.course_id,h.max_score
                   FROM center_homework_submissions s
                   JOIN center_homework h ON h.id=s.homework_id
                   WHERE s.id=%s AND s.context_id=%s FOR UPDATE""",
                (submission_id, request.context_id),
            )
            submission = cur.fetchone()
            if not submission:
                raise HTTPException(status_code=404, detail="Topshiriq topilmadi")
            course_access(
                cur, request.context_id, submission["course_id"], user_id, write=True
            )
            if request.score > submission["max_score"]:
                raise HTTPException(status_code=422, detail="Ball maksimal balldan oshdi")
            cur.execute(
                """UPDATE center_homework_submissions SET status='graded',
                     score=%s,feedback=%s,graded_at=NOW(),graded_by_user_id=%s,
                     updated_at=NOW()
                   WHERE id=%s RETURNING *""",
                (request.score, request.feedback, user_id, submission_id),
            )
            item = cur.fetchone()
            audit(
                cur, request.context_id, user_id, "homework.grade",
                "homework_submission", submission_id,
            )
        return {"item": item}

    @router.post("/homework-submissions/{submission_id}/return")
    def return_homework_submission(
        submission_id: int,
        request: HumanConfirm,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(
            request.confirmation,
            "Topshiriqni tuzatishga qaytarish uchun inson tasdig'i kerak.",
        )
        with database() as (_, cur):
            ensure_schema(cur)
            cur.execute(
                """SELECT s.*,h.course_id FROM center_homework_submissions s
                   JOIN center_homework h ON h.id=s.homework_id
                   WHERE s.id=%s FOR UPDATE""",
                (submission_id,),
            )
            submission = cur.fetchone()
            if not submission:
                raise HTTPException(
                    status_code=404, detail="Topshiriq topilmadi"
                )
            require_operational(cur, submission["context_id"])
            course_access(
                cur, submission["context_id"], submission["course_id"],
                user_id, write=True,
            )
            if submission["status"] not in {"submitted", "graded"}:
                raise HTTPException(
                    status_code=409,
                    detail="Bu topshiriq allaqachon tuzatishga qaytarilgan",
                )
            cur.execute(
                """UPDATE center_homework_submissions SET status='returned',
                     updated_at=NOW()
                   WHERE id=%s RETURNING *""",
                (submission_id,),
            )
            item = cur.fetchone()
            audit(
                cur, submission["context_id"], user_id, "homework.return",
                "homework_submission", submission_id,
            )
        return {"item": item}

    def optional_datetime(value: datetime | str | None, field: str) -> datetime | None:
        if value in (None, ""):
            return None
        if isinstance(value, datetime):
            parsed = value
        else:
            try:
                parsed = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
            except ValueError as exc:
                raise HTTPException(
                    status_code=422, detail=f"{field} ISO sana-vaqt bo'lsin"
                ) from exc
        if parsed.tzinfo is None:
            parsed = parsed.replace(tzinfo=BUSINESS_TIMEZONE)
        return parsed

    def validate_assessment_item(
        question_source: str,
        question_ref: str,
        metadata: dict[str, Any],
        item_order: int,
    ) -> None:
        if question_source not in {"manual_center", "generated_tests"}:
            raise HTTPException(
                status_code=422,
                detail=f"{item_order}-savol manbasi qo'llab-quvvatlanmaydi",
            )
        if question_source == "generated_tests":
            if not str(question_ref).isdigit():
                raise HTTPException(
                    status_code=422,
                    detail=f"{item_order}-savol uchun test bank ID raqam bo'lsin",
                )
            return
        prompt = str(metadata.get("prompt") or "").strip()
        if not prompt:
            raise HTTPException(
                status_code=422,
                detail=f"{item_order}-savol matni bo'sh",
            )
        question_type = str(metadata.get("question_type") or "")
        if question_type not in {"multiple_choice", "short_answer"}:
            raise HTTPException(
                status_code=422,
                detail=f"{item_order}-savol turi noto'g'ri",
            )
        correct_answer = str(metadata.get("correct_answer") or "").strip()
        if question_type == "short_answer":
            if not correct_answer:
                raise HTTPException(
                    status_code=422,
                    detail=f"{item_order}-qisqa savolda javob kaliti kerak",
                )
            return
        raw_options = metadata.get("options")
        if not isinstance(raw_options, list) or len(raw_options) > 20:
            raise HTTPException(
                status_code=422,
                detail=f"{item_order}-variantlar ro'yxati noto'g'ri",
            )
        nonempty_indexes = {
            index
            for index, option in enumerate(raw_options)
            if str(option or "").strip()
        }
        if len(nonempty_indexes) < 2:
            raise HTTPException(
                status_code=422,
                detail=f"{item_order}-savolda kamida 2 ta variant bo'lsin",
            )
        if not correct_answer.isdigit() or int(correct_answer) not in nonempty_indexes:
            raise HTTPException(
                status_code=422,
                detail=f"{item_order}-savolda to'g'ri variantni tanlang",
            )

    def normalize_assessment_item_metadata(
        question_source: str, metadata: dict[str, Any],
    ) -> dict[str, Any]:
        normalized = dict(metadata)
        if (
            question_source == "manual_center"
            and normalized.get("question_type") == "multiple_choice"
            and isinstance(normalized.get("options"), list)
        ):
            raw_options = list(normalized["options"])
            correct_raw = str(normalized.get("correct_answer") or "").strip()
            if correct_raw.isdigit() and int(correct_raw) < len(raw_options):
                correct_value = str(raw_options[int(correct_raw)] or "").strip()
                cleaned = [
                    str(option).strip()
                    for option in raw_options
                    if str(option or "").strip()
                ]
                if correct_value in cleaned:
                    normalized["options"] = cleaned
                    normalized["correct_answer"] = str(
                        cleaned.index(correct_value)
                    )
        return normalized

    def public_assessment_items(
        cur: Any, context_id: int, assessment_id: int,
    ) -> list[dict[str, Any]]:
        cur.execute(
            """SELECT id,question_ref,question_source,item_order,points,
                      section_key,metadata
               FROM center_assessment_items
               WHERE context_id=%s AND assessment_id=%s
               ORDER BY item_order,id""",
            (context_id, assessment_id),
        )
        items = [dict(row) for row in cur.fetchall()]
        generated_ids = [
            int(item["question_ref"])
            for item in items
            if item["question_source"] == "generated_tests"
            and str(item["question_ref"]).isdigit()
        ]
        generated: dict[int, dict[str, Any]] = {}
        if generated_ids:
            cur.execute(
                """SELECT id,question,option_a,option_b,option_c,option_d,
                          question_type,is_latex,time_limit,difficulty,
                          CASE WHEN rasm_malumot IS NOT NULL
                            THEN '/api/test_rasmi/'||id::TEXT
                            ELSE COALESCE(NULLIF(image_url,''),
                                         NULLIF(image_file_id,''))
                          END image_url
                   FROM generated_tests WHERE id=ANY(%s)""",
                (generated_ids,),
            )
            generated = {int(row["id"]): dict(row) for row in cur.fetchall()}
        for item in items:
            raw_metadata = dict(item.get("metadata") or {})
            allowed_metadata = {
                "prompt", "question_type", "content_latex", "formula_latex",
                "options", "image_url",
            }
            safe_metadata = {
                key: value for key, value in raw_metadata.items()
                if key in allowed_metadata
            }
            item["metadata"] = safe_metadata
            rendered = None
            if (
                item["question_source"] == "generated_tests"
                and str(item["question_ref"]).isdigit()
            ):
                rendered = generated.get(int(item["question_ref"]))
            item["question"] = rendered
            if rendered:
                item["question_text"] = rendered.get("question")
                item["question_type"] = rendered.get("question_type")
                item["formula_latex"] = (
                    rendered.get("question") if rendered.get("is_latex") else None
                )
                item["options"] = [
                    value for value in (
                        rendered.get("option_a"), rendered.get("option_b"),
                        rendered.get("option_c"), rendered.get("option_d"),
                    ) if value is not None
                ]
        return items

    @router.get("/assessments")
    def list_assessments(
        context_id: int = Query(ge=1),
        course_id: int | None = Query(default=None, ge=1),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=50, ge=1, le=200),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            roles = require_roles(cur, context_id, user_id, VIEW_ROLES)
            if course_id:
                course_access(cur, context_id, course_id, user_id)
            cur.execute(
                """SELECT a.id,a.course_id,c.name course_name,a.assessment_type,
                          a.title,a.instructions,a.framework,a.duration_minutes,
                          a.max_attempts,a.total_points,a.total_points total_score,
                          a.opens_at,a.closes_at,
                          a.status,a.settings,a.published_at,
                          (SELECT COUNT(*) FROM center_assessment_items i
                           WHERE i.context_id=a.context_id
                             AND i.assessment_id=a.id) item_count,
                          (SELECT COUNT(*) FROM center_assessment_items i
                           WHERE i.context_id=a.context_id
                             AND i.assessment_id=a.id) question_count
                   FROM center_assessments a
                   JOIN center_courses c ON c.id=a.course_id
                   WHERE a.context_id=%s AND a.id>%s
                     AND (%s IS NULL OR a.course_id=%s)
                     AND (
                       %s OR (
                         c.teacher_user_id=%s AND EXISTS(
                           SELECT 1 FROM center_role_assignments teacher_role
                           WHERE teacher_role.context_id=a.context_id
                             AND teacher_role.user_id=%s
                             AND teacher_role.role_key='teacher'
                             AND teacher_role.status='active'
                             AND teacher_role.starts_at<=NOW()
                             AND (
                               teacher_role.ends_at IS NULL
                               OR teacher_role.ends_at>NOW()
                             )
                             AND (
                               teacher_role.branch_id IS NULL
                               OR teacher_role.branch_id IS NOT DISTINCT FROM c.branch_id
                             )
                         )
                       )
                       OR EXISTS(
                         SELECT 1 FROM center_enrollments e
                         WHERE e.context_id=a.context_id AND e.course_id=a.course_id
                           AND e.student_user_id=%s AND e.status IN ('active','completed')
                           AND (
                             e.starts_on IS NULL
                             OR e.starts_on<=
                               (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                           )
                           AND a.status='published'
                       )
                       OR EXISTS(
                         SELECT 1 FROM center_parent_links p
                         JOIN center_enrollments e
                           ON e.context_id=p.context_id
                          AND e.student_user_id=p.student_user_id
                         WHERE p.context_id=a.context_id AND p.parent_user_id=%s
                           AND p.status='active' AND e.course_id=a.course_id
                           AND e.status IN ('active','completed')
                           AND (
                             e.starts_on IS NULL
                             OR e.starts_on<=
                               (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                           )
                           AND a.status='published'
                       )
                       OR EXISTS(
                         SELECT 1 FROM center_role_assignments r
                         WHERE r.context_id=a.context_id AND r.user_id=%s
                           AND r.status='active' AND r.starts_at<=NOW()
                           AND (r.ends_at IS NULL OR r.ends_at>NOW())
                           AND r.role_key=ANY(%s)
                           AND (
                             r.branch_id IS NULL
                             OR (c.branch_id IS NOT NULL AND r.branch_id=c.branch_id)
                           )
                       )
                     )
                   ORDER BY a.id LIMIT %s""",
                (
                    context_id, after_id or 0, course_id, course_id,
                    "system_admin" in roles, user_id, user_id,
                    user_id, user_id, user_id,
                    list(MANAGER_ROLES | {"methodist"}), limit + 1,
                ),
            )
            rows = cur.fetchall()
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
        }

    @router.post("/assessments")
    def create_assessment(
        request: AssessmentCreate,
        idempotency_key: str | None = Header(
            default=None, alias="Idempotency-Key"
        ),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        idempotency_key = validate_request_key(idempotency_key)
        required_framework = {
            "ielts_mock": "ielts",
            "cefr_mock": "cefr",
        }.get(request.assessment_type)
        if required_framework and request.framework != required_framework:
            raise HTTPException(
                status_code=422,
                detail=(
                    f"{request.assessment_type} uchun framework "
                    f"'{required_framework}' bo'lishi shart"
                ),
            )
        computed_total = sum(
            (item.points for item in request.items), Decimal("0")
        )
        if (
            request.total_score is not None
            and request.total_score != computed_total
        ):
            raise HTTPException(
                status_code=422,
                detail=(
                    "total_score savollar ballari yig'indisiga teng bo'lishi "
                    f"kerak ({computed_total})"
                ),
            )
        require_encoded_size(
            {
                "settings": request.settings,
                "items": [item.model_dump(mode="json") for item in request.items],
                "instructions": request.instructions,
            },
            maximum=400_000,
            label="Imtihon tuzilmasi",
        )
        opens_at = optional_datetime(request.opens_at, "opens_at")
        closes_at = optional_datetime(request.closes_at, "closes_at")
        if opens_at and closes_at and closes_at <= opens_at:
            raise HTTPException(status_code=422, detail="Imtihon yopilish vaqti kechroq bo'lsin")
        for index, item in enumerate(request.items, start=1):
            item.metadata = normalize_assessment_item_metadata(
                item.question_source, item.metadata
            )
            validate_assessment_item(
                item.question_source, item.question_ref, item.metadata, index
            )
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            course_access(cur, request.context_id, request.course_id, user_id, write=True)
            if idempotency_key:
                existing_target = claim_request(
                    cur, context_id=request.context_id, actor=user_id,
                    key=idempotency_key, action="assessment.create",
                )
                if existing_target:
                    cur.execute(
                        """SELECT * FROM center_assessments
                           WHERE id=%s AND context_id=%s""",
                        (existing_target, request.context_id),
                    )
                    return {
                        "item": cur.fetchone(),
                        "idempotent_replay": True,
                    }
            cur.execute(
                """INSERT INTO center_assessments(
                     context_id,course_id,created_by_user_id,assessment_type,
                     title,instructions,framework,duration_minutes,max_attempts,
                     total_points,opens_at,closes_at,status,settings
                   ) VALUES(
                     %s,%s,%s,%s,%s,%s,%s,%s,%s,0,%s,%s,'draft',%s::jsonb
                   ) RETURNING *""",
                (
                    request.context_id, request.course_id, user_id,
                    request.assessment_type, request.title.strip(),
                    request.instructions, request.framework,
                    request.duration_minutes, request.max_attempts,
                    opens_at, closes_at,
                    json.dumps(
                        {
                            **request.settings,
                            "formula_latex": request.formula_latex,
                        },
                        ensure_ascii=False,
                    ),
                ),
            )
            assessment = cur.fetchone()
            total = Decimal("0")
            for index, item in enumerate(request.items, start=1):
                cur.execute(
                    """INSERT INTO center_assessment_items(
                         context_id,assessment_id,question_ref,question_source,
                         item_order,points,section_key,metadata
                       ) VALUES(%s,%s,%s,%s,%s,%s,%s,%s::jsonb)""",
                    (
                        request.context_id, assessment["id"], item.question_ref,
                        item.question_source, index, item.points, item.section_key,
                        json.dumps(item.metadata, ensure_ascii=False),
                    ),
                )
                total += item.points
            if idempotency_key:
                finish_request(
                    cur, request.context_id, idempotency_key,
                    assessment["id"],
                )
            if request.items:
                cur.execute(
                    """UPDATE center_assessments SET total_points=%s,updated_at=NOW()
                       WHERE id=%s RETURNING *""",
                    (total, assessment["id"]),
                )
                assessment = cur.fetchone()
            audit(
                cur, request.context_id, user_id, "assessment.create",
                "assessment", assessment["id"], {"item_count": len(request.items)},
            )
        return {"item": assessment}

    @router.post("/assessments/{assessment_id}/publish")
    def publish_assessment(
        assessment_id: int, request: HumanConfirm,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(request.confirmation, "Testni e'lon qilish uchun tasdiq kerak.")
        with database() as (_, cur):
            ensure_schema(cur)
            cur.execute(
                """SELECT * FROM center_assessments WHERE id=%s FOR UPDATE""",
                (assessment_id,),
            )
            assessment = cur.fetchone()
            if not assessment:
                raise HTTPException(status_code=404, detail="Test topilmadi")
            require_operational(cur, assessment["context_id"])
            course_access(
                cur, assessment["context_id"], assessment["course_id"],
                user_id, write=True,
            )
            cur.execute(
                """SELECT COUNT(*) item_count,COALESCE(SUM(points),0) total
                   FROM center_assessment_items
                   WHERE context_id=%s AND assessment_id=%s""",
                (assessment["context_id"], assessment_id),
            )
            totals = cur.fetchone()
            if int(totals["item_count"]) < 1:
                raise HTTPException(
                    status_code=409,
                    detail="Savolsiz test e'lon qilinmaydi; savol banki havolalarini qo'shing.",
                )
            cur.execute(
                """SELECT question_source,question_ref,metadata,item_order
                   FROM center_assessment_items
                   WHERE context_id=%s AND assessment_id=%s
                   ORDER BY item_order,id""",
                (assessment["context_id"], assessment_id),
            )
            stored_items = cur.fetchall()
            generated_ids: list[int] = []
            for stored in stored_items:
                validate_assessment_item(
                    stored["question_source"], stored["question_ref"],
                    dict(stored["metadata"] or {}), int(stored["item_order"]),
                )
                if stored["question_source"] == "generated_tests":
                    generated_ids.append(int(stored["question_ref"]))
            if generated_ids:
                cur.execute(
                    "SELECT COUNT(*) found FROM generated_tests WHERE id=ANY(%s)",
                    (generated_ids,),
                )
                if int(cur.fetchone()["found"]) != len(set(generated_ids)):
                    raise HTTPException(
                        status_code=409,
                        detail="Test bankidagi ayrim savollar topilmadi",
                    )
            cur.execute(
                """UPDATE center_assessments SET status='published',
                     total_points=%s,published_at=NOW(),updated_at=NOW()
                   WHERE id=%s RETURNING *""",
                (totals["total"], assessment_id),
            )
            item = cur.fetchone()
            audit(
                cur, assessment["context_id"], user_id, "assessment.publish",
                "assessment", assessment_id,
            )
        return {"item": item}

    def read_attempt_for_user(
        cur: Any, attempt_id: int, context_id: int, user_id: int,
        *, allow_linked_parent: bool = False,
    ) -> tuple[dict[str, Any], dict[str, Any]]:
        cur.execute(
            """SELECT at.*,a.course_id,a.title assessment_title,a.framework,
                      a.duration_minutes,a.total_points,a.closes_at,
                      a.status assessment_status
               FROM center_assessment_attempts at
               JOIN center_assessments a ON a.id=at.assessment_id
               WHERE at.id=%s AND at.context_id=%s""",
            (attempt_id, context_id),
        )
        attempt = cur.fetchone()
        if not attempt:
            raise HTTPException(status_code=404, detail="Urinish topilmadi")
        if int(attempt["student_user_id"]) != user_id:
            linked_parent = False
            if allow_linked_parent:
                cur.execute(
                    """SELECT 1 FROM center_parent_links
                       WHERE context_id=%s AND parent_user_id=%s
                         AND student_user_id=%s AND status='active'""",
                    (context_id, user_id, attempt["student_user_id"]),
                )
                linked_parent = cur.fetchone() is not None
            if not linked_parent:
                course_access(
                    cur, context_id, attempt["course_id"], user_id, write=True
                )
        return dict(attempt), {
            "id": attempt["assessment_id"],
            "course_id": attempt["course_id"],
            "title": attempt["assessment_title"],
            "framework": attempt["framework"],
            "duration_minutes": attempt["duration_minutes"],
            "total_points": attempt["total_points"],
            "closes_at": attempt["closes_at"],
        }

    def attempt_deadline(attempt: dict[str, Any]) -> datetime | None:
        deadlines: list[datetime] = []
        if attempt.get("closes_at") is not None:
            deadlines.append(attempt["closes_at"])
        if (
            attempt.get("duration_minutes") is not None
            and attempt.get("started_at") is not None
        ):
            deadlines.append(
                attempt["started_at"]
                + timedelta(minutes=int(attempt["duration_minutes"]))
            )
        return min(deadlines) if deadlines else None

    def expire_attempt_if_late(
        cur: Any, attempt: dict[str, Any], *, now: datetime,
    ) -> dict[str, Any] | None:
        if attempt.get("status") != "in_progress":
            return None
        deadline = attempt_deadline(attempt)
        if deadline is None or now <= deadline:
            return None
        cur.execute(
            """UPDATE center_assessment_attempts SET status='expired',
                 submitted_at=COALESCE(submitted_at,NOW()),updated_at=NOW()
               WHERE id=%s AND context_id=%s AND status='in_progress'
               RETURNING *""",
            (attempt["id"], attempt["context_id"]),
        )
        expired = cur.fetchone()
        return dict(expired) if expired else None

    @router.post("/assessments/{assessment_id}/attempts")
    def start_attempt(
        assessment_id: int,
        idempotency_key: str | None = Header(default=None, alias="Idempotency-Key"),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        idempotency_key = validate_request_key(idempotency_key)
        with database() as (_, cur):
            ensure_schema(cur)
            cur.execute(
                """SELECT * FROM center_assessments
                   WHERE id=%s AND status='published'""",
                (assessment_id,),
            )
            assessment = cur.fetchone()
            if not assessment:
                raise HTTPException(status_code=404, detail="E'lon qilingan test topilmadi")
            require_operational(cur, assessment["context_id"])
            cur.execute(
                """SELECT pg_advisory_xact_lock(
                     74130,
                     hashtext('center-attempt-'||%s::TEXT||'-'||%s::TEXT||'-'||%s::TEXT)
                   )""",
                (assessment["context_id"], assessment_id, user_id),
            )
            now = datetime.now(timezone.utc)
            if assessment["opens_at"] and assessment["opens_at"] > now:
                raise HTTPException(status_code=409, detail="Test hali ochilmagan")
            cur.execute(
                """SELECT * FROM center_assessment_attempts
                   WHERE context_id=%s AND assessment_id=%s
                     AND student_user_id=%s AND status='in_progress'
                   ORDER BY id DESC LIMIT 1""",
                (assessment["context_id"], assessment_id, user_id),
            )
            in_progress = cur.fetchone()
            if in_progress:
                attempt_window = {
                    **dict(in_progress),
                    "duration_minutes": assessment["duration_minutes"],
                    "closes_at": assessment["closes_at"],
                }
                expired = expire_attempt_if_late(
                    cur, attempt_window, now=now
                )
                if expired:
                    audit(
                        cur, assessment["context_id"], user_id,
                        "attempt.expire", "assessment_attempt", expired["id"],
                        {"source": "attempt.start"},
                    )
                    return JSONResponse(
                        status_code=409,
                        content={
                            "detail": "Oldingi urinish muddati tugagan",
                            "code": "attempt_expired",
                            "attempt_id": int(expired["id"]),
                        },
                    )
            if assessment["closes_at"] and assessment["closes_at"] < now:
                raise HTTPException(status_code=409, detail="Test muddati tugagan")
            enrollment = active_enrollment(
                cur, assessment["context_id"], assessment["course_id"], user_id
            )
            if idempotency_key:
                existing = claim_request(
                    cur, context_id=assessment["context_id"], actor=user_id,
                    key=idempotency_key, action="attempt.start",
                )
                if existing:
                    attempt, assessment_view = read_attempt_for_user(
                        cur, existing, assessment["context_id"], user_id
                    )
                    if attempt["status"] == "expired":
                        return JSONResponse(
                            status_code=409,
                            content={
                                "detail": "Oldingi urinish muddati tugagan",
                                "code": "attempt_expired",
                                "attempt_id": int(attempt["id"]),
                            },
                        )
                    return {
                        "attempt": attempt,
                        "assessment": assessment_view,
                        "items": public_assessment_items(
                            cur, assessment["context_id"], assessment_id
                        ),
                        "idempotent_replay": True,
                    }
            if in_progress:
                if idempotency_key:
                    finish_request(
                        cur, assessment["context_id"], idempotency_key,
                        in_progress["id"],
                    )
                return {
                    "attempt": in_progress,
                    "assessment": {
                        "id": assessment_id,
                        "course_id": assessment["course_id"],
                        "title": assessment["title"],
                        "framework": assessment["framework"],
                        "duration_minutes": assessment["duration_minutes"],
                        "total_points": assessment["total_points"],
                        "closes_at": assessment["closes_at"],
                    },
                    "items": public_assessment_items(
                        cur, assessment["context_id"], assessment_id
                    ),
                    "resumed": True,
                    "answer_key_included": False,
                }
            cur.execute(
                """SELECT COUNT(*) count,COALESCE(MAX(attempt_no),0)+1 next_no
                   FROM center_assessment_attempts
                   WHERE context_id=%s AND assessment_id=%s
                     AND student_user_id=%s""",
                (assessment["context_id"], assessment_id, user_id),
            )
            count = cur.fetchone()
            if int(count["count"]) >= int(assessment["max_attempts"]):
                raise HTTPException(status_code=409, detail="Urinishlar soni tugagan")
            cur.execute(
                """INSERT INTO center_assessment_attempts(
                     context_id,assessment_id,course_id,enrollment_id,
                     student_user_id,attempt_no,status,max_score
                   ) VALUES(%s,%s,%s,%s,%s,%s,'in_progress',%s)
                   RETURNING *""",
                (
                    assessment["context_id"], assessment_id,
                    assessment["course_id"], enrollment["id"], user_id,
                    count["next_no"], assessment["total_points"],
                ),
            )
            attempt = cur.fetchone()
            if idempotency_key:
                finish_request(
                    cur, assessment["context_id"], idempotency_key, attempt["id"]
                )
            audit(
                cur, assessment["context_id"], user_id, "attempt.start",
                "assessment_attempt", attempt["id"],
            )
            items = public_assessment_items(
                cur, assessment["context_id"], assessment_id
            )
        return {
            "attempt": attempt,
            "assessment": {
                "id": assessment_id,
                "course_id": assessment["course_id"],
                "title": assessment["title"],
                "framework": assessment["framework"],
                "duration_minutes": assessment["duration_minutes"],
                "total_points": assessment["total_points"],
            },
            "items": items,
            "answer_key_included": False,
        }

    @router.get("/assessment-attempts")
    def list_assessment_attempts(
        context_id: int = Query(ge=1),
        course_id: int | None = Query(default=None, ge=1),
        assessment_id: int | None = Query(default=None, ge=1),
        status: Literal["submitted", "scored"] = Query(default="submitted"),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=100, ge=1, le=300),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            require_roles(
                cur, context_id, user_id, roles_for_permission("assessments.write")
            )
            if course_id is not None:
                course_access(cur, context_id, course_id, user_id, write=True)
            cur.execute(
                """SELECT at.id,at.assessment_id,a.title assessment_title,
                          at.course_id,c.name course_name,at.student_user_id,
                          u.full_name student_name,at.attempt_no,at.status,
                          at.started_at,at.submitted_at,at.score,at.max_score
                   FROM center_assessment_attempts at
                   JOIN center_assessments a ON a.id=at.assessment_id
                   JOIN center_courses c ON c.id=at.course_id
                   JOIN users u ON u.user_id=at.student_user_id
                   WHERE at.context_id=%s AND at.id>%s AND at.status=%s
                     AND (%s IS NULL OR at.course_id=%s)
                     AND (%s IS NULL OR at.assessment_id=%s)
                     AND (
                       (
                         c.teacher_user_id=%s AND EXISTS(
                           SELECT 1 FROM center_role_assignments teacher_role
                           WHERE teacher_role.context_id=at.context_id
                             AND teacher_role.user_id=%s
                             AND teacher_role.role_key='teacher'
                             AND teacher_role.status='active'
                             AND teacher_role.starts_at<=NOW()
                             AND (
                               teacher_role.ends_at IS NULL
                               OR teacher_role.ends_at>NOW()
                             )
                             AND (
                               teacher_role.branch_id IS NULL
                               OR teacher_role.branch_id IS NOT DISTINCT FROM c.branch_id
                             )
                         )
                       )
                       OR EXISTS(
                         SELECT 1 FROM center_role_assignments r
                         WHERE r.context_id=at.context_id AND r.user_id=%s
                           AND r.status='active' AND r.starts_at<=NOW()
                           AND (r.ends_at IS NULL OR r.ends_at>NOW())
                           AND r.role_key=ANY(%s)
                           AND (
                             r.branch_id IS NULL
                             OR r.branch_id IS NOT DISTINCT FROM c.branch_id
                           )
                       )
                     )
                   ORDER BY at.id LIMIT %s""",
                (
                    context_id, after_id or 0, status, course_id, course_id,
                    assessment_id, assessment_id, user_id, user_id, user_id,
                    list(roles_for_permission("assessments.write")), limit + 1,
                ),
            )
            rows = cur.fetchall()
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
        }

    @router.get("/my-assessment-attempts")
    def list_my_assessment_attempts(
        context_id: int = Query(ge=1),
        student_user_id: int | None = Query(default=None, ge=1),
        course_id: int | None = Query(default=None, ge=1),
        assessment_id: int | None = Query(default=None, ge=1),
        status: Literal[
            "in_progress", "submitted", "scored", "expired"
        ] | None = Query(default=None),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=100, ge=1, le=300),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        """Return the learner's or an actively linked child's safe summaries."""
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            require_roles(cur, context_id, user_id, VIEW_ROLES)
            target_student_id = student_user_id or user_id
            scope = "self"
            if target_student_id != user_id:
                cur.execute(
                    """SELECT 1 FROM center_parent_links
                       WHERE context_id=%s AND parent_user_id=%s
                         AND student_user_id=%s AND status='active'""",
                    (context_id, user_id, target_student_id),
                )
                if not cur.fetchone():
                    raise HTTPException(
                        status_code=403,
                        detail="Faqat faol bog'langan farzand natijasi ko'rinadi",
                    )
                scope = "linked_child"
            cur.execute(
                """SELECT at.id,at.assessment_id,a.title assessment_title,
                          at.course_id,c.name course_name,at.attempt_no,at.status,
                          at.started_at,at.submitted_at,at.scored_at,
                          at.score,at.max_score,a.duration_minutes,a.closes_at
                   FROM center_assessment_attempts at
                   JOIN center_assessments a ON a.id=at.assessment_id
                   JOIN center_courses c ON c.id=at.course_id
                   WHERE at.context_id=%s AND at.student_user_id=%s
                     AND at.id>%s
                     AND (%s IS NULL OR at.status=%s)
                     AND (%s IS NULL OR at.course_id=%s)
                     AND (%s IS NULL OR at.assessment_id=%s)
                     AND at.status IN ('in_progress','submitted','scored','expired')
                   ORDER BY at.id LIMIT %s""",
                (
                    context_id, target_student_id, after_id or 0, status, status,
                    course_id, course_id, assessment_id, assessment_id,
                    limit + 1,
                ),
            )
            rows = cur.fetchall()
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
            "answer_key_included": False,
            "scope": scope,
            "student_user_id": target_student_id,
        }

    @router.get("/attempts/{attempt_id}/review")
    def review_attempt(
        attempt_id: int,
        context_id: int = Query(ge=1),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            attempt, assessment = read_attempt_for_user(
                cur, attempt_id, context_id, user_id
            )
            course_access(
                cur, context_id, assessment["course_id"], user_id, write=True
            )
            if attempt["status"] not in {"submitted", "scored"}:
                raise HTTPException(
                    status_code=409,
                    detail=(
                        "Davom etayotgan yoki topshirilmagan urinish "
                        "xodim tekshiruviga ochilmaydi"
                    ),
                )
            cur.execute(
                """SELECT ai.id assessment_item_id,ai.item_order,ai.points,
                          ai.section_key,ai.question_source,ai.question_ref,
                          ai.metadata,aa.response,aa.answered_at,
                          gt.question,gt.option_a,gt.option_b,gt.option_c,
                          gt.option_d,gt.correct_answer generated_correct_answer,
                          gt.explanation generated_explanation,
                          gt.question_type generated_question_type
                   FROM center_assessment_items ai
                   LEFT JOIN center_attempt_answers aa
                     ON aa.context_id=ai.context_id
                    AND aa.assessment_item_id=ai.id
                    AND aa.attempt_id=%s
                   LEFT JOIN generated_tests gt
                     ON ai.question_source='generated_tests'
                    AND ai.question_ref~'^[0-9]+$'
                    AND gt.id=ai.question_ref::BIGINT
                   WHERE ai.context_id=%s AND ai.assessment_id=%s
                   ORDER BY ai.item_order,ai.id""",
                (attempt_id, context_id, attempt["assessment_id"]),
            )
            answers = cur.fetchall()
        return {
            "attempt": attempt,
            "assessment": assessment,
            "answers": answers,
            "answer_key_included": True,
            "grader_view": True,
        }

    @router.get("/attempts/{attempt_id}")
    def get_attempt(
        attempt_id: int,
        context_id: int = Query(ge=1),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database() as (_, cur):
            ensure_schema(cur)
            attempt, assessment = read_attempt_for_user(
                cur, attempt_id, context_id, user_id
            )
            if (
                attempt["status"] == "in_progress"
                and int(attempt["student_user_id"]) != user_id
            ):
                raise HTTPException(
                    status_code=403,
                    detail="Davom etayotgan test javobini faqat o'quvchi ko'radi",
                )
            expired = expire_attempt_if_late(
                cur, attempt, now=datetime.now(timezone.utc)
            )
            if expired:
                attempt = expired
                audit(
                    cur, context_id, user_id, "attempt.expire",
                    "assessment_attempt", attempt_id,
                    {"source": "attempt.get"},
                )
            items = public_assessment_items(
                cur, context_id, attempt["assessment_id"]
            )
            cur.execute(
                """SELECT assessment_item_id,response,answered_at
                   FROM center_attempt_answers
                   WHERE context_id=%s AND attempt_id=%s
                   ORDER BY assessment_item_id""",
                (context_id, attempt_id),
            )
            draft_answers = [dict(row) for row in cur.fetchall()]
        return {
            "attempt": attempt,
            "assessment": assessment,
            "items": items,
            "draft_answers": draft_answers,
            "answer_key_included": False,
        }

    @router.patch("/attempts/{attempt_id}/draft")
    def save_attempt_draft(
        attempt_id: int, request: AttemptDraft,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        """Persist partial learner answers without submitting the attempt."""
        require_encoded_size(
            request.answers, maximum=300_000, label="Urinish qoralama javoblari"
        )
        with database() as (_, cur):
            ensure_schema(cur)
            cur.execute(
                """SELECT context_id FROM center_assessment_attempts
                   WHERE id=%s FOR UPDATE""",
                (attempt_id,),
            )
            attempt_scope = cur.fetchone()
            if not attempt_scope:
                raise HTTPException(status_code=404, detail="Urinish topilmadi")
            context_id = int(attempt_scope["context_id"])
            require_operational(cur, context_id)
            if request.context_id is not None and request.context_id != context_id:
                raise HTTPException(status_code=404, detail="Urinish topilmadi")
            attempt, _ = read_attempt_for_user(
                cur, attempt_id, context_id, user_id
            )
            if int(attempt["student_user_id"]) != user_id:
                raise HTTPException(
                    status_code=403,
                    detail="Faqat o'z urinishingiz qoralamasini saqlaysiz",
                )
            if attempt["status"] != "in_progress":
                raise HTTPException(
                    status_code=409, detail="Urinish allaqachon yopilgan"
                )
            expired = expire_attempt_if_late(
                cur, attempt, now=datetime.now(timezone.utc)
            )
            if expired:
                audit(
                    cur, context_id, user_id, "attempt.expire",
                    "assessment_attempt", attempt_id,
                    {"source": "attempt.draft"},
                )
                return JSONResponse(
                    status_code=409,
                    content={
                        "detail": "Urinish saqlash muddati tugagan",
                        "code": "attempt_expired",
                        "attempt_id": attempt_id,
                    },
                )
            cur.execute(
                """SELECT id FROM center_assessment_items
                   WHERE context_id=%s AND assessment_id=%s""",
                (context_id, attempt["assessment_id"]),
            )
            allowed_items = {int(row["id"]) for row in cur.fetchall()}
            seen: set[int] = set()
            for answer in request.answers:
                try:
                    item_id = int(answer["assessment_item_id"])
                except (KeyError, TypeError, ValueError) as exc:
                    raise HTTPException(
                        status_code=422, detail="Javob item ID noto'g'ri"
                    ) from exc
                if item_id not in allowed_items or item_id in seen:
                    raise HTTPException(
                        status_code=422,
                        detail="Begona yoki takror savol javobi",
                    )
                response = answer.get("response")
                if not isinstance(response, dict):
                    raise HTTPException(
                        status_code=422, detail="response obyekt bo'lsin"
                    )
                seen.add(item_id)
                cur.execute(
                    """INSERT INTO center_attempt_answers(
                         context_id,attempt_id,assessment_id,
                         assessment_item_id,response
                       ) VALUES(%s,%s,%s,%s,%s::jsonb)
                       ON CONFLICT(
                         context_id,attempt_id,assessment_item_id
                       ) DO UPDATE SET
                         response=EXCLUDED.response,answered_at=NOW()""",
                    (
                        context_id, attempt_id, attempt["assessment_id"],
                        item_id, json.dumps(response, ensure_ascii=False),
                    ),
                )
            audit(
                cur, context_id, user_id, "attempt.draft.save",
                "assessment_attempt", attempt_id,
                {"answer_count": len(seen)},
            )
        return {
            "attempt_id": attempt_id,
            "saved_count": len(seen),
            "saved_at": datetime.now(timezone.utc),
            "submitted": False,
            "answer_key_included": False,
        }

    @router.post("/attempts/{attempt_id}/submit")
    def submit_attempt(
        attempt_id: int, request: AttemptSubmit,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_encoded_size(
            request.answers, maximum=300_000, label="Urinish javoblari"
        )
        with database() as (_, cur):
            ensure_schema(cur)
            cur.execute(
                """SELECT context_id FROM center_assessment_attempts
                   WHERE id=%s FOR UPDATE""",
                (attempt_id,),
            )
            attempt_scope = cur.fetchone()
            if not attempt_scope:
                raise HTTPException(status_code=404, detail="Urinish topilmadi")
            context_id = int(attempt_scope["context_id"])
            require_operational(cur, context_id)
            if request.context_id is not None and request.context_id != context_id:
                raise HTTPException(status_code=404, detail="Urinish topilmadi")
            existing = claim_request(
                cur, context_id=context_id, actor=user_id,
                key=request.idempotency_key, action="attempt.submit",
            )
            if existing:
                cur.execute(
                    """SELECT * FROM center_assessment_attempts
                       WHERE id=%s AND context_id=%s""",
                    (existing, context_id),
                )
                existing_attempt = cur.fetchone()
                if existing_attempt["status"] == "expired":
                    return JSONResponse(
                        status_code=409,
                        content={
                            "detail": "Urinish muddati tugagan",
                            "code": "attempt_expired",
                            "attempt_id": int(existing_attempt["id"]),
                        },
                    )
                return {
                    "attempt": existing_attempt,
                    "idempotent_replay": True,
                }
            attempt, _ = read_attempt_for_user(
                cur, attempt_id, context_id, user_id
            )
            if int(attempt["student_user_id"]) != user_id:
                raise HTTPException(status_code=403, detail="Faqat o'z urinishingizni topshirasiz")
            if attempt["status"] != "in_progress":
                raise HTTPException(status_code=409, detail="Urinish allaqachon yopilgan")
            now = datetime.now(timezone.utc)
            expired = expire_attempt_if_late(cur, attempt, now=now)
            if expired:
                finish_request(
                    cur, context_id, request.idempotency_key, attempt_id
                )
                audit(
                    cur, context_id, user_id, "attempt.expire",
                    "assessment_attempt", attempt_id,
                    {"source": "attempt.submit"},
                )
                return JSONResponse(
                    status_code=409,
                    content={
                        "detail": "Urinish topshirish muddati tugagan",
                        "code": "attempt_expired",
                        "attempt_id": attempt_id,
                    },
                )
            cur.execute(
                """SELECT ai.id,ai.question_source,ai.question_ref,ai.points,
                          ai.metadata,gt.correct_answer generated_correct_answer,
                          gt.question_type generated_question_type,
                          gt.option_a,gt.option_b,gt.option_c,gt.option_d
                   FROM center_assessment_items ai
                   LEFT JOIN generated_tests gt
                     ON ai.question_source='generated_tests'
                    AND ai.question_ref~'^[0-9]+$'
                    AND gt.id=ai.question_ref::BIGINT
                   WHERE ai.context_id=%s AND ai.assessment_id=%s
                   ORDER BY ai.item_order,ai.id""",
                (context_id, attempt["assessment_id"]),
            )
            raw_items = {int(row["id"]): dict(row) for row in cur.fetchall()}
            allowed_items = set(raw_items)
            seen: set[int] = set()
            responses: dict[int, dict[str, Any]] = {}
            for answer in request.answers:
                try:
                    item_id = int(answer["assessment_item_id"])
                except (KeyError, TypeError, ValueError) as exc:
                    raise HTTPException(status_code=422, detail="Javob item ID noto'g'ri") from exc
                if item_id not in allowed_items or item_id in seen:
                    raise HTTPException(status_code=422, detail="Begona yoki takror savol javobi")
                response = answer.get("response")
                if not isinstance(response, dict):
                    raise HTTPException(status_code=422, detail="response obyekt bo'lsin")
                seen.add(item_id)
                responses[item_id] = response
                cur.execute(
                    """INSERT INTO center_attempt_answers(
                         context_id,attempt_id,assessment_id,assessment_item_id,
                         response
                       ) VALUES(%s,%s,%s,%s,%s::jsonb)
                       ON CONFLICT(
                         context_id,attempt_id,assessment_item_id
                       ) DO UPDATE SET
                         response=EXCLUDED.response,answered_at=NOW()""",
                    (
                        context_id, attempt_id, attempt["assessment_id"],
                        item_id, json.dumps(response, ensure_ascii=False),
                    ),
                )
            if seen != allowed_items:
                raise HTTPException(
                    status_code=422,
                    detail="Barcha test savollariga bittadan javob bering.",
                )
            auto_score = Decimal("0")
            auto_max = Decimal("0")
            all_auto_scorable = True
            for item_id, item in raw_items.items():
                metadata = dict(item.get("metadata") or {})
                response = responses[item_id]
                if item["question_source"] == "manual_center":
                    correct = str(
                        metadata.get("correct_answer") or ""
                    ).strip()
                    if not correct:
                        all_auto_scorable = False
                        continue
                    if metadata.get("question_type") == "multiple_choice":
                        given = str(
                            response.get("selected")
                            if response.get("selected") is not None
                            else response.get("selected_index") or ""
                        ).strip()
                    else:
                        given = str(response.get("text") or "").strip()
                    normalized_given = " ".join(given.casefold().split())
                    normalized_correct = " ".join(
                        correct.casefold().split()
                    )
                    is_correct = normalized_given == normalized_correct
                elif (
                    item["question_source"] == "generated_tests"
                    and str(item.get("generated_question_type") or "")
                    != "write_answer"
                    and item.get("generated_correct_answer") is not None
                ):
                    options = [
                        str(item.get(key) or "").strip()
                        for key in ("option_a", "option_b", "option_c", "option_d")
                    ]

                    def choice_index(value: Any) -> int | None:
                        text_value = " ".join(
                            str(value or "").casefold().split()
                        )
                        if not text_value:
                            return None
                        letter = text_value.upper()
                        if letter in {"A", "B", "C", "D"}:
                            return ord(letter) - ord("A")
                        for option_index, option in enumerate(options):
                            if text_value == " ".join(option.casefold().split()):
                                return option_index
                        if text_value.isdigit():
                            numeric = int(text_value)
                            if 0 <= numeric <= 3:
                                return numeric
                            if 1 <= numeric <= 4:
                                return numeric - 1
                        return None

                    given_value = (
                        response.get("selected")
                        if response.get("selected") is not None
                        else response.get("selected_index")
                        if response.get("selected_index") is not None
                        else response.get("answer")
                    )
                    is_correct = choice_index(given_value) == choice_index(
                        item["generated_correct_answer"]
                    )
                else:
                    all_auto_scorable = False
                    continue
                points = Decimal(str(item["points"]))
                auto_max += points
                if is_correct:
                    auto_score += points
            final_status = "scored" if all_auto_scorable else "submitted"
            cur.execute(
                """UPDATE center_assessment_attempts SET status=%s,
                     submitted_at=NOW(),submission_key=%s,
                     score=CASE WHEN %s THEN %s ELSE NULL END,
                     max_score=CASE WHEN %s THEN %s ELSE max_score END,
                     scored_at=CASE WHEN %s THEN NOW() ELSE NULL END,
                     updated_at=NOW()
                   WHERE id=%s RETURNING *""",
                (
                    final_status, request.idempotency_key, all_auto_scorable,
                    auto_score, all_auto_scorable, auto_max,
                    all_auto_scorable, attempt_id,
                ),
            )
            submitted = cur.fetchone()
            finish_request(cur, context_id, request.idempotency_key, attempt_id)
            audit(
                cur, context_id, user_id, "attempt.submit",
                "assessment_attempt", attempt_id, {"answer_count": len(seen)},
            )
        return {
            "attempt": submitted,
            "auto_scored": all_auto_scorable,
            "message": (
                "Test avtomatik tekshirildi."
                if all_auto_scorable
                else "Javoblar topshirildi; yozma javoblar tekshirilgach natija ko'rinadi."
            ),
        }

    @router.post("/attempts/{attempt_id}/score")
    def score_attempt(
        attempt_id: int, request: AttemptScore,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(request.confirmation, "Imtihon bahosi uchun inson tasdig'i kerak.")
        require_encoded_size(
            request.component_scores, maximum=10_000,
            label="Komponent ballari",
        )
        if request.score > request.max_score:
            raise HTTPException(status_code=422, detail="Ball maksimal balldan oshdi")
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            existing = claim_request(
                cur, context_id=request.context_id, actor=user_id,
                key=request.idempotency_key, action="attempt.score",
            )
            if existing:
                cur.execute(
                    "SELECT * FROM center_assessment_attempts WHERE id=%s",
                    (existing,),
                )
                return {"attempt": cur.fetchone(), "idempotent_replay": True}
            attempt, assessment = read_attempt_for_user(
                cur, attempt_id, request.context_id, user_id
            )
            course_access(
                cur, request.context_id, assessment["course_id"], user_id, write=True
            )
            if Decimal(str(request.max_score)) != Decimal(
                str(assessment["total_points"])
            ):
                raise HTTPException(
                    status_code=422,
                    detail="Maksimal ball e'lon qilingan imtihon jami balliga teng bo'lsin",
                )
            if attempt["status"] not in {"submitted", "scored"}:
                raise HTTPException(status_code=409, detail="Topshirilmagan urinish baholanmaydi")
            if assessment["framework"] == "ielts":
                if set(request.component_scores) != set(IELTS_COMPONENTS):
                    raise HTTPException(
                        status_code=422,
                        detail="IELTS uchun listening, reading, writing va speaking ballari kerak",
                    )
                for value in request.component_scores.values():
                    if (
                        value < 0 or value > 9
                        or value * 2 != (value * 2).to_integral_value()
                    ):
                        raise HTTPException(status_code=422, detail="IELTS balli 0–9, 0.5 qadam")
            cur.execute(
                """UPDATE center_assessment_attempts SET status='scored',
                     score=%s,max_score=%s,component_scores=%s::jsonb,
                     scored_at=NOW(),scored_by_user_id=%s,updated_at=NOW()
                   WHERE id=%s RETURNING *""",
                (
                    request.score, request.max_score,
                    json.dumps(request.component_scores, default=str),
                    user_id, attempt_id,
                ),
            )
            scored = cur.fetchone()
            finish_request(cur, request.context_id, request.idempotency_key, attempt_id)
            audit(
                cur, request.context_id, user_id, "attempt.score",
                "assessment_attempt", attempt_id,
            )
        return {"attempt": scored}

    @router.get("/attempts/{attempt_id}/result")
    def attempt_result(
        attempt_id: int,
        context_id: int = Query(ge=1),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            attempt, assessment = read_attempt_for_user(
                cur, attempt_id, context_id, user_id,
                allow_linked_parent=True,
            )
            if attempt["status"] == "in_progress":
                raise HTTPException(status_code=409, detail="Urinish hali topshirilmagan")
            cur.execute(
                """SELECT COUNT(*) answer_count
                   FROM center_attempt_answers
                   WHERE context_id=%s AND attempt_id=%s""",
                (context_id, attempt_id),
            )
            answer_count = cur.fetchone()["answer_count"]
        return {
            "attempt_id": attempt_id,
            "assessment": assessment,
            "status": attempt["status"],
            "score": attempt["score"],
            "max_score": attempt["max_score"],
            "component_scores": attempt["component_scores"],
            "answer_count": answer_count,
            "answer_key_included": False,
        }

    @router.get("/billing/plans")
    def list_billing_plans(
        context_id: int = Query(ge=1),
        course_id: int | None = Query(default=None, ge=1),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=50, ge=1, le=200),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            roles = require_roles(cur, context_id, user_id, FINANCE_VIEW_ROLES)
            cur.execute(
                """SELECT p.*,c.name course_name
                   FROM center_billing_plans p
                   LEFT JOIN center_courses c ON c.id=p.course_id
                   WHERE p.context_id=%s AND p.id>%s
                     AND (%s IS NULL OR p.course_id=%s)
                     AND (
                       %s OR EXISTS(
                         SELECT 1 FROM center_role_assignments r
                         WHERE r.context_id=p.context_id AND r.user_id=%s
                           AND r.status='active' AND r.starts_at<=NOW()
                           AND (r.ends_at IS NULL OR r.ends_at>NOW())
                           AND r.role_key=ANY(%s)
                           AND (
                             r.branch_id IS NULL
                             OR (c.branch_id IS NOT NULL AND r.branch_id=c.branch_id)
                           )
                         )
                       OR (
                         p.course_id IS NOT NULL AND (
                           EXISTS(
                             SELECT 1 FROM center_enrollments e
                             WHERE e.context_id=p.context_id
                               AND e.course_id=p.course_id
                               AND e.student_user_id=%s
                               AND e.status IN ('active','completed')
                               AND (
                                 e.starts_on IS NULL
                                 OR e.starts_on<=
                                   (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                               )
                           )
                           OR EXISTS(
                             SELECT 1 FROM center_parent_links pl
                             JOIN center_enrollments e
                               ON e.context_id=pl.context_id
                              AND e.student_user_id=pl.student_user_id
                             WHERE pl.context_id=p.context_id
                               AND pl.parent_user_id=%s
                               AND pl.status='active'
                               AND e.course_id=p.course_id
                               AND e.status IN ('active','completed')
                               AND (
                                 e.starts_on IS NULL
                                 OR e.starts_on<=
                                   (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                               )
                           )
                         )
                       )
                     )
                   ORDER BY p.id LIMIT %s""",
                (
                    context_id, after_id or 0, course_id, course_id,
                    "system_admin" in roles, user_id,
                    list(FINANCE_VIEW_STAFF_ROLES), user_id, user_id,
                    limit + 1,
                ),
            )
            rows = cur.fetchall()
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
        }

    @router.post("/billing/plans")
    def create_billing_plan(
        request: BillingPlanCreate,
        idempotency_key: str | None = Header(
            default=None, alias="Idempotency-Key"
        ),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        idempotency_key = validate_request_key(idempotency_key)
        require_human(request.confirmation, "To'lov rejasi uchun inson tasdig'i kerak.")
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            branch_id = None
            if request.course_id is not None:
                course = course_resource(
                    cur, request.context_id, request.course_id
                )
                branch_id = course["branch_id"]
            require_permission(
                cur, request.context_id, user_id, "billing.manage",
                branch_id=branch_id, require_global=request.course_id is None,
            )
            if idempotency_key:
                existing_target = claim_request(
                    cur, context_id=request.context_id, actor=user_id,
                    key=idempotency_key, action="billing_plan.create",
                )
                if existing_target:
                    cur.execute(
                        """SELECT * FROM center_billing_plans
                           WHERE id=%s AND context_id=%s""",
                        (existing_target, request.context_id),
                    )
                    return {
                        "item": cur.fetchone(),
                        "idempotent_replay": True,
                    }
            cur.execute(
                """INSERT INTO center_billing_plans(
                     context_id,course_id,name,amount,billing_cycle,currency,
                     billing_day,created_by_user_id
                   ) VALUES(%s,%s,%s,%s,%s,%s,%s,%s)
                   RETURNING *""",
                (
                    request.context_id, request.course_id, request.name.strip(),
                    request.amount, request.billing_cycle, request.currency.upper(),
                    request.billing_day, user_id,
                ),
            )
            item = cur.fetchone()
            if idempotency_key:
                finish_request(
                    cur, request.context_id, idempotency_key, item["id"]
                )
            audit(
                cur, request.context_id, user_id, "billing_plan.create",
                "billing_plan", item["id"],
            )
        return {"item": item}

    @router.post("/billing/discounts")
    def create_discount(
        request: DiscountCreate, user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(request.confirmation, "Chegirma uchun inson tasdig'i kerak.")
        if request.discount_type == "percent" and request.value > 100:
            raise HTTPException(status_code=422, detail="Foizli chegirma 100 dan oshmaydi")
        if request.starts_on and request.ends_on and request.ends_on < request.starts_on:
            raise HTTPException(status_code=422, detail="Chegirma sanalari noto'g'ri")
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            branch_id = None
            if request.course_id:
                course = course_resource(
                    cur, request.context_id, request.course_id
                )
                branch_id = course["branch_id"]
            require_permission(
                cur, request.context_id, user_id, "billing.manage",
                branch_id=branch_id, require_global=request.course_id is None,
            )
            cur.execute(
                """INSERT INTO center_discounts(
                     context_id,name,discount_type,value,student_user_id,
                     course_id,starts_on,ends_on,created_by_user_id
                   ) VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s)
                   RETURNING *""",
                (
                    request.context_id, request.name.strip(),
                    request.discount_type, request.value, request.student_user_id,
                    request.course_id, request.starts_on, request.ends_on, user_id,
                ),
            )
            item = cur.fetchone()
            audit(
                cur, request.context_id, user_id, "discount.create",
                "discount", item["id"],
            )
        return {"item": item}

    @router.get("/billing/invoices")
    def list_invoices(
        context_id: int = Query(ge=1),
        status: str | None = Query(default=None, max_length=30),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=100, ge=1, le=300),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            roles = require_roles(cur, context_id, user_id, VIEW_ROLES)
            cur.execute(
                """SELECT i.*,i.amount-i.paid_amount balance_due,
                          u.full_name student_name,c.name course_name,
                          p.name plan_name
                   FROM center_invoices i
                   JOIN users u ON u.user_id=i.student_user_id
                   JOIN center_courses c ON c.id=i.course_id
                   JOIN center_billing_plans p ON p.id=i.plan_id
                   WHERE i.context_id=%s AND i.id>%s
                     AND (%s IS NULL OR i.status=%s)
                     AND (
                       %s OR i.student_user_id=%s
                       OR EXISTS(
                         SELECT 1 FROM center_parent_links pl
                         WHERE pl.context_id=i.context_id
                           AND pl.parent_user_id=%s
                           AND pl.student_user_id=i.student_user_id
                           AND pl.status='active'
                       )
                       OR EXISTS(
                         SELECT 1 FROM center_role_assignments r
                         WHERE r.context_id=i.context_id AND r.user_id=%s
                           AND r.status='active' AND r.starts_at<=NOW()
                           AND (r.ends_at IS NULL OR r.ends_at>NOW())
                           AND r.role_key=ANY(%s)
                           AND (
                             r.branch_id IS NULL
                             OR (c.branch_id IS NOT NULL AND r.branch_id=c.branch_id)
                           )
                       )
                     )
                   ORDER BY i.id LIMIT %s""",
                (
                    context_id, after_id or 0, status, status,
                    "system_admin" in roles, user_id, user_id, user_id,
                    list(FINANCE_VIEW_STAFF_ROLES), limit + 1,
                ),
            )
            rows = cur.fetchall()
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
        }

    @router.post("/billing/invoices")
    def create_invoice(
        request: InvoiceCreate,
        idempotency_key: str | None = Header(
            default=None, alias="Idempotency-Key"
        ),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        idempotency_key = validate_request_key(idempotency_key)
        require_human(request.confirmation, "Hisob yaratish uchun inson tasdig'i kerak.")
        if request.period_end < request.period_start:
            raise HTTPException(status_code=422, detail="Hisob davri noto'g'ri")
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            cur.execute(
                """SELECT p.*,
                          p.course_id plan_course_id,
                          e.course_id enrollment_course_id,
                          e.student_user_id,c.branch_id
                   FROM center_billing_plans p
                   JOIN center_enrollments e
                     ON e.id=%s AND e.context_id=p.context_id
                   JOIN center_courses c ON c.id=e.course_id
                   WHERE p.id=%s AND p.context_id=%s AND p.active=TRUE
                     AND e.status IN ('active','paused')""",
                (request.enrollment_id, request.plan_id, request.context_id),
            )
            source = cur.fetchone()
            if not source:
                raise HTTPException(status_code=404, detail="Reja yoki qabul topilmadi")
            if (
                source["plan_course_id"] is not None
                and source["plan_course_id"] != source["enrollment_course_id"]
            ):
                raise HTTPException(status_code=409, detail="To'lov rejasi boshqa kursniki")
            require_permission(
                cur, request.context_id, user_id, "billing.manage",
                branch_id=source["branch_id"],
            )
            if idempotency_key:
                existing_target = claim_request(
                    cur, context_id=request.context_id, actor=user_id,
                    key=idempotency_key, action="invoice.create",
                )
                if existing_target:
                    cur.execute(
                        """SELECT * FROM center_invoices
                           WHERE id=%s AND context_id=%s""",
                        (existing_target, request.context_id),
                    )
                    return {
                        "item": cur.fetchone(),
                        "idempotent_replay": True,
                    }
            if request.discount_amount > source["amount"]:
                raise HTTPException(status_code=422, detail="Chegirma hisobdan katta")
            final_amount = source["amount"] - request.discount_amount
            cur.execute(
                """INSERT INTO center_invoices(
                     context_id,plan_id,enrollment_id,course_id,student_user_id,
                     period_start,period_end,due_date,subtotal,discount_amount,
                     amount,created_by_user_id
                   ) VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                   ON CONFLICT(
                     context_id,plan_id,enrollment_id,period_start,period_end
                   ) DO NOTHING
                   RETURNING *""",
                (
                    request.context_id, request.plan_id, request.enrollment_id,
                    source["enrollment_course_id"], source["student_user_id"],
                    request.period_start, request.period_end, request.due_date,
                    source["amount"], request.discount_amount, final_amount, user_id,
                ),
            )
            item = cur.fetchone()
            if not item:
                raise HTTPException(status_code=409, detail="Bu davr hisobi mavjud")
            if idempotency_key:
                finish_request(
                    cur, request.context_id, idempotency_key, item["id"]
                )
            audit(
                cur, request.context_id, user_id, "invoice.create",
                "invoice", item["id"],
            )
        return {"item": item}

    @router.post("/billing/payments")
    def create_payment(
        request: PaymentCreate, user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(request.confirmation, "To'lov yozuvi uchun inson tasdig'i kerak.")
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            existing = claim_request(
                cur, context_id=request.context_id, actor=user_id,
                key=request.idempotency_key, action="payment.create",
            )
            if existing:
                cur.execute("SELECT * FROM center_payments WHERE id=%s", (existing,))
                return {"item": cur.fetchone(), "idempotent_replay": True}
            cur.execute(
                """SELECT i.*,c.branch_id FROM center_invoices i
                   JOIN center_courses c ON c.id=i.course_id
                   WHERE i.id=%s AND i.context_id=%s FOR UPDATE""",
                (request.invoice_id, request.context_id),
            )
            invoice = cur.fetchone()
            if not invoice or invoice["status"] in {"cancelled", "written_off"}:
                raise HTTPException(status_code=404, detail="Faol hisob topilmadi")
            require_permission(
                cur, request.context_id, user_id, "billing.manage",
                branch_id=invoice["branch_id"],
            )
            remaining = invoice["amount"] - invoice["paid_amount"]
            if request.amount > remaining:
                raise HTTPException(status_code=422, detail="To'lov qoldiqdan oshdi")
            if isinstance(request.paid_at, date) and not isinstance(request.paid_at, datetime):
                paid_at = datetime.combine(
                    request.paid_at, time(12), tzinfo=BUSINESS_TIMEZONE
                )
            else:
                paid_at = request.paid_at or datetime.now(timezone.utc)
            cur.execute(
                """INSERT INTO center_payments(
                     context_id,invoice_id,amount,payment_method,idempotency_key,
                     reference,received_by_user_id,paid_at
                   ) VALUES(%s,%s,%s,%s,%s,%s,%s,%s)
                   RETURNING *""",
                (
                    request.context_id, request.invoice_id, request.amount,
                    request.payment_method, request.idempotency_key,
                    request.reference, user_id, paid_at,
                ),
            )
            payment = cur.fetchone()
            new_paid = invoice["paid_amount"] + request.amount
            new_status = "paid" if new_paid == invoice["amount"] else "partial"
            cur.execute(
                """UPDATE center_invoices SET paid_amount=%s,status=%s,updated_at=NOW()
                   WHERE id=%s""",
                (new_paid, new_status, request.invoice_id),
            )
            finish_request(
                cur, request.context_id, request.idempotency_key, payment["id"]
            )
            audit(
                cur, request.context_id, user_id, "payment.create",
                "payment", payment["id"],
                {"invoice_id": request.invoice_id, "amount": str(request.amount)},
            )
        return {"item": payment, "invoice_status": new_status}

    @router.get("/billing/debts")
    def list_debts(
        context_id: int = Query(ge=1),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=100, ge=1, le=300),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        result = list_invoices(
            context_id=context_id, status=None, after_id=after_id, limit=limit,
            user_id=user_id,
        )
        result["items"] = [
            {
                **dict(item),
                "debt_amount": item["amount"] - item["paid_amount"],
                "overdue": item["due_date"] < business_today(),
            }
            for item in result["items"]
            if item["status"] in {"unpaid", "partial"}
        ]
        return result

    def analytics_scope(
        cur: Any, context_id: int, user_id: int,
    ) -> tuple[list[int] | None, list[int] | None]:
        """Return allowed course IDs and, for learners, allowed student IDs.

        ``None`` for course IDs means a global staff assignment.  Learners and
        parents always receive a concrete student filter so aggregate endpoints
        cannot disclose another learner's results.
        """

        if system_admin(cur, user_id):
            return None, None
        roles, assignments = active_roles(cur, context_id, user_id)
        broad_roles = MANAGER_ROLES | {
            "receptionist", "accountant", "methodist",
        }
        broad_assignments = [
            row for row in assignments if row["role_key"] in broad_roles
        ]
        if any(row["branch_id"] is None for row in broad_assignments):
            return None, None
        broad_branches = sorted(
            {
                int(row["branch_id"])
                for row in broad_assignments
                if row["branch_id"] is not None
            }
        )
        teacher_assignments = [
            row for row in assignments if row["role_key"] == "teacher"
        ]
        teacher_global = any(
            row["branch_id"] is None for row in teacher_assignments
        )
        teacher_branches = sorted(
            {
                int(row["branch_id"])
                for row in teacher_assignments
                if row["branch_id"] is not None
            }
        )
        staff_user = bool(roles & STAFF_ROLES)
        student_ids: list[int] | None = None
        if not staff_user:
            student_set: set[int] = set()
            if "student" in roles:
                student_set.add(user_id)
            if "parent" in roles:
                cur.execute(
                    """SELECT student_user_id FROM center_parent_links
                       WHERE context_id=%s AND parent_user_id=%s
                         AND status='active'""",
                    (context_id, user_id),
                )
                student_set.update(int(row["student_user_id"]) for row in cur.fetchall())
            student_ids = sorted(student_set)
        cur.execute(
            """SELECT DISTINCT c.id
               FROM center_courses c
               WHERE c.context_id=%s AND (
                 (
                   %s AND c.teacher_user_id=%s
                   AND (
                     %s
                     OR (
                       %s::BIGINT[] IS NOT NULL
                       AND c.branch_id=ANY(%s)
                     )
                   )
                 )
                 OR (%s::BIGINT[] IS NOT NULL AND c.branch_id=ANY(%s))
                 OR (%s AND EXISTS(
                   SELECT 1 FROM center_enrollments e
                   WHERE e.context_id=c.context_id AND e.course_id=c.id
                     AND e.student_user_id=%s
                     AND e.status IN ('active','completed','paused')
                     AND (
                       e.starts_on IS NULL
                       OR e.starts_on<=
                         (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                     )
                 ))
                 OR (%s AND EXISTS(
                   SELECT 1 FROM center_parent_links p
                   JOIN center_enrollments e
                     ON e.context_id=p.context_id
                    AND e.student_user_id=p.student_user_id
                   WHERE p.context_id=c.context_id AND p.parent_user_id=%s
                     AND p.status='active' AND e.course_id=c.id
                     AND e.status IN ('active','completed','paused')
                     AND (
                       e.starts_on IS NULL
                       OR e.starts_on<=
                         (CURRENT_TIMESTAMP AT TIME ZONE 'Asia/Tashkent')::DATE
                     )
                 ))
               )
               ORDER BY c.id""",
            (
                context_id, "teacher" in roles, user_id,
                teacher_global,
                teacher_branches or None, teacher_branches or None,
                broad_branches or None, broad_branches or None,
                not staff_user, user_id, not staff_user, user_id,
            ),
        )
        return [int(row["id"]) for row in cur.fetchall()], student_ids

    @router.get("/teacher-workload")
    def teacher_workload(
        context_id: int = Query(ge=1),
        teacher_user_id: int | None = Query(default=None, ge=1),
        from_date: date | None = Query(default=None),
        to_date: date | None = Query(default=None),
        after_id: int | None = Query(default=None, ge=1),
        limit: int = Query(default=100, ge=1, le=300),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        period_end = to_date or business_today()
        period_start = from_date or period_end.replace(day=1)
        if period_end < period_start or (period_end - period_start).days > 366:
            raise HTTPException(status_code=422, detail="Yuklama davri 367 kundan oshmasin")
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            roles = require_roles(
                cur, context_id, user_id,
                MANAGER_ROLES | {"accountant", "methodist", "teacher"},
            )
            course_ids, _ = analytics_scope(cur, context_id, user_id)
            requested_teacher = teacher_user_id
            can_view_other_teachers = (
                "system_admin" in roles
                or bool(roles & (MANAGER_ROLES | {"accountant", "methodist"}))
            )
            if (
                requested_teacher is not None
                and requested_teacher != user_id
                and not can_view_other_teachers
            ):
                raise HTTPException(status_code=403, detail="Boshqa o'qituvchi yuklamasi yopiq")
            if requested_teacher is None and not can_view_other_teachers:
                requested_teacher = user_id
            cur.execute(
                """SELECT w.*,u.full_name teacher_name,c.name course_name,
                          b.name branch_name
                   FROM center_teacher_work_logs w
                   JOIN users u ON u.user_id=w.teacher_user_id
                   JOIN center_courses c ON c.id=w.course_id
                   LEFT JOIN center_branches b ON b.id=c.branch_id
                   WHERE w.context_id=%s AND w.id>%s
                     AND w.work_date BETWEEN %s AND %s
                     AND (%s::BIGINT[] IS NULL OR w.course_id=ANY(%s))
                     AND (%s IS NULL OR w.teacher_user_id=%s)
                   ORDER BY w.id LIMIT %s""",
                (
                    context_id, after_id or 0, period_start, period_end,
                    course_ids, course_ids, requested_teacher, requested_teacher,
                    limit + 1,
                ),
            )
            rows = [dict(row) for row in cur.fetchall()]
            cur.execute(
                """SELECT COUNT(*) entry_count,
                          COALESCE(SUM(minutes_worked),0) minutes_total,
                          COALESCE(SUM(minutes_worked)
                            FILTER(WHERE status='approved'),0) approved_minutes,
                          COALESCE(SUM(amount)
                            FILTER(WHERE status IN ('approved','paid')),0) approved_amount
                   FROM center_teacher_work_logs
                   WHERE context_id=%s AND work_date BETWEEN %s AND %s
                     AND (%s::BIGINT[] IS NULL OR course_id=ANY(%s))
                     AND (%s IS NULL OR teacher_user_id=%s)""",
                (
                    context_id, period_start, period_end, course_ids, course_ids,
                    requested_teacher, requested_teacher,
                ),
            )
            summary = dict(cur.fetchone())
            can_view_pay = bool(
                "system_admin" in roles
                or roles & FINANCE_VIEW_STAFF_ROLES
                or (
                    requested_teacher == user_id
                    and "teacher" in roles
                )
            )
            if not can_view_pay:
                for row in rows:
                    row["rate"] = None
                    row["amount"] = None
                summary["approved_amount"] = None
        return {
            "items": rows[:limit],
            "next_cursor": rows[limit - 1]["id"] if len(rows) > limit else None,
            "period": {"from": period_start, "to": period_end},
            "summary": summary,
        }

    @router.post("/teacher-worklogs")
    def create_teacher_worklog(
        request: WorkLogCreate,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            course = course_access(
                cur, request.context_id, request.course_id, user_id, write=True
            )
            if request.teacher_user_id != user_id:
                require_roles(
                    cur, request.context_id, user_id, MANAGER_ROLES | {"methodist"},
                    branch_id=course["branch_id"],
                )
            if (
                course["teacher_user_id"] is not None
                and int(course["teacher_user_id"]) != request.teacher_user_id
            ):
                raise HTTPException(
                    status_code=409,
                    detail="Ish yozuvi kursga biriktirilgan o'qituvchiga mos emas",
                )
            if request.schedule_slot_id is not None:
                cur.execute(
                    """SELECT 1 FROM center_schedule_slots
                       WHERE id=%s AND context_id=%s AND course_id=%s
                         AND teacher_user_id=%s""",
                    (
                        request.schedule_slot_id, request.context_id,
                        request.course_id, request.teacher_user_id,
                    ),
                )
                if not cur.fetchone():
                    raise HTTPException(status_code=404, detail="Mos dars jadvali topilmadi")
            if request.idempotency_key:
                existing = claim_request(
                    cur, context_id=request.context_id, actor=user_id,
                    key=request.idempotency_key, action="worklog.create",
                )
                if existing:
                    cur.execute(
                        "SELECT * FROM center_teacher_work_logs WHERE id=%s",
                        (existing,),
                    )
                    return {"item": cur.fetchone(), "idempotent_replay": True}
            amount: Decimal | None = None
            if request.rate is not None:
                amount = (
                    request.rate * Decimal(request.minutes_worked) / Decimal(60)
                    if request.pay_unit == "hour"
                    else request.rate
                ).quantize(Decimal("0.01"))
            cur.execute(
                """INSERT INTO center_teacher_work_logs(
                     context_id,teacher_user_id,course_id,schedule_slot_id,
                     work_date,minutes_worked,pay_unit,rate,amount,note,
                     idempotency_key
                   ) VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                   RETURNING *""",
                (
                    request.context_id, request.teacher_user_id,
                    request.course_id, request.schedule_slot_id,
                    request.work_date, request.minutes_worked, request.pay_unit,
                    request.rate, amount, request.note, request.idempotency_key,
                ),
            )
            item = cur.fetchone()
            if request.idempotency_key:
                finish_request(
                    cur, request.context_id, request.idempotency_key, item["id"]
                )
            audit(
                cur, request.context_id, user_id, "worklog.create",
                "teacher_work_log", item["id"],
            )
        return {"item": item}

    @router.post("/teacher-worklogs/{worklog_id}/decision")
    def decide_teacher_worklog(
        worklog_id: int, request: WorkLogDecision,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        require_human(request.confirmation, "Ish yuklamasi qarori uchun tasdiq kerak.")
        with database() as (_, cur):
            ensure_schema(cur)
            require_operational(cur, request.context_id)
            cur.execute(
                """SELECT w.*,c.branch_id FROM center_teacher_work_logs w
                   JOIN center_courses c ON c.id=w.course_id
                   WHERE w.id=%s AND w.context_id=%s FOR UPDATE""",
                (worklog_id, request.context_id),
            )
            item = cur.fetchone()
            if not item:
                raise HTTPException(status_code=404, detail="Ish yozuvi topilmadi")
            if int(item["teacher_user_id"]) == user_id:
                raise HTTPException(
                    status_code=403,
                    detail="O'qituvchi o'z ish yozuvini o'zi tasdiqlay olmaydi",
                )
            require_roles(
                cur, request.context_id, user_id,
                MANAGER_ROLES | {"accountant"},
                branch_id=item["branch_id"],
                require_global=item["branch_id"] is None,
            )
            cur.execute(
                """UPDATE center_teacher_work_logs SET status=%s,
                     approved_by_user_id=CASE WHEN %s='approved' THEN %s ELSE NULL END,
                     approved_at=CASE WHEN %s='approved' THEN NOW() ELSE NULL END,
                     updated_at=NOW()
                   WHERE id=%s RETURNING *""",
                (
                    request.status, request.status, user_id,
                    request.status, worklog_id,
                ),
            )
            updated = cur.fetchone()
            audit(
                cur, request.context_id, user_id, f"worklog.{request.status}",
                "teacher_work_log", worklog_id,
            )
        return {"item": updated}

    @router.get("/analytics/summary")
    def analytics_summary(
        context_id: int = Query(ge=1),
        period: Literal["7d", "30d", "90d", "year"] = Query(default="30d"),
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        today = business_today()
        if period == "year":
            start = date(today.year if today.month >= 9 else today.year - 1, 9, 1)
        else:
            days = {"7d": 7, "30d": 30, "90d": 90}[period]
            start = today.fromordinal(today.toordinal() - days + 1)
        with database(readonly=True) as (_, cur):
            ensure_schema(cur)
            roles = require_roles(cur, context_id, user_id, VIEW_ROLES)
            course_ids, student_ids = analytics_scope(cur, context_id, user_id)
            cur.execute(
                """SELECT
                     (SELECT COUNT(DISTINCT e.student_user_id)
                      FROM center_enrollments e
                      WHERE e.context_id=%s AND e.status='active'
                        AND (%s::BIGINT[] IS NULL OR e.course_id=ANY(%s))
                        AND (%s::BIGINT[] IS NULL OR e.student_user_id=ANY(%s))
                     ) students_active,
                     (SELECT COALESCE(ROUND(
                        100.0*COUNT(*) FILTER(WHERE a.status IN ('present','late'))
                        /NULLIF(COUNT(*),0),1),0)
                      FROM center_attendance a
                      WHERE a.context_id=%s AND a.lesson_date BETWEEN %s AND %s
                        AND (%s::BIGINT[] IS NULL OR a.course_id=ANY(%s))
                        AND (%s::BIGINT[] IS NULL OR a.student_user_id=ANY(%s))
                     ) attendance_rate,
                     (SELECT COALESCE(ROUND(
                        100.0*AVG(g.score/NULLIF(g.max_score,0)),1),0)
                      FROM center_grade_entries g
                      WHERE g.context_id=%s AND g.graded_at::DATE BETWEEN %s AND %s
                        AND (%s::BIGINT[] IS NULL OR g.course_id=ANY(%s))
                        AND (%s::BIGINT[] IS NULL OR g.student_user_id=ANY(%s))
                     ) average_score,
                     (SELECT COUNT(*)
                      FROM center_assessment_attempts at
                      WHERE at.context_id=%s
                        AND at.status IN ('submitted','scored')
                        AND at.submitted_at::DATE BETWEEN %s AND %s
                        AND (%s::BIGINT[] IS NULL OR at.course_id=ANY(%s))
                        AND (%s::BIGINT[] IS NULL OR at.student_user_id=ANY(%s))
                     ) assessments_completed,
                     (SELECT COUNT(*)
                      FROM center_enrollments e
                      WHERE e.context_id=%s AND e.created_at::DATE BETWEEN %s AND %s
                        AND (%s::BIGINT[] IS NULL OR e.course_id=ANY(%s))
                        AND (%s::BIGINT[] IS NULL OR e.student_user_id=ANY(%s))
                     ) new_enrollments,
                     (SELECT COALESCE(SUM(i.amount-i.paid_amount),0)
                      FROM center_invoices i
                      WHERE i.context_id=%s AND i.status IN ('unpaid','partial')
                        AND (%s::BIGINT[] IS NULL OR i.course_id=ANY(%s))
                        AND (%s::BIGINT[] IS NULL OR i.student_user_id=ANY(%s))
                     ) debt_total""",
                (
                    context_id, course_ids, course_ids, student_ids, student_ids,
                    context_id, start, today, course_ids, course_ids,
                    student_ids, student_ids,
                    context_id, start, today, course_ids, course_ids,
                    student_ids, student_ids,
                    context_id, start, today, course_ids, course_ids,
                    student_ids, student_ids,
                    context_id, start, today, course_ids, course_ids,
                    student_ids, student_ids,
                    context_id, course_ids, course_ids, student_ids, student_ids,
                ),
            )
            summary = dict(cur.fetchone())
            can_view_staff_debt = bool(
                "system_admin" in roles or roles & FINANCE_VIEW_STAFF_ROLES
            )
            is_pure_learner_view = bool(
                roles & {"student", "parent"} and not roles & STAFF_ROLES
            )
            if not (can_view_staff_debt or is_pure_learner_view):
                summary["debt_total"] = Decimal("0")
            cur.execute(
                """SELECT c.id course_id,c.name course_name,
                          COUNT(DISTINCT e.student_user_id)
                            FILTER(WHERE e.status='active') students,
                          COALESCE(ROUND(
                            100.0*AVG(g.score/NULLIF(g.max_score,0)),1
                          ),0) average_score
                   FROM center_courses c
                   LEFT JOIN center_enrollments e
                     ON e.context_id=c.context_id AND e.course_id=c.id
                    AND (%s::BIGINT[] IS NULL OR e.student_user_id=ANY(%s))
                   LEFT JOIN center_grade_entries g
                     ON g.context_id=c.context_id AND g.course_id=c.id
                    AND g.graded_at::DATE BETWEEN %s AND %s
                    AND (%s::BIGINT[] IS NULL OR g.student_user_id=ANY(%s))
                   WHERE c.context_id=%s
                     AND (%s::BIGINT[] IS NULL OR c.id=ANY(%s))
                   GROUP BY c.id,c.name
                   ORDER BY c.id LIMIT 20""",
                (
                    student_ids, student_ids, start, today,
                    student_ids, student_ids, context_id, course_ids, course_ids,
                ),
            )
            summary["course_progress"] = cur.fetchall()
        summary["period"] = {"key": period, "from": start, "to": today}
        return summary

    @router.post("/assistant/sessions")
    def start_assistant_session(
        request: AssistantStart,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        if request.context_id is None and request.draft_id is None:
            raise HTTPException(
                status_code=422,
                detail="Avatar uchun markaz yoki onboarding qoralamasi kerak",
            )
        with database() as (_, cur):
            ensure_schema(cur)
            if request.context_id is not None:
                roles = require_roles(cur, request.context_id, user_id, VIEW_ROLES)
                if "assistant.use" not in permissions_for(roles):
                    raise HTTPException(status_code=403, detail="AI avatar ruxsati yo'q")
            if request.draft_id is not None:
                cur.execute(
                    """SELECT id,confirmed_context_id FROM center_setup_drafts
                       WHERE id=%s AND creator_user_id=%s
                         AND status IN ('draft','confirmed')""",
                    (request.draft_id, user_id),
                )
                draft = cur.fetchone()
                if not draft:
                    raise HTTPException(status_code=404, detail="Qoralama topilmadi")
                if (
                    request.context_id is not None
                    and draft["confirmed_context_id"] is not None
                    and int(draft["confirmed_context_id"]) != request.context_id
                ):
                    raise HTTPException(status_code=409, detail="Qoralama boshqa markazniki")
            cur.execute(
                """INSERT INTO center_assistant_sessions(
                     user_id,context_id,draft_id,workflow_key,current_step,
                     avatar_enabled,speech_enabled,avatar_variant,state_payload
                   ) VALUES(%s,%s,%s,%s,'welcome',%s,%s,%s,%s::jsonb)
                   RETURNING *""",
                (
                    user_id, request.context_id, request.draft_id,
                    request.workflow_key, request.avatar_enabled,
                    request.speech_enabled, request.avatar_variant,
                    json.dumps(
                        {
                            "mode": "guided_ui_only",
                            "can_publish": False,
                            "can_record_payment": False,
                            "can_assign_roles": False,
                            "can_score": False,
                        }
                    ),
                ),
            )
            session = cur.fetchone()
            audit(
                cur, request.context_id, user_id, "assistant.session.start",
                "assistant_session", session["id"],
            )
        return {
            "session": session,
            "allowed_actions": sorted(ASSISTANT_ACTIONS),
            "autonomous_mutations": False,
        }

    @router.post("/assistant/sessions/{session_id}/actions")
    def assistant_action(
        session_id: int, request: AssistantAction,
        user_id: int = Depends(authenticated_user),
    ) -> dict[str, Any]:
        if request.action_id not in ASSISTANT_ACTIONS:
            raise HTTPException(
                status_code=422,
                detail="Avatar bu amalni bajara olmaydi; inson tasdig'i kerak",
            )
        require_encoded_size(
            request.payload, maximum=20_000, label="Avatar amali"
        )
        encoded = json.dumps(request.payload, ensure_ascii=False, default=str)
        with database() as (_, cur):
            ensure_schema(cur)
            cur.execute(
                """SELECT * FROM center_assistant_sessions
                   WHERE id=%s AND user_id=%s AND state<>'cancelled'
                   FOR UPDATE""",
                (session_id, user_id),
            )
            session = cur.fetchone()
            if not session:
                raise HTTPException(status_code=404, detail="Avatar sessiyasi topilmadi")
            if session["context_id"] is not None:
                roles = require_roles(
                    cur, int(session["context_id"]), user_id, VIEW_ROLES
                )
                if "assistant.use" not in permissions_for(roles):
                    raise HTTPException(status_code=403, detail="AI avatar ruxsati yo'q")
            cur.execute(
                """SELECT COALESCE(MAX(sequence_no),0)+1 next_no
                   FROM center_assistant_actions WHERE session_id=%s""",
                (session_id,),
            )
            sequence = int(cur.fetchone()["next_no"])
            state = session["state"]
            completed_at = None
            if request.action_id == "PAUSE":
                state = "paused"
            elif request.action_id == "MINIMIZE":
                state = "minimized"
            elif request.action_id in {"RESUME", "RESTORE"}:
                state = "active"
            elif request.action_id == "COMPLETE_TOUR":
                state = "completed"
                completed_at = datetime.now(timezone.utc)
            current_step = (
                str(
                    request.payload.get("step_key")
                    or request.payload.get("target_key")
                    or session["current_step"]
                )[:120]
            )
            result_payload = {
                "accepted": True,
                "ui_only": True,
                "action_id": request.action_id,
                "requires_human_for_commit": True,
            }
            if request.action_id == "SET_DRAFT_VALUE":
                result_payload["draft_mutated"] = False
                result_payload["message"] = (
                    "Avatar maydonni interfeysda ko'rsatadi; server qoralamasini "
                    "faqat alohida PATCH so'rovi saqlaydi."
                )
            cur.execute(
                """INSERT INTO center_assistant_actions(
                     session_id,sequence_no,action_id,ui_anchor,reversible,
                     input_payload,result_payload
                   ) VALUES(%s,%s,%s,%s,%s,%s::jsonb,%s::jsonb)
                   RETURNING *""",
                (
                    session_id, sequence, request.action_id, request.ui_anchor,
                    request.action_id in REVERSIBLE_ACTIONS, encoded,
                    json.dumps(result_payload, ensure_ascii=False),
                ),
            )
            action = cur.fetchone()
            cur.execute(
                """UPDATE center_assistant_sessions SET state=%s,current_step=%s,
                     completed_at=%s,updated_at=NOW()
                   WHERE id=%s RETURNING *""",
                (state, current_step, completed_at, session_id),
            )
            updated = cur.fetchone()
            audit(
                cur, session["context_id"], user_id, "assistant.action",
                "assistant_session", session_id,
                {"action_id": request.action_id, "sequence": sequence},
            )
        return {
            "session": updated,
            "action": action,
            "result": result_payload,
            "autonomous_mutations": False,
        }

    return router
